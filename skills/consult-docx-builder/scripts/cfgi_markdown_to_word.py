#!/usr/bin/env python3
"""
cfgi_markdown_to_word.py — Convert finalized consult-drafter process-documentation
Markdown into a CFGI-styled Word (.docx) deliverable.

House style (green CFGI palette) is applied unconditionally; this converter is
opinionated by design and takes no JSON style config. It is aligned to the output
of the consult-drafter skill:

  - Canonical heading hierarchy is preserved straight-through
    (#=H1, ##=H2, ###=H3, ####=H4). The first H1 is treated as the document
    title and lifted to the cover page.
  - The "Document Profile" table is lifted onto the cover as a summary card and
    suppressed inline (no duplicate). Disable with --no-cover.
  - Callouts are matched by the drafter's labels: CONTROL, VALIDATION REQUIRED,
    PAIN POINT, IMPROVEMENT OPPORTUNITY, SCREENSHOT PLACEHOLDER — each shaded a
    distinct color. Screenshot placeholders always render as text callouts;
    images are never inserted.
  - Tables are auto-styled by kind: field, control, pain, gap, screenshot,
    standard.

CLI:
  python cfgi_markdown_to_word.py input.md
  python cfgi_markdown_to_word.py input.md -o output.docx
  python cfgi_markdown_to_word.py input.md --include-toc
  python cfgi_markdown_to_word.py input.md --landscape
  python cfgi_markdown_to_word.py input.md --no-cover
"""
from __future__ import annotations
import argparse
import html
import re
from dataclasses import dataclass, field as dc_field
from html.parser import HTMLParser
from pathlib import Path
from typing import Dict, List, Optional, Tuple

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

# --------------------------------------------------------------------------- #
# Style constants (CFGI green house style)
# --------------------------------------------------------------------------- #
PALETTE = {
    "dark_green": "#0F4A22", "green": "#1A7A3D", "gray_text": "#595959",
    "black": "#000000", "white": "#FFFFFF",
    "light_red": "#FBEBEB", "red": "#C00000",
    "light_green": "#F3F8F4", "light_yellow": "#FCF7CC",
    "label_gray": "#F2F2F2", "border_gray": "#BFBFBF",
    # Improvement-opportunity callouts/tables get their own soft blue so they
    # read distinctly from CONTROL (green), PAIN POINT (red), and GAP (yellow).
    "io_fill": "#EAF1F8", "io_text": "#1F5C99",
}
BODY_FONT = "Calibri"
BODY_SIZE = 10
SMALL_SIZE = 9

# Standalone section-label lines that should render as a subtitle rather than a
# heading. The consult-drafter A–H subsections are ### headings and handled by
# the heading path; these cover free-standing labels a draft might still use.
SECTION_LABELS = {
    "Purpose", "Scope", "Procedure Header",
    "A. Process Overview", "B. Quick Reference", "C. Pre-Requisites",
    "D. Inputs", "E. Step-by-Step Procedure", "F. Key Controls",
    "G. Outputs", "H. Known Issues / Improvement Notes",
}
# Sections lifted to the cover page and suppressed inline when a cover is built.
COVER_SECTIONS = {"document profile"}


@dataclass
class TableBlock:
    header: List[str]
    rows: List[List[str]]
    context: str = ""


@dataclass
class Seg:
    text: str
    bold: bool = False
    italic: bool = False
    code: bool = False


# --------------------------------------------------------------------------- #
# Low-level docx / OOXML helpers
# --------------------------------------------------------------------------- #
def rgb(x: str) -> RGBColor:
    x = x.strip().lstrip("#")
    return RGBColor(int(x[:2], 16), int(x[2:4], 16), int(x[4:], 16))


def nohash(x: str) -> str:
    return x.strip().lstrip("#").upper()


def cell_shading(cell, fill: str) -> None:
    tcPr = cell._tc.get_or_add_tcPr()
    shd = tcPr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tcPr.append(shd)
    shd.set(qn("w:fill"), nohash(fill))


def cell_borders(cell, color: str = PALETTE["border_gray"], size: str = "4") -> None:
    tcPr = cell._tc.get_or_add_tcPr()
    b = tcPr.first_child_found_in("w:tcBorders")
    if b is None:
        b = OxmlElement("w:tcBorders")
        tcPr.append(b)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        e = b.find(qn("w:" + edge))
        if e is None:
            e = OxmlElement("w:" + edge)
            b.append(e)
        e.set(qn("w:val"), "single")
        e.set(qn("w:sz"), size)
        e.set(qn("w:space"), "0")
        e.set(qn("w:color"), nohash(color))


def cell_margins(cell, twips: int = 80) -> None:
    tcPr = cell._tc.get_or_add_tcPr()
    mar = tcPr.first_child_found_in("w:tcMar")
    if mar is None:
        mar = OxmlElement("w:tcMar")
        tcPr.append(mar)
    for name in ("top", "start", "bottom", "end"):
        n = mar.find(qn("w:" + name))
        if n is None:
            n = OxmlElement("w:" + name)
            mar.append(n)
        n.set(qn("w:w"), str(twips))
        n.set(qn("w:type"), "dxa")


def para_bottom_border(p, color: str = PALETTE["green"], size: str = "4", space: str = "2") -> None:
    pPr = p._p.get_or_add_pPr()
    b = pPr.find(qn("w:pBdr"))
    if b is None:
        b = OxmlElement("w:pBdr")
        pPr.append(b)
    bot = b.find(qn("w:bottom"))
    if bot is None:
        bot = OxmlElement("w:bottom")
        b.append(bot)
    bot.set(qn("w:val"), "single")
    bot.set(qn("w:sz"), size)
    bot.set(qn("w:space"), space)
    bot.set(qn("w:color"), nohash(color))


def repeat_header(row) -> None:
    trPr = row._tr.get_or_add_trPr()
    h = OxmlElement("w:tblHeader")
    h.set(qn("w:val"), "true")
    trPr.append(h)


def add_page_num(p) -> None:
    r = p.add_run()
    b = OxmlElement("w:fldChar")
    b.set(qn("w:fldCharType"), "begin")
    i = OxmlElement("w:instrText")
    i.set(qn("xml:space"), "preserve")
    i.text = "PAGE"
    e = OxmlElement("w:fldChar")
    e.set(qn("w:fldCharType"), "end")
    r._r.append(b)
    r._r.append(i)
    r._r.append(e)


def add_toc(p) -> None:
    r = p.add_run()
    b = OxmlElement("w:fldChar")
    b.set(qn("w:fldCharType"), "begin")
    i = OxmlElement("w:instrText")
    i.set(qn("xml:space"), "preserve")
    i.text = 'TOC \\o "1-3" \\h \\z \\u'
    s = OxmlElement("w:fldChar")
    s.set(qn("w:fldCharType"), "separate")
    e = OxmlElement("w:fldChar")
    e.set(qn("w:fldCharType"), "end")
    for el in (b, i, s, e):
        r._r.append(el)


def add_hr(doc) -> None:
    """Render a Markdown horizontal rule as a subtle full-width hairline."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(6)
    para_bottom_border(p, PALETTE["border_gray"], size="4", space="1")


# --------------------------------------------------------------------------- #
# Styles
# --------------------------------------------------------------------------- #
def set_style(st, size=None, color=None, bold=None, italic=None) -> None:
    st.font.name = BODY_FONT
    st._element.rPr.rFonts.set(qn("w:eastAsia"), BODY_FONT)
    if size:
        st.font.size = Pt(size)
    if color:
        st.font.color.rgb = rgb(color)
    if bold is not None:
        st.font.bold = bold
    if italic is not None:
        st.font.italic = italic


def define_styles(doc) -> None:
    set_style(doc.styles["Normal"], BODY_SIZE, PALETTE["black"])
    doc.styles["Normal"].paragraph_format.space_after = Pt(6)
    specs = {
        "Title": (40, PALETTE["dark_green"], True, False),
        "Heading 1": (14, PALETTE["dark_green"], False, False),
        "Heading 2": (13, PALETTE["green"], True, False),
        "Heading 3": (12, PALETTE["green"], False, True),
        "Heading 4": (11, PALETTE["black"], True, False),
    }
    for name, (sz, col, b, it) in specs.items():
        st = doc.styles[name]
        set_style(st, sz, col, b, it)
        st.paragraph_format.space_before = Pt(8)
        st.paragraph_format.space_after = Pt(5)
        st.paragraph_format.keep_with_next = True
    existing = [x.name for x in doc.styles]
    for name, sz, col, b, it in [
        ("Section Subtitle", 12, PALETTE["green"], False, True),
        ("Caption Text", 9, PALETTE["gray_text"], False, True),
    ]:
        st = doc.styles[name] if name in existing else doc.styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
        set_style(st, sz, col, b, it)
        st.paragraph_format.space_after = Pt(3)
    for name in ("List Bullet", "List Number", "List Paragraph"):
        try:
            set_style(doc.styles[name], BODY_SIZE, PALETTE["black"])
        except KeyError:
            pass


def configure(doc, landscape: bool = False) -> None:
    s = doc.sections[0]
    s.top_margin = s.bottom_margin = s.left_margin = s.right_margin = Inches(1)
    if landscape:
        s.orientation = WD_ORIENT.LANDSCAPE
        s.page_width = Inches(11)
        s.page_height = Inches(8.5)
    define_styles(doc)
    f = s.footer
    p = f.paragraphs[0] if f.paragraphs else f.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Page ")
    r.font.name = BODY_FONT
    r.font.size = Pt(8)
    r.font.color.rgb = rgb(PALETTE["gray_text"])
    add_page_num(p)


# --------------------------------------------------------------------------- #
# Inline text / segments
# --------------------------------------------------------------------------- #
def clean(t: str) -> str:
    t = html.unescape(t)
    t = re.sub(r"<br\s*/?>", "\n", t, flags=re.I)
    t = re.sub(r"</?span[^>]*>", "", t, flags=re.I)
    t = re.sub(r"<[^>]+>", "", t)
    return t.replace("\\_", "_").replace("\\*", "*").replace("\\|", "|").strip()


def strip_links(t: str) -> str:
    return re.sub(r"\[([^\]]+)\]\([^)]+\)", r"\1", t)


_BOLD_RE = re.compile(r"\*\*([^*]+?)\*\*")
# Underscore italics only open at a word boundary and close before one, so
# snake_case identifiers are never treated as emphasis.
_ITALIC_RE = re.compile(r"\*([^*]+?)\*|(?<!\w)_([A-Za-z](?:`[^`]*`|[^_])*)_(?!\w)")
_CODE_RE = re.compile(r"`([^`]+?)`")


def _emit_code(t: str, out: List[Seg], bold: bool = False, italic: bool = False) -> None:
    for k, part in enumerate(_CODE_RE.split(t)):
        if part:
            out.append(Seg(part, bold=bold, italic=italic, code=(k % 2 == 1)))


def _emit_italic(t: str, out: List[Seg], bold: bool = False) -> None:
    pos = 0
    for m in _ITALIC_RE.finditer(t):
        if m.start() > pos:
            _emit_code(t[pos:m.start()], out, bold)
        _emit_code(m.group(1) or m.group(2), out, bold, italic=True)
        pos = m.end()
    if pos < len(t):
        _emit_code(t[pos:], out, bold)


def segs(text: str) -> List[Seg]:
    # Nested inline pass: bold outermost, then italics, then code spans, so
    # markers inside a bold span (e.g. **Notes on `A?` rows:**) still resolve.
    text = strip_links(clean(text))
    out: List[Seg] = []
    pos = 0
    for m in _BOLD_RE.finditer(text):
        if m.start() > pos:
            _emit_italic(text[pos:m.start()], out)
        _emit_italic(m.group(1), out, bold=True)
        pos = m.end()
    if pos < len(text):
        _emit_italic(text[pos:], out)
    return [s for s in out if s.text]


def styled_run(p, sg: Seg, color: str = PALETTE["black"], size: int = BODY_SIZE) -> None:
    r = p.add_run(sg.text)
    r.font.name = BODY_FONT
    r._element.rPr.rFonts.set(qn("w:eastAsia"), r.font.name)
    r.font.size = Pt(SMALL_SIZE if sg.code else size)
    r.font.color.rgb = rgb(color)
    r.bold = sg.bold
    r.italic = sg.italic


# Direct run-level formatting for section headings. Heading runs previously
# received body-default direct formatting (10pt / black / bold=False /
# italic=False) from styled_run, which overrides the Heading 2/3 paragraph
# STYLES in Word. Direct-format the runs to match the style instead.
HEADING_RUN_FMT = {
    "Heading 2": (13, PALETTE["green"], True, False),   # size, color, bold, italic
    "Heading 3": (12, PALETTE["green"], False, True),
}


def para(doc, text: str, style: Optional[str] = None,
         color: str = PALETTE["black"], size: int = BODY_SIZE):
    p = doc.add_paragraph(style=style)
    fmt = HEADING_RUN_FMT.get(style or "")
    for s in segs(text):
        if fmt:
            size, color, s.bold, s.italic = fmt[0], fmt[1], fmt[2], fmt[3]
            s.code = False   # keep heading runs at heading size, not code size
        styled_run(p, s, color, size)
    return p


# --------------------------------------------------------------------------- #
# Table parsing
# --------------------------------------------------------------------------- #
class TableParser(HTMLParser):
    def __init__(self):
        super().__init__(convert_charrefs=True)
        self.rows: List[List[Tuple[str, bool]]] = []
        self.row: Optional[list] = None
        self.cell: Optional[list] = None
        self.is_header = False
        self.in_cell = False

    def handle_starttag(self, tag, attrs):
        tag = tag.lower()
        if tag == "tr":
            self.row = []
        elif tag in ("td", "th"):
            self.cell = []
            self.is_header = tag == "th"
            self.in_cell = True
        elif tag == "br" and self.in_cell:
            self.cell.append("\n")

    def handle_endtag(self, tag):
        tag = tag.lower()
        if tag in ("td", "th") and self.row is not None and self.cell is not None:
            txt = re.sub(r"[ \t]+", " ", " ".join("".join(self.cell).splitlines()).strip())
            self.row.append((txt, self.is_header))
            self.cell = None
            self.in_cell = False
        elif tag == "tr" and self.row is not None:
            if self.row:
                self.rows.append(self.row)
            self.row = None

    def handle_data(self, data):
        if self.in_cell and self.cell is not None:
            self.cell.append(data)


def html_table(block: str, context: str = "") -> TableBlock:
    p = TableParser()
    p.feed(block)
    rr = p.rows
    if not rr:
        return TableBlock([], [], context)
    if all(h for _, h in rr[0]):
        header = [x for x, _ in rr[0]]
        rows = [[x for x, _ in r] for r in rr[1:]]
    else:
        header = []
        rows = [[x for x, _ in r] for r in rr]
    n = max([len(header)] + [len(r) for r in rows] + [1])
    if header:
        header += [""] * (n - len(header))
    rows = [r + [""] * (n - len(r)) for r in rows]
    return TableBlock(header, rows, context)


def md_table(lines: List[str], i: int, context: str = "") -> Tuple[Optional[TableBlock], int]:
    if i + 1 >= len(lines) or "|" not in lines[i]:
        return None, i
    sep = [c.strip() for c in lines[i + 1].strip().strip("|").split("|")]
    if not sep or not all(re.fullmatch(r":?-{3,}:?", c or "") for c in sep):
        return None, i

    def split(line: str) -> List[str]:
        return [c.strip() for c in line.strip().strip("|").split("|")]

    header = split(lines[i])
    rows: List[List[str]] = []
    i += 2
    while i < len(lines) and lines[i].strip() and "|" in lines[i]:
        rows.append(split(lines[i]))
        i += 1
    n = len(header)
    rows = [r + [""] * (n - len(r)) for r in rows]
    return TableBlock(header, rows, context), i


# --------------------------------------------------------------------------- #
# Table rendering
# --------------------------------------------------------------------------- #
def table_kind(tb: TableBlock, ctx: str) -> str:
    # Kind is driven by the table's own header, not the surrounding context, so
    # adjacent appendices (e.g. a Gap log next to a Screenshot index) can't bleed
    # their keywords into each other. Context only supplies the summary card.
    h = " | ".join(tb.header).lower()
    c = ctx.lower()
    if len(tb.header) == 2 and {x.lower() for x in tb.header} & {"field", "value", "detail"}:
        return "field"
    if "summary card" in c:
        return "field"
    if re.search(r"\bsc id\b", h) or "screenshot" in h:
        return "screenshot"
    if re.search(r"\bgap\b", h):
        return "gap"
    if "control" in h:
        return "control"
    # Pain-point-only tables get red; a mixed risks/pain/opportunity register
    # (Appendix A) has no "pain" in its header, so it stays standard and its
    # improvement rows aren't miscolored as pain.
    if "pain" in h or "known issue" in h:
        return "pain"
    return "standard"


def render_cell(cell, value: str, fill: str, color: str, bold: bool = False,
                border: str = PALETTE["border_gray"]) -> None:
    cell.text = ""
    cell_shading(cell, fill)
    cell_borders(cell, border)
    cell_margins(cell, 80)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.0
    if re.fullmatch(r"[A-Z]{1,4}|v?\d+(\.\d+)?|Yes|No|TBD|[-\u2014]", value.strip(), flags=re.I):
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for s in segs(value):
        if bold:
            s.bold = True
        styled_run(p, s, color, BODY_SIZE)


def _column_widths(tb: TableBlock, n: int, usable) -> List[int]:
    """Content-proportional column widths (EMU), min 0.6\" per column.

    Weight per column = max per-cell score, where a cell scores
    min(len(text), 60) so one huge cell can't take the whole page."""

    def score(txt: str) -> int:
        return min(len(strip_links(clean(txt))), 60)

    weights: List[float] = []
    for j in range(n):
        col = ([tb.header[j]] if tb.header and j < len(tb.header) else [])
        col += [r[j] for r in tb.rows if j < len(r)]
        weights.append(float(max([score(c) for c in col] + [4])))
    min_w = int(Inches(0.6))
    total = sum(weights)
    widths = [max(min_w, int(usable * w / total)) for w in weights]
    # Rescale the above-minimum columns so the table fills the usable width.
    flex = [j for j in range(n) if widths[j] > min_w]
    spare = usable - min_w * (n - len(flex))
    if flex and spare > min_w * len(flex):
        fsum = sum(widths[j] for j in flex)
        for j in flex:
            widths[j] = max(min_w, int(spare * widths[j] / fsum))
    widths[widths.index(max(widths))] += usable - sum(widths)  # rounding remainder
    return widths


def _fixed_layout(t) -> None:
    tblPr = t._tbl.tblPr
    layout = tblPr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tblPr.append(layout)
    layout.set(qn("w:type"), "fixed")


def add_table(doc, tb: TableBlock, ctx: str = "") -> None:
    if not tb.header and not tb.rows:
        return
    n = max([len(tb.header)] + [len(r) for r in tb.rows] + [1])
    kind = table_kind(tb, ctx)
    nrows = len(tb.rows) + (1 if tb.header else 0)
    t = doc.add_table(rows=nrows, cols=n)
    t.style = "Normal Table"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Content-proportional fixed column widths: set tblGrid (column.width) and
    # tcW on every cell, with autofit off + fixed layout so Word honors them.
    t.autofit = False
    _fixed_layout(t)
    sec = doc.sections[-1]
    usable = int(sec.page_width - sec.left_margin - sec.right_margin)
    widths = _column_widths(tb, n, usable)
    for j, w in enumerate(widths):
        t.columns[j].width = w
        for cell in t.columns[j].cells:
            cell.width = w
    off = 0
    if tb.header:
        for j, v in enumerate(tb.header):
            render_cell(t.cell(0, j), v, PALETTE["green"], PALETTE["white"], True, PALETTE["green"])
        repeat_header(t.rows[0])
        off = 1
    for i, r in enumerate(tb.rows):
        for j in range(n):
            v = r[j] if j < len(r) else ""
            fill, color, bold, border = PALETTE["white"], PALETTE["black"], False, PALETTE["border_gray"]
            if kind == "field" and j == 0:
                fill, color, bold, border = PALETTE["green"], PALETTE["white"], True, PALETTE["green"]
            elif kind == "pain":
                fill = PALETTE["light_red"]
                color = PALETTE["red"] if j < 2 else PALETTE["black"]
            render_cell(t.cell(i + off, j), v, fill, color, bold, border)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)


# --------------------------------------------------------------------------- #
# Callouts
# --------------------------------------------------------------------------- #
def callout_style(text: str) -> Tuple[str, str]:
    u = text.upper()
    if "PAIN POINT" in u:
        return PALETTE["light_red"], PALETTE["red"]
    if "IMPROVEMENT" in u or "OPPORTUNITY" in u:
        return PALETTE["io_fill"], PALETTE["io_text"]
    if "VALIDATION" in u or "GAP" in u or "UNKNOWN" in u or "DECISION REQUIRED" in u:
        return PALETTE["light_yellow"], PALETTE["dark_green"]
    if "SCREENSHOT" in u or "SC-" in u:
        return PALETTE["label_gray"], PALETTE["gray_text"]
    if "CONTROL" in u or "FUTURE" in u:
        return PALETTE["light_green"], PALETTE["dark_green"]
    return PALETTE["light_green"], PALETTE["dark_green"]


def enable_track_changes(doc) -> None:
    """Turn tracked changes ON by default (settings.xml <w:trackChanges/>).

    Reviewers can still toggle it off, but a kit doc opens recording — the
    review contract without anyone having to remember the ribbon."""
    settings = doc.settings.element
    if settings.find(qn("w:trackChanges")) is None:
        settings.append(OxmlElement("w:trackChanges"))


def add_image(doc, path: Path, caption: str = "") -> None:
    """Embed an image scaled to the usable width, with a small italic caption
    paragraph below it (final-mode screenshot rendering)."""
    sec = doc.sections[-1]
    usable = sec.page_width - sec.left_margin - sec.right_margin
    p = doc.add_paragraph()
    run = p.add_run()
    shape = run.add_picture(str(path))
    if shape.width > usable:
        shape.height = int(shape.height * usable / shape.width)
        shape.width = usable
    if caption:
        cp = doc.add_paragraph()
        cp.paragraph_format.space_before = Pt(2)
        cp.paragraph_format.space_after = Pt(10)
        r = cp.add_run(caption)
        r.italic = True
        r.font.name = BODY_FONT
        r.font.size = Pt(SMALL_SIZE)
        r.font.color.rgb = rgb(PALETTE["gray_text"])


def add_callout(doc, text: str) -> List:
    # Strip blockquote markers up front; the callout box carries the emphasis.
    lines = [re.sub(r"^\s*>\s?", "", ln) for ln in text.split("\n")]
    fill, col = callout_style("\n".join(lines))
    t = doc.add_table(rows=1, cols=1)
    c = t.cell(0, 0)
    c.text = ""
    c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    cell_shading(c, fill)
    cell_borders(c, fill)
    cell_margins(c, 120)
    # Group lines: "- " lines become sub-bullets; consecutive plain lines wrap
    # into a single paragraph; blank quote lines separate paragraphs.
    blocks: List[Tuple[str, str]] = []
    buf: List[str] = []

    def flush() -> None:
        if buf:
            blocks.append(("para", " ".join(buf)))
            buf.clear()

    for ln in lines:
        s = ln.strip()
        if not s:
            flush()
        elif re.match(r"^[-*+]\s+", s):
            flush()
            blocks.append(("bullet", re.sub(r"^[-*+]\s+", "", s)))
        else:
            buf.append(s)
    flush()
    out_paras = []
    for k, (kind, txt) in enumerate(blocks):
        p = c.paragraphs[0] if k == 0 else c.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        if kind == "bullet":
            p.paragraph_format.left_indent = Inches(0.39)
            p.paragraph_format.first_line_indent = Inches(-0.19)
            styled_run(p, Seg("•  ", italic=True), col, BODY_SIZE)
        for s in segs(txt):
            s.italic = True
            styled_run(p, s, col, BODY_SIZE)
        out_paras.append(p)
    doc.add_paragraph()
    return out_paras


CALLOUT_PREFIXES = (
    "PAIN POINT", "CONTROL", "VALIDATION REQUIRED", "IMPROVEMENT OPPORTUNITY",
    "SCREENSHOT PLACEHOLDER", "CALLOUT", "WARNING", "NOTE", "DECISION REQUIRED",
)


def is_callout(line: str) -> bool:
    if line.strip().startswith(">"):
        return True
    u = line.strip().upper().strip("*_ ")
    # Prefix must end at a word boundary so prose like "Controller ..." is not
    # mistaken for a CONTROL callout.
    return any(
        u.startswith(p) and (len(u) == len(p) or not u[len(p)].isalpha())
        for p in CALLOUT_PREFIXES
    ) or "\u26a0" in line


def _is_list_item(line: str) -> bool:
    return bool(re.match(r"^\s*[-*+]\s+", line) or re.match(r"^\s*\d+[.)]\s+", line))


def _is_structural(line: str) -> bool:
    """Lines that must start their own block and never be joined into a
    preceding paragraph or list item: headings, table rows, blockquotes /
    callouts, fences, horizontal rules, images, HTML tables."""
    st = line.strip()
    return bool(
        re.match(r"^#{1,6}\s+", st)
        or st.startswith("|") or "|" in st and st.count("|") >= 2
        or st.startswith(">")
        or st.startswith("```") or st.startswith("~~~")
        or re.fullmatch(r"(\*\s*){3,}|(-\s*){3,}|(_\s*){3,}", st)
        or re.match(r"^!\[", st)
        or re.match(r"^<table\b", st, flags=re.I)
        or is_callout(st)
    )


# --------------------------------------------------------------------------- #
# Cover page (built from the Document Profile table)
# --------------------------------------------------------------------------- #
def extract_cover_data(lines: List[str]) -> Tuple[str, str, List[List[str]]]:
    """Return (title, subtitle, profile_rows) pulled from the top of the doc."""
    title, subtitle = "", ""
    profile_rows: List[List[str]] = []
    # Title = first H1; subtitle = first italic line after it.
    for idx, ln in enumerate(lines):
        m = re.match(r"^#\s+(.*)$", ln.strip())
        if m:
            title = clean(strip_links(m.group(1)))
            for nxt in lines[idx + 1:idx + 4]:
                s = nxt.strip()
                if not s:
                    continue
                if s.startswith("_") and s.endswith("_"):
                    subtitle = clean(s).strip("_* ")
                break
            break
    # Profile rows = first table under a "Document Profile" heading.
    for idx, ln in enumerate(lines):
        m = re.match(r"^#{1,6}\s+(.*)$", ln.strip())
        if m and clean(m.group(1)).lower() in COVER_SECTIONS:
            tb, _ = md_table(lines, _next_nonblank(lines, idx + 1))
            if tb:
                if tb.header:
                    profile_rows.append(tb.header)
                profile_rows.extend(tb.rows)
            break
    return title, subtitle, profile_rows


def _next_nonblank(lines: List[str], i: int) -> int:
    while i < len(lines) and not lines[i].strip():
        i += 1
    return i


def build_cover(doc, title: str, subtitle: str, profile_rows: List[List[str]]) -> None:
    title = title or "Process Documentation"
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("_" * 86)
    r.font.color.rgb = rgb(PALETTE["green"])
    r.font.size = Pt(10)

    p = doc.add_paragraph()
    r = p.add_run(title)
    r.bold = True
    r.font.name = BODY_FONT
    r.font.size = Pt(36 if len(title) < 60 else 28)
    r.font.color.rgb = rgb(PALETTE["dark_green"])

    # Entity line from the profile's Client/Organization value, if meaningful.
    org = ""
    for row in profile_rows:
        if len(row) >= 2 and re.search(r"client|organi", row[0], flags=re.I):
            if row[1].strip() and row[1].strip().upper() != "TBD":
                org = row[1].strip()
            break
    for line in filter(None, [org, subtitle]):
        p = doc.add_paragraph()
        r = p.add_run(line)
        r.italic = True
        r.font.size = Pt(18)
        r.font.color.rgb = rgb(PALETTE["green"])

    doc.add_paragraph()
    if profile_rows:
        # Render the profile as a 2-col field/value summary card.
        header = ["Field", "Value"]
        body = [row for row in profile_rows if not (len(row) == 2 and row[0].lower() == "field")]
        add_table(doc, TableBlock(header, [r[:2] + [""] * (2 - len(r)) for r in body], "summary card"), "summary card")
    doc.add_page_break()


# --------------------------------------------------------------------------- #
# Main conversion
# --------------------------------------------------------------------------- #
def _emit_toc(doc) -> None:
    para(doc, "Table of Contents", "Heading 1")
    p = doc.add_paragraph()
    add_toc(p)
    doc.add_page_break()


def render_body(doc, lines: List[str], do_cover: bool = True, prov=None) -> None:
    """Render assembled Markdown body lines into `doc`.

    Shared by the single-file `convert` and the pre-assembled folder
    `convert_assembled` paths. `do_cover` only governs inline suppression of
    the title H1 and lifted cover sections; the cover itself is built by the
    caller.

    `prov` (optional) is a provenance collector with a
    `mark(paras, i_start, i_end, kind)` method — called with the emitted
    paragraph(s) and the half-open source-line span [i_start, i_end) that
    produced them, for paragraphs, list items, and callout blocks (the units
    the M10 review-apply pipeline can splice back). Headings and tables are
    deliberately not marked: heading edits are manifest identity and table
    edits route to notes/triage."""
    ctx: List[str] = []
    i = 0
    title_skipped = not do_cover      # if no cover, don't swallow the first H1
    subtitle_skipped = not do_cover
    skip_until_level: Optional[int] = None

    while i < len(lines):
        line = lines[i]
        st = line.strip()
        if not st:
            i += 1
            continue

        # Suppress a section that was lifted to the cover (e.g. Document Profile).
        if skip_until_level is not None:
            hm = re.match(r"^(#{1,6})\s+", st)
            if hm and len(hm.group(1)) <= skip_until_level:
                skip_until_level = None      # fall through to render this heading
            else:
                i += 1
                continue

        # Horizontal rule.
        if re.fullmatch(r"(\*\s*){3,}|(-\s*){3,}|(_\s*){3,}", st):
            add_hr(doc)
            i += 1
            continue

        # HTML table.
        if re.match(r"^<table\b", st, flags=re.I):
            block = [line]
            i += 1
            while i < len(lines):
                block.append(lines[i])
                if re.search(r"</table>", lines[i], flags=re.I):
                    i += 1
                    break
                i += 1
            add_table(doc, html_table("\n".join(block), " | ".join(ctx[-3:])), " | ".join(ctx[-3:]))
            continue

        # Markdown table.
        tb, ni = md_table(lines, i, " | ".join(ctx[-3:]))
        if tb:
            add_table(doc, tb, " | ".join(ctx[-3:]))
            i = ni
            continue

        # Heading.
        m = re.match(r"^(#{1,6})\s+(.*)$", st)
        if m:
            level = len(m.group(1))
            txt = clean(strip_links(m.group(2)))
            # First H1 is the document title -> already on the cover.
            if do_cover and not title_skipped and level == 1:
                title_skipped = True
                i += 1
                continue
            # Lift-to-cover sections are suppressed inline.
            if do_cover and txt.lower() in COVER_SECTIONS:
                skip_until_level = level
                i += 1
                continue
            ctx.append(txt)
            style = f"Heading {min(level, 4)}"      # straight-through: #->H1 ... ####->H4
            p = para(doc, txt, style)
            # H1->H2 weight remap: under the flat-H2 template the section
            # weight formerly carried by H1 moves to H2, so the green rule is
            # drawn under both. (Single-file docs still have one inline H1;
            # folder docs hold their single H1/title on the cover and start
            # every section at H2.)
            if style in ("Heading 1", "Heading 2"):
                para_bottom_border(p)
            i += 1
            continue

        # Free-standing section label -> subtitle.
        label = clean(st).strip("*_ ")
        if label in SECTION_LABELS or re.fullmatch(r"[A-H]\.\s+.+", label):
            para(doc, re.sub(r"^[A-H]\.\s+", "", label), "Section Subtitle", PALETTE["green"])
            i += 1
            continue

        # Skip the italic subtitle line directly under the (lifted) title.
        if do_cover and title_skipped and not subtitle_skipped and st.startswith("_") and st.endswith("_"):
            subtitle_skipped = True
            i += 1
            continue

        # Blockquote / labeled callout (screenshots always stay text callouts).
        if is_callout(st):
            i0 = i
            block = [st]
            i += 1
            while i < len(lines) and lines[i].strip().startswith(">"):
                block.append(lines[i].strip())
                i += 1
            paras = add_callout(doc, "\n".join(block))
            if prov is not None:
                prov.mark(paras, i0, i, "callout")
            continue

        # Image links: embed the file (scaled, italic caption below) when it
        # exists on disk — the final-mode screenshot path. A dangling path
        # renders as a screenshot placeholder callout, never a broken image.
        im = re.match(r"^!\[([^\]]*)\]\(([^)]+)\)", st)
        if im:
            cap = im.group(1).strip()
            img = Path(im.group(2).strip())
            if img.is_file():
                add_image(doc, img, cap)
            else:
                add_callout(doc, "SCREENSHOT PLACEHOLDER: " + (cap or "insert screenshot"))
            i += 1
            continue

        # Lists. A list item absorbs following indented or plain continuation
        # lines (hard-wrapped prose) into the same bullet paragraph until a
        # blank line, a new list item, or a structural line.
        lm = re.match(r"^\s*[-*+]\s+", line) or re.match(r"^\s*\d+[.)]\s+", line)
        if lm:
            i0 = i
            style = "List Bullet" if re.match(r"^\s*[-*+]\s+", line) else "List Number"
            buf = [line[lm.end():].strip()]
            i += 1
            while i < len(lines):
                nxt = lines[i]
                ns = nxt.strip()
                if not ns or _is_structural(nxt) or _is_list_item(nxt):
                    break
                buf.append(ns)
                i += 1
            p = para(doc, " ".join(buf), style)
            p.paragraph_format.left_indent = Inches(0.25)
            p.paragraph_format.first_line_indent = Inches(-0.15)
            if prov is not None:
                prov.mark([p], i0, i, "bullet")
            continue

        # Plain paragraph. Consecutive plain lines (hard-wrapped prose) join
        # into one paragraph; a blank line or structural line ends it.
        i0 = i
        buf = [st]
        i += 1
        while i < len(lines):
            nxt = lines[i]
            ns = nxt.strip()
            if (not ns or _is_structural(nxt) or _is_list_item(nxt)
                    or (do_cover and title_skipped and not subtitle_skipped
                        and ns.startswith("_") and ns.endswith("_"))):
                break
            lbl = clean(ns).strip("*_ ")
            if lbl in SECTION_LABELS or re.fullmatch(r"[A-H]\.\s+.+", lbl):
                break
            buf.append(ns)
            i += 1
        p = para(doc, " ".join(buf))
        if prov is not None:
            prov.mark([p], i0, i, "para")


def convert(inp: Path, out: Path, include_toc: bool = False,
            landscape: bool = False, do_cover: bool = True) -> None:
    text = inp.read_text(encoding="utf-8").replace("\r\n", "\n").replace("\r", "\n")
    lines = text.split("\n")
    doc = Document()
    configure(doc, landscape)
    if do_cover:
        title, subtitle, profile_rows = extract_cover_data(lines)
        build_cover(doc, title, subtitle, profile_rows)
    if include_toc:
        _emit_toc(doc)
    render_body(doc, lines, do_cover)
    doc.save(str(out))


def convert_assembled(body_md: str, out: Path, *, title: str, subtitle: str,
                      profile_md: str = "", include_toc: bool = False,
                      landscape: bool = False, do_cover: bool = True,
                      prov=None, track_changes: bool = False) -> None:
    """Render pre-assembled folder content produced by scripts/render.py.

    Additive hook for the M4 folder path. Title/subtitle come from the
    manifest (no inline H1 is scanned — a folder document has none). When a
    cover is built the Document Profile card is parsed from `profile_md`; the
    caller omits that section from `body_md` so it is not duplicated inline.
    With do_cover False the caller leaves the profile inline in `body_md`."""
    doc = Document()
    configure(doc, landscape)
    if do_cover:
        rows: List[List[str]] = []
        plines = profile_md.split("\n")
        tb, _ = md_table(plines, _next_nonblank(plines, 0)) if profile_md.strip() else (None, 0)
        if tb:
            if tb.header:
                rows.append(tb.header)
            rows.extend(tb.rows)
        build_cover(doc, title, subtitle, rows)
    if include_toc:
        _emit_toc(doc)
    render_body(doc, body_md.split("\n"), do_cover=False, prov=prov)
    if track_changes:
        enable_track_changes(doc)
    doc.save(str(out))


def infer_output(inp: Path, output_arg: Optional[str]) -> Path:
    if output_arg:
        return Path(output_arg)
    return inp.with_name(inp.stem + "_process-doc.docx")


def main(argv=None) -> int:
    ap = argparse.ArgumentParser(description="Convert consult-drafter process-doc Markdown to a CFGI-styled .docx")
    ap.add_argument("input_md")
    ap.add_argument("-o", "--output")
    ap.add_argument("--include-toc", action="store_true", help="Insert a generated Table of Contents field")
    ap.add_argument("--landscape", action="store_true", help="Use landscape orientation")
    ap.add_argument("--no-cover", action="store_true", help="Skip the generated cover page")
    a = ap.parse_args(argv)
    inp = Path(a.input_md)
    if not inp.exists():
        ap.error(f"input file not found: {inp}")
    out = infer_output(inp, a.output)
    convert(inp, out, a.include_toc, a.landscape, not a.no_cover)
    print("Wrote " + str(out))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
