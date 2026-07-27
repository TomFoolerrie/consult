#!/usr/bin/env python3
"""
kits.py — per-procedure (L3) review kit emitter (M9).

    python3 scripts/kits.py <area> [-o {area}/_review/kits]

After a working-mode render, this builds ONE FOLDER PER PROCEDURE — review
travels by topic, and an L3 usually goes to a single person (when several
people are involved they share the same kit; every ask names its contact):

    _review/kits/
      index.md                       kit -> send-to checklist, grouped by person
      <num>_<slug>/
        README.md                    instructions (usable as the email body)
        <num>_<slug>.docx            the procedure render, tracked changes ON
        gaps_<slug>.xlsx             ALL of this procedure's gap rows; the
                                     Contact column names who each row is for
        screenshots_<slug>.docx      its SC items + paste boxes

Contacts are deterministic string-work over data the pipeline already has:
  - procedure contact = the At a Glance card "Preparer" row → roles.yaml
    `people:` → LOWEST org-chart rank (deepest reports_to = closest to the
    work); their manager is listed as the escalation.
  - gap row contact   = the gap's "Owner to confirm" role, same resolution;
    falls back to the procedure contact. A row can name someone other than
    the procedure contact — intended: the kit stays whole and the reader
    forwards that one ask instead of the ask living in another kit.
  - screenshot contact = the procedure contact (SC callouts carry no owner).
A procedure whose preparer resolves to nobody gets a kit with an
"(unassigned)" contact, flagged in index.md for human triage.

Consistency guarantees: kit docs use the SAME display numbers and global
callout display IDs as the full draft (computed from the full manifest), so
"GAP-07" means the same thing in a workbook, a kit doc, and the master copy.

Python 3, stdlib + pyyaml + python-docx.
"""

from __future__ import annotations

import argparse
import json
import re
import shutil
import sys
import uuid
from pathlib import Path

_SCRIPTS_DIR = Path(__file__).resolve().parent
if str(_SCRIPTS_DIR) not in sys.path:
    sys.path.insert(0, str(_SCRIPTS_DIR))

import doc_model  # noqa: E402
import render as render_mod  # noqa: E402
from aggregate import parse_callouts, parse_bullets, split_subsections, _pick  # noqa: E402
from people import People  # noqa: E402
from xlsx_min import write_xlsx  # noqa: E402

from docx import Document  # noqa: E402
from docx.enum.table import WD_ALIGN_VERTICAL  # noqa: E402
from docx.oxml import OxmlElement  # noqa: E402
from docx.oxml.ns import qn  # noqa: E402
from docx.shared import Inches, Pt, RGBColor  # noqa: E402

GREEN = RGBColor(0x1A, 0x7A, 0x3D)
DARK = RGBColor(0x0F, 0x4A, 0x22)
GRAY = RGBColor(0x59, 0x59, 0x59)
FONT = "Calibri"

GAP_HEADER = ["Gap ID", "Procedure", "Question to Confirm", "Nature",
              "Contact", "Escalation", "Answer", "Status",
              "Ref (do not edit)"]
GAP_WIDTHS = [10, 30, 55, 20, 20, 20, 55, 12, 8]

STEP_RE = re.compile(r"^####\s+(.*)$")


# --------------------------------------------------------------------------- #
# extraction
# --------------------------------------------------------------------------- #

def _step_of(raw_text: str) -> dict[str, str]:
    """{callout-id: enclosing '#### Step …' heading} for one fragment."""
    out: dict[str, str] = {}
    current = ""
    for line in raw_text.splitlines():
        m = STEP_RE.match(line)
        if m:
            current = m.group(1).strip()
            continue
        cm = re.match(r"^\s*>\s*\*\*.*?[-–—]\s*([A-Z]+-[A-Z0-9-]+)\s*:\*\*", line)
        if cm:
            out[cm.group(1)] = current
    return out


def collect(folder: Path):
    """Walk the area once; return (procs, gaps, screens).

    procs:   {slug: {slug,title,number,label,file,preparer_text,owner,role}}
    gaps:    [{slug, local, disp, text, nature, owner, escalation, proc_label}]
    screens: [{slug, local, disp, text, step, owner, proc_label}]
    """
    manifest = doc_model.load_manifest(folder)
    numbers = doc_model.display_numbers(manifest)
    disp = doc_model.callout_display_ids(folder)
    ppl = People(folder)

    procs: dict[str, dict] = {}
    gaps: list[dict] = []
    screens: list[dict] = []

    for comp in manifest.get("components", []):
        if comp.get("role") != "procedure":
            continue
        slug = comp["slug"]
        fpath = folder / comp["file"]
        if not fpath.is_file():
            continue
        raw = fpath.read_text(encoding="utf-8")
        subs = split_subsections(raw)
        # M23: the At a Glance card, by SLUG — this was `subs["B"]`, a
        # letter (a render position) standing in for the section's identity.
        preparer = _pick(parse_bullets(subs.get("quick-reference", "")), "preparer")
        owner, role = ppl.contact_for_text(preparer)
        label = f"{numbers.get(slug, '')} {comp.get('heading', '')}".strip()
        procs[slug] = {
            "slug": slug, "title": comp.get("heading", ""),
            "number": numbers.get(slug, ""), "label": label,
            "file": comp["file"], "preparer_text": preparer,
            "owner": owner, "role": role,
        }
        steps = _step_of(raw)
        for c in parse_callouts(slug, raw):
            d = disp.get((slug, c["id"]), c["id"])
            if c["prefix"] == "GAP":
                o_text = c["fields"].get("Owner to confirm", "")
                g_owner, _ = ppl.contact_for_text(o_text)
                g_owner = g_owner or owner
                gaps.append({
                    "slug": slug, "local": c["id"], "disp": d,
                    "text": c["text"],
                    "nature": c["fields"].get("Nature", ""),
                    "owner": g_owner,
                    "escalation": ppl.manager_of(g_owner) if g_owner else "",
                    "proc_label": label,
                })
            elif c["prefix"] == "SC":
                screens.append({
                    "slug": slug, "local": c["id"], "disp": d,
                    "text": c["text"], "step": steps.get(c["id"], ""),
                    "owner": owner, "proc_label": label,
                })
    return procs, gaps, screens


# --------------------------------------------------------------------------- #
# screenshot template
# --------------------------------------------------------------------------- #

def _bookmark(p, name: str, bid: int) -> None:
    start = OxmlElement("w:bookmarkStart")
    start.set(qn("w:id"), str(bid))
    start.set(qn("w:name"), name)
    end = OxmlElement("w:bookmarkEnd")
    end.set(qn("w:id"), str(bid))
    pPr = p._p.find(qn("w:pPr"))
    p._p.insert(1 if pPr is not None else 0, start)
    p._p.append(end)


def _cell_border(cell, color: str = "BFBFBF") -> None:
    tcPr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    for side in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"), "dashed")
        el.set(qn("w:sz"), "6")
        el.set(qn("w:color"), color)
        borders.append(el)
    tcPr.append(borders)


def write_screenshot_template(path: Path, label: str, contact: str,
                              items: list[dict], area_title: str) -> dict:
    """One entry per SC item: heading, italic what-to-capture, a dashed paste
    box anchored with a bookmark (screens_ingest keys extraction on it).
    Returns the anchor map {bookmark: {slug, local, disp}}."""
    doc = Document()
    st = doc.styles["Normal"]
    st.font.name = FONT
    st.font.size = Pt(10)

    h = doc.add_paragraph()
    r = h.add_run(f"Screenshot Collection — {label}")
    r.bold = True
    r.font.size = Pt(16)
    r.font.color.rgb = DARK
    sub = doc.add_paragraph()
    r = sub.add_run(area_title)
    r.italic = True
    r.font.color.rgb = GREEN
    cp = doc.add_paragraph()
    r = cp.add_run(f"Contact: {contact}")
    r.font.color.rgb = GRAY

    intro = doc.add_paragraph()
    r = intro.add_run(
        "For each item below: capture the screen described, paste it INSIDE "
        "the dashed box (click in the box, then paste), and save this file. "
        "Please don't delete or reorder the entries — an empty box just means "
        "\"not available\", which is also useful to know."
    )
    r.font.color.rgb = GRAY
    doc.add_paragraph()

    entries: dict[str, dict] = {}
    for n, it in enumerate(items):
        hp = doc.add_paragraph()
        r = hp.add_run(f"{it['disp']} — {it['proc_label']}"
                       + (f"  ·  {it['step']}" if it["step"] else ""))
        r.bold = True
        r.font.color.rgb = GREEN
        dp = doc.add_paragraph()
        r = dp.add_run(it["text"])
        r.italic = True

        t = doc.add_table(rows=1, cols=1)
        c = t.cell(0, 0)
        c.width = Inches(6.3)
        c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        _cell_border(c)
        p = c.paragraphs[0]
        name = f"scr_{n:03d}"
        _bookmark(p, name, 9500 + n)
        r = p.add_run("Paste screenshot here")
        r.font.color.rgb = GRAY
        r.italic = True
        # give the box some breathing room
        trPr = t.rows[0]._tr.get_or_add_trPr()
        trH = OxmlElement("w:trHeight")
        trH.set(qn("w:val"), "2400")
        trPr.append(trH)
        doc.add_paragraph()
        entries[name] = {"slug": it["slug"], "local": it["local"], "disp": it["disp"]}

    doc_id = uuid.uuid4().hex[:12]
    doc.core_properties.category = f"cw-screens:{doc_id}"
    doc.save(str(path))
    return {"schema": "consult-screens-map/v1", "doc_id": doc_id,
            "docx": path.name, "entries": entries}


# --------------------------------------------------------------------------- #
# kit assembly
# --------------------------------------------------------------------------- #

def _readme(kit: dict, area_title: str) -> str:
    extra = kit["extra_contacts"]
    parts = [
        f"# Review kit — {kit['label']}",
        "",
        f"_{area_title}_",
        "",
        f"**Send to: {kit['contact']}**"
        + (f" · also involves: {', '.join(extra)}" if extra else ""),
        "",
        "Hi — this folder covers one procedure in the current-state process",
        "documentation. Three quick asks, all in this folder:",
        "",
        "1. **Review the procedure document** and correct anything that is",
        "   wrong or missing — edit directly in Word. The document opens locked",
        "   in *Reviewing* mode: every edit is recorded automatically, even if",
        "   it displays as normal text (Word's *Simple Markup* view). Just",
        "   type — nothing extra to turn on. To see your edits as redlines,",
        "   pick **All Markup** in the Review ribbon. Add comments for anything",
        "   you want to explain rather than fix.",
        f"- `{kit['doc']}`",
    ]
    if kit["gaps"]:
        parts += [
            "",
            f"2. **Answer the open questions** in `{kit['xlsx']}` — fill the",
            "   *Answer* column (and *Status* if you like). The *Contact*",
            "   column names who each question is for — usually you; if a row",
            "   names someone else, forward it or answer together. \"I don't",
            "   know\" or \"ask X instead\" are useful answers too.",
        ]
    if kit["screens"]:
        parts += [
            "",
            f"3. **Capture the screenshots** listed in `{kit['screens_doc']}` —",
            "   paste each one inside its dashed box and save.",
        ]
    parts += [
        "",
        "When you're done, send the whole folder back. Thank you!",
        "",
    ]
    return "\n".join(parts)


def build_kits(area: str, out_dir: str | None = None) -> int:
    folder = Path(area).resolve()
    if not (folder / "manifest.json").is_file():
        raise SystemExit(f"error: no manifest.json in {folder}")
    manifest = doc_model.load_manifest(folder)
    area_title = manifest.get("title", folder.name)

    procs, gaps, screens = collect(folder)

    out = Path(out_dir) if out_dir else folder / "_review" / "kits"
    if out.exists():
        shutil.rmtree(out)   # kits are derived artifacts; regenerate whole
    out.mkdir(parents=True)

    # One kit per PROCEDURE (L3): review travels by topic. Contacts are
    # annotations inside the kit, not the grouping key.
    def _numkey(n: str):
        try:
            return tuple(int(x) for x in n.split("."))
        except ValueError:
            return (99,)

    kits: dict[str, dict] = {}
    for slug, p in procs.items():
        contact = (p["owner"]
                   or (p["role"] and f"({p['role']} — no person mapped)")
                   or "(unassigned)")
        kits[slug] = {
            "key": f"{p['number'].replace('.', '-')}_{slug}" if p["number"] else slug,
            "slug": slug, "label": p["label"], "number": p["number"],
            "contact": contact, "assigned": bool(p["owner"]),
            "extra_contacts": [],
            "proc": p, "gap_rows": [], "screen_items": [],
            "doc": "", "xlsx": "", "screens_doc": "", "gaps": 0, "screens": 0,
        }
    for g in gaps:
        kits[g["slug"]]["gap_rows"].append(g)
    for s in screens:
        kits[s["slug"]]["screen_items"].append(s)
    for kit in kits.values():
        kit["extra_contacts"] = sorted(
            {g["owner"] for g in kit["gap_rows"]
             if g["owner"] and g["owner"] != kit["proc"]["owner"]})

    maps_dir = folder / "_review" / ".maps"
    maps_dir.mkdir(parents=True, exist_ok=True)

    ordered = sorted(kits.values(), key=lambda k: (_numkey(k["number"]), k["slug"]))
    for kit in ordered:
        kdir = out / kit["key"]
        kdir.mkdir(parents=True)
        slug = kit["slug"]

        kit["doc"] = f"{kit['number']}_{slug}.docx"
        render_mod.render_folder(
            folder, kdir / kit["doc"], mode="working", slugs=[slug],
            track_changes=True, emit_signal=False)

        if kit["gap_rows"]:
            kit["gaps"] = len(kit["gap_rows"])
            kit["xlsx"] = f"gaps_{slug}.xlsx"
            rows = [[g["disp"], g["proc_label"], g["text"], g["nature"],
                     g["owner"], g["escalation"], "", "",
                     f"{g['slug']}#{g['local']}"]
                    for g in sorted(kit["gap_rows"], key=lambda g: g["disp"])]
            write_xlsx(kdir / kit["xlsx"], GAP_HEADER, rows,
                       sheet_name="Open Questions", widths=GAP_WIDTHS)

        if kit["screen_items"]:
            kit["screens"] = len(kit["screen_items"])
            kit["screens_doc"] = f"screenshots_{slug}.docx"
            items = sorted(kit["screen_items"], key=lambda s: s["disp"])
            smap = write_screenshot_template(
                kdir / kit["screens_doc"], kit["label"], kit["contact"],
                items, area_title)
            (maps_dir / f"{smap['doc_id']}.json").write_text(
                json.dumps(smap, indent=1, ensure_ascii=False) + "\n",
                encoding="utf-8")

        (kdir / "README.md").write_text(_readme(kit, area_title),
                                        encoding="utf-8")

    # index — the send checklist, kit per row plus a by-person rollup
    lines = [f"# Review kits — {area_title}", "",
             "| Kit | Procedure | Send to | Gaps | Screenshots |",
             "|---|---|---|---|---|"]
    for kit in ordered:
        extra = (" (+ " + ", ".join(kit["extra_contacts"]) + ")"
                 if kit["extra_contacts"] else "")
        lines.append(f"| `{kit['key']}/` | {kit['label']} | "
                     f"{kit['contact']}{extra} | {kit['gaps']} | {kit['screens']} |")
    by_person: dict[str, list[str]] = {}
    for kit in ordered:
        by_person.setdefault(kit["contact"], []).append(kit["key"])
    lines += ["", "## By person", ""]
    for person in sorted(by_person):
        lines.append(f"- **{person}**: " +
                     ", ".join(f"`{k}/`" for k in by_person[person]))
    lines += ["", "Send each kit folder to its person (one person may get "
              "several); returned files go to `_review/returned/`.", ""]
    (out / "index.md").write_text("\n".join(lines), encoding="utf-8")

    print(f"kits: {len(kits)} kit(s) under {out}")
    for kit in ordered:
        print(f"  {kit['key']}: -> {kit['contact']}, {kit['gaps']} gap(s), "
              f"{kit['screens']} screenshot item(s)")
    return 0


def main(argv=None) -> int:
    ap = argparse.ArgumentParser(description="Emit per-owner review kits (M9)")
    ap.add_argument("area", help="area folder")
    ap.add_argument("-o", "--out", default=None, help="kits output dir")
    a = ap.parse_args(argv)
    return build_kits(a.area, a.out)


if __name__ == "__main__":
    sys.exit(main())
