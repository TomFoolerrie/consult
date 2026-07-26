#!/usr/bin/env python3
"""
review_apply.py — deterministic tracked-changes apply (M10).

    python3 scripts/review_apply.py <area> [reviewed.docx ...]

With no explicit files, sweeps `{area}/_review/returned/` for reviewed docs
(anything whose core category reads `cw-map:<doc_id>` — renders stamp it, so
renamed files still self-identify). Applies reviewers' tracked INSERTIONS /
DELETIONS back into the markdown fragments with ZERO token spend, under a
verified-or-fallback guarantee:

  precision is structural — nothing unverified is ever written;
  recall failures cost a drafter note (today's workflow), never corruption.

The pipeline, per edited docx paragraph:

  1. ANCHOR   bookmark (`cw_NNNNN`) → the sidecar review map written at render
              time (`_review/.maps/<doc_id>.json`: file, slug, l2, source line
              span, block kind, sub index, text hash).
  2. VERIFY   the paragraph's reject-all-changes text must equal BOTH the
              recorded hash AND the text re-synthesized from the CURRENT
              fragment (an aligned re-run of the render transforms: local→
              display IDs, [[slug]] resolution, gap-tag flag, markdown marker
              strip). Any mismatch — fragment changed since render, anchor
              tampered, or a transform this synthesizer can't align — falls
              back.
  3. SPLICE   common-prefix/suffix diff between original and accepted text,
              mapped through the alignment back to markdown coordinates, so
              formatting/tokens OUTSIDE the edit survive verbatim; display IDs
              typed by the reviewer INSIDE the edit are mapped back to local.
  4. RE-VERIFY the whole edited block is forward-rendered through the REAL
              converter and its paragraph texts must equal the reviewer's
              accepted texts exactly; then the fragment must not gain any new
              reconcile grammar error. Either failure reverts the file and
              routes every one of its edits to notes.

Routed to notes (`notes_util` merge into `_review/{slug}.notes.yaml`) instead
of applied: unanchored/ambiguous paragraphs, hash mismatches, paragraph-mark
insertions/deletions (splits/merges), table edits (tables are never anchored:
derived tables are regenerated — triaged separately — and agent-owned sections
belong to their agents), callout label/ID line tampering, and any failed
verification. COMMENTS are never touched here — run
`review_extract.py --comments-only` after (it also archives the docx; this
script deliberately leaves it in place so extract still finds it).

Python 3, stdlib + pyyaml + python-docx (+ the bundled converter).
"""

from __future__ import annotations

import argparse
import hashlib
import json
import re
import sys
from pathlib import Path

_SCRIPTS_DIR = Path(__file__).resolve().parent
_REPO_ROOT = _SCRIPTS_DIR.parent
_SKILL_SCRIPTS = _REPO_ROOT / "skills" / "consult-docx-builder" / "scripts"
for _p in (_SCRIPTS_DIR, _SKILL_SCRIPTS):
    if str(_p) not in sys.path:
        sys.path.insert(0, str(_p))

import doc_model  # noqa: E402
import render as render_mod  # noqa: E402
import reconcile as reconcile_mod  # noqa: E402
import cfgi_markdown_to_word as cfgi  # noqa: E402
from notes_util import append_items  # noqa: E402

from docx import Document  # noqa: E402
from docx.oxml.ns import qn  # noqa: E402

UNASSIGNED = "_unassigned"

_XREF_RE = re.compile(r"\[\[(#?)([a-z0-9][a-z0-9-]*)\]\]")
_GAPTAG_RE = re.compile(r"\[\[\s*(GAP-\d+[^\]]*?)\s*\]\]")
_LINK_RE = re.compile(r"\[([^\]]+)\]\(([^)]+)\)")
_CALLOUT_DEF_RE = re.compile(
    r"^(?P<pre>\*\*[A-Z][A-Z ]+?\s*[-–—]\s*[A-Z0-9-]+:\*\*)")
_LIST_RE = re.compile(r"^\s*(?:[-*+]|\d+[.)])\s+")


# --------------------------------------------------------------------------- #
# docx side: per-paragraph original/accepted text + anchors + authorship
# --------------------------------------------------------------------------- #

def _ln(el) -> str:
    return el.tag.rsplit("}", 1)[-1]


def paragraph_texts(p_el):
    """(original, accepted, meta) for one w:p.

    original = reject all changes; accepted = accept all changes.
    meta: {edited, mark_changed, author, date, bookmarks}."""
    orig: list[str] = []
    fin: list[str] = []
    meta = {"edited": False, "mark_changed": False, "author": "", "date": "",
            "bookmarks": []}

    def note_authorship(el):
        if not meta["author"]:
            meta["author"] = el.get(qn("w:author")) or ""
            meta["date"] = el.get(qn("w:date")) or ""

    def walk(el, in_ins: bool, in_del: bool):
        for child in el:
            tag = _ln(child)
            if tag == "pPr":
                # inserted/deleted PARAGRAPH MARK = split/merge → not appliable
                rpr = child.find(qn("w:rPr"))
                if rpr is not None and (
                        rpr.find(qn("w:ins")) is not None
                        or rpr.find(qn("w:del")) is not None):
                    meta["mark_changed"] = True
                    meta["edited"] = True
                continue
            if tag == "bookmarkStart":
                nm = child.get(qn("w:name")) or ""
                if nm.startswith("cw_"):
                    meta["bookmarks"].append(nm)
                continue
            if tag == "ins":
                meta["edited"] = True
                note_authorship(child)
                walk(child, True, in_del)
                continue
            if tag == "del":
                meta["edited"] = True
                note_authorship(child)
                walk(child, in_ins, True)
                continue
            if tag == "r":
                for sub in child:
                    st = _ln(sub)
                    if st == "t":
                        txt = sub.text or ""
                        if in_ins:
                            fin.append(txt)
                        elif in_del:
                            orig.append(txt)   # defensive; del uses delText
                        else:
                            orig.append(txt)
                            fin.append(txt)
                    elif st == "delText":
                        orig.append(sub.text or "")
                    elif st == "tab":
                        (fin if in_ins else orig).append("\t")
                        if not (in_ins or in_del):
                            fin.append("\t")
                continue
            # containers (hyperlink, smartTag, sdt …) — recurse
            walk(child, in_ins, in_del)

    walk(p_el, False, False)
    return "".join(orig), "".join(fin), meta


def _in_table(p_el) -> bool:
    el = p_el.getparent()
    while el is not None:
        if _ln(el) == "tc":
            return True
        el = el.getparent()
    return False


# --------------------------------------------------------------------------- #
# aligned rendered-text synthesis (mirrors render.py + converter segs())
# --------------------------------------------------------------------------- #

def _aligned_sub(text: str, amap: list, regex, repl) -> tuple[str, list]:
    """Regex substitute while keeping a rendered→markdown index map.

    Replacement characters map to the markdown index of the match START
    (i.e. an edit touching them splices from the whole matched construct)."""
    out: list[str] = []
    omap: list[int] = []
    last = 0
    for m in regex.finditer(text):
        out.append(text[last:m.start()])
        omap.extend(amap[last:m.start()])
        r = repl(m) if callable(repl) else m.expand(repl)
        anchor = amap[m.start()] if m.start() < len(amap) else -1
        out.append(r)
        omap.extend([anchor] * len(r))
        last = m.end()
    out.append(text[last:])
    omap.extend(amap[last:])
    return "".join(out), omap


def synth_rendered(logical_md: str, submap: dict, labels: dict,
                   synthetic_prefix: str = "") -> tuple[str, list] | None:
    """(rendered_text, map rendered-idx→logical-md-idx), or None when this
    synthesizer can't faithfully reproduce the converter (clean() not
    identity) — the caller falls back to a note."""
    if cfgi.clean(logical_md) != logical_md:
        return None
    text = logical_md
    amap = list(range(len(text)))
    # 1. local → display callout IDs (same regex render.py uses)
    text, amap = _aligned_sub(
        text, amap, render_mod._CALLOUT_ID_RE,
        lambda m: submap.get(m.group(0), m.group(0)))
    # 2. [[slug]] / [[#slug]] cross-references → label / bare number
    def xref(m):
        label = labels.get(m.group(2))
        if label is None:
            return m.group(0)
        return label.split(" ", 1)[0] if m.group(1) else label
    text, amap = _aligned_sub(text, amap, _XREF_RE, xref)
    # 3. body gap tags: [[GAP-x — y]] → **[GAP-x — y]** → (bold strip) [·]
    text, amap = _aligned_sub(text, amap, _GAPTAG_RE, lambda m: f"[{m.group(1)}]")
    # 4-7. converter inline pass: links, then bold/italic/code marker strip
    text, amap = _aligned_sub(text, amap, _LINK_RE, lambda m: m.group(1))
    text, amap = _aligned_sub(text, amap, cfgi._BOLD_RE, lambda m: m.group(1))
    text, amap = _aligned_sub(text, amap, cfgi._ITALIC_RE,
                              lambda m: m.group(1) or m.group(2) or "")
    text, amap = _aligned_sub(text, amap, cfgi._CODE_RE, lambda m: m.group(1))
    if synthetic_prefix:
        text = synthetic_prefix + text
        amap = [-1] * len(synthetic_prefix) + amap
    return text, amap


def _md_index(amap: list, rendered_idx: int, side: str) -> int:
    """Nearest real markdown index at a rendered position; synthetic (-1)
    positions roll left/right depending on which diff side we're on."""
    if rendered_idx >= len(amap):
        # one past the end: markdown end
        for k in range(len(amap) - 1, -1, -1):
            if amap[k] >= 0:
                return amap[k] + 1
        return 0
    j = rendered_idx
    if side == "start":
        while j < len(amap) and amap[j] < 0:
            j += 1
        if j == len(amap):
            return _md_index(amap, len(amap), "start")
        return amap[j]
    while j >= 0 and amap[j] < 0:
        j -= 1
    return amap[j] + 1 if j >= 0 else 0


# --------------------------------------------------------------------------- #
# markdown block model (mirrors the converter's paragraph grouping)
# --------------------------------------------------------------------------- #

class Segment:
    """One docx paragraph's markdown source inside a block."""

    def __init__(self, sub, rel_start, rel_end, prefix, logical, synthetic=""):
        self.sub = sub
        self.rel_start = rel_start   # relative line range within the block
        self.rel_end = rel_end       # inclusive
        self.prefix = prefix         # structural line prefix to re-emit
        self.logical = logical       # joined, stripped markdown content
        self.synthetic = synthetic   # rendered-only prefix ("•  ")
        self.deleted = False
        self.dirty = False           # only dirty segments are re-emitted


def block_segments(block_lines: list[str], kind: str) -> list[Segment]:
    if kind in ("para", "bullet"):
        first = block_lines[0]
        if kind == "bullet":
            m = _LIST_RE.match(first)
            prefix = first[:m.end()] if m else ""
            head = first[m.end():].strip() if m else first.strip()
        else:
            prefix = ""
            head = first.strip()
        parts = [head] + [ln.strip() for ln in block_lines[1:] if ln.strip()]
        return [Segment(0, 0, len(block_lines) - 1, prefix, " ".join(parts))]

    # callout: reproduce add_callout's grouping over the '>'-stripped lines
    segs: list[Segment] = []
    buf: list[str] = []
    buf_start = 0
    sub = 0

    def flush(end_rel: int):
        nonlocal sub, buf
        if buf:
            segs.append(Segment(sub, buf_start, end_rel, "> ", " ".join(buf)))
            sub += 1
            buf = []

    for rel, raw in enumerate(block_lines):
        s = re.sub(r"^\s*>\s?", "", raw).strip()
        if not s:
            flush(rel - 1)
        elif re.match(r"^[-*+]\s+", s):
            flush(rel - 1)
            content = re.sub(r"^[-*+]\s+", "", s)
            segs.append(Segment(sub, rel, rel, "> - ", content, "•  "))
            sub += 1
        else:
            if not buf:
                buf_start = rel
            buf.append(s)
    flush(len(block_lines) - 1)
    return segs


def rebuild_block(block_lines: list[str], segments: list[Segment]) -> list[str]:
    """Replace each EDITED segment's line range with its single-line form;
    untouched segments and separator lines survive byte-for-byte (no cosmetic
    reflow of hard-wrapped source). Descending order keeps ranges valid."""
    out = list(block_lines)
    for seg in sorted(segments, key=lambda s: s.rel_start, reverse=True):
        if not (seg.dirty or seg.deleted):
            continue
        repl = [] if seg.deleted else [seg.prefix + seg.logical]
        out[seg.rel_start:seg.rel_end + 1] = repl
    return out


# --------------------------------------------------------------------------- #
# forward verification: render the edited block through the REAL converter
# --------------------------------------------------------------------------- #

def forward_paragraph_texts(block_lines: list[str], submap: dict,
                            labels: dict) -> list[str]:
    text = "\n".join(block_lines)
    text = render_mod._rewrite_callout_ids(text, submap)
    text = render_mod._resolve_tokens(text, labels)
    text = render_mod._flag_gap_tags(text)
    doc = Document()
    cfgi.render_body(doc, text.split("\n"), do_cover=False)
    texts = []
    for t in doc.tables:            # callout boxes
        for row in t.rows:
            for c in row.cells:
                texts.extend(p.text for p in c.paragraphs if p.text.strip())
    for p in doc.paragraphs:        # plain paragraphs / list items
        if p.text.strip():
            texts.append(p.text)
    return texts


# --------------------------------------------------------------------------- #
# grammar guard (procedure fragments only)
# --------------------------------------------------------------------------- #

def _grammar_errors(slug: str, file: str, text: str) -> set[str]:
    frag = reconcile_mod.parse_procedure(slug, file, text)
    # line numbers shift on splice; compare on normalized messages
    return {re.sub(r":\d+:", ":", e) for e in frag.errors}


# --------------------------------------------------------------------------- #
# main apply
# --------------------------------------------------------------------------- #

class Edit:
    def __init__(self, entry: dict, orig: str, final: str, meta: dict):
        self.entry = entry
        self.orig = orig
        self.final = final
        self.meta = meta
        self.reason = ""   # set when falling back


def load_map(area: Path, doc) -> dict | None:
    cat = doc.core_properties.category or ""
    if not cat.startswith("cw-map:"):
        return None
    f = area / "_review" / ".maps" / f"{cat.split(':', 1)[1].strip()}.json"
    if not f.is_file():
        return None
    try:
        return json.loads(f.read_text(encoding="utf-8"))
    except json.JSONDecodeError:
        return None


def apply_docx(area: Path, path: Path, report: dict) -> None:
    doc = Document(str(path))
    rmap = load_map(area, doc)
    if rmap is None:
        report.setdefault("skipped", []).append(
            f"{path.name}: no review map (not a rendered working doc, or "
            f"_review/.maps/ entry missing) — leave to review_extract")
        return

    manifest = doc_model.load_manifest(area)
    numbers = doc_model.display_numbers(manifest)
    labels = dict(numbers)
    role_of: dict[str, str] = {}
    slug_of_file: dict[str, str] = {}
    for comp in manifest.get("components", []):
        s = comp.get("slug")
        if s and s in labels and comp.get("heading"):
            labels[s] = f"{labels[s]} {comp['heading']}"
        if comp.get("file"):
            role_of[comp["file"]] = comp.get("role", "")
            if s:
                slug_of_file[comp["file"]] = s
    id_map = doc_model.callout_display_ids(area)
    sub_by_slug: dict[str, dict] = {}
    for (s, local), disp in id_map.items():
        sub_by_slug.setdefault(s, {})[local] = disp
    disp2local: dict[str, dict] = {
        s: {d: l for l, d in m.items()} for s, m in sub_by_slug.items()}

    entries = rmap.get("entries", {})
    edits: list[Edit] = []
    fallbacks: list[Edit] = []

    # Iterate EVERY w:p in document order — doc.paragraphs skips table cells,
    # and callout boxes are 1×1 tables whose cell paragraphs ARE anchored.
    for p_el in doc.element.body.iter(qn("w:p")):
        orig, fin, meta = paragraph_texts(p_el)
        if not meta["edited"] or (orig == fin and not meta["mark_changed"]):
            continue
        entry = None
        for bm in meta["bookmarks"]:
            if bm in entries:
                entry = dict(entries[bm])
                break
        e = Edit(entry or {}, orig, fin, meta)
        if meta["mark_changed"]:
            e.reason = "paragraph split/merge — routed to notes"
            fallbacks.append(e)
        elif entry is None:
            # Anchored callout paragraphs live in (their own 1×1) tables;
            # an UNanchored table edit is a derived/agent table — not ours.
            e.reason = ("edit inside a derived/agent table — the real fix is "
                        "the fragment/registry" if _in_table(p_el)
                        else "no verified anchor (new/unanchored paragraph)")
            fallbacks.append(e)
        elif hashlib.sha1(orig.encode("utf-8")).hexdigest() != entry.get("hash"):
            e.reason = "anchor hash mismatch (paragraph moved/rebuilt?)"
            fallbacks.append(e)
        else:
            edits.append(e)

    # ---- group appliable edits by file, then block --------------------------
    by_file: dict[str, list[Edit]] = {}
    for e in edits:
        by_file.setdefault(e.entry["file"], []).append(e)

    applied: list[str] = []
    for fname, file_edits in sorted(by_file.items()):
        fpath = area / fname
        if role_of.get(fname) not in ("procedure", "static") or not fpath.is_file():
            for e in file_edits:
                e.reason = f"target {fname} not an appliable file"
                fallbacks.append(e)
            continue
        original_text = fpath.read_text(encoding="utf-8")
        lines = original_text.split("\n")
        slug = slug_of_file.get(fname, "")
        submap = sub_by_slug.get(slug, {})
        d2l = disp2local.get(slug, {})

        blocks: dict[tuple, list[Edit]] = {}
        for e in file_edits:
            blocks.setdefault(tuple(e.entry["lines"]), []).append(e)

        file_applied: list[str] = []
        file_failed: list[Edit] = []
        # descending start line keeps earlier blocks' recorded lines valid
        for (a, b), b_edits in sorted(blocks.items(), reverse=True):
            kind = b_edits[0].entry.get("kind", "para")
            block_lines = lines[a:b + 1]
            segments = block_segments(block_lines, kind)
            seg_by_sub = {s.sub: s for s in segments}
            ok = True
            for e in b_edits:
                seg = seg_by_sub.get(e.entry.get("sub", 0))
                if seg is None:
                    e.reason = "block shape drifted (sub index unmatched)"
                    ok = False
                    break
                synth = synth_rendered(seg.logical, submap, labels, seg.synthetic)
                if synth is None or synth[0] != e.orig:
                    e.reason = "fragment changed since render (re-synthesis mismatch)"
                    ok = False
                    break
                rendered, amap = synth
                if not e.final.strip():
                    seg.deleted = True
                    seg.dirty = True
                    continue
                # common prefix / suffix diff on rendered text
                pre = 0
                while (pre < len(rendered) and pre < len(e.final)
                       and rendered[pre] == e.final[pre]):
                    pre += 1
                suf = 0
                while (suf < len(rendered) - pre and suf < len(e.final) - pre
                       and rendered[len(rendered) - 1 - suf]
                       == e.final[len(e.final) - 1 - suf]):
                    suf += 1
                ins = e.final[pre:len(e.final) - suf]
                # reviewer-typed display IDs → this procedure's local IDs
                ins = render_mod._CALLOUT_ID_RE.sub(
                    lambda m: d2l.get(m.group(0), m.group(0)), ins)
                md_a = _md_index(amap, pre, "start")
                md_b = _md_index(amap, len(rendered) - suf, "start") \
                    if suf else len(seg.logical)
                if md_a > md_b:
                    md_a = md_b
                new_logical = seg.logical[:md_a] + ins + seg.logical[md_b:]
                # callout definition line: label/ID prefix must be untouched
                if kind == "callout" and seg.sub == 0:
                    m_old = _CALLOUT_DEF_RE.match(seg.logical)
                    m_new = _CALLOUT_DEF_RE.match(new_logical)
                    if not (m_old and m_new
                            and m_old.group("pre") == m_new.group("pre")):
                        e.reason = "callout label/ID line changed — needs a drafter"
                        ok = False
                        break
                seg.logical = new_logical
                seg.dirty = True
            if not ok:
                for e in b_edits:
                    if not e.reason:
                        e.reason = "sibling edit in the same block failed"
                    file_failed.append(e)
                continue
            new_block = rebuild_block(block_lines, segments)
            # forward verification: converter must reproduce the accepted texts
            expected = []
            for s in sorted(segments, key=lambda s: s.sub):
                if s.deleted:
                    continue
                edited = next((e for e in b_edits
                               if e.entry.get("sub", 0) == s.sub), None)
                if edited is not None:
                    expected.append(edited.final)
                else:
                    sy = synth_rendered(s.logical, submap, labels, s.synthetic)
                    expected.append(sy[0] if sy else None)
            got = forward_paragraph_texts(new_block, submap, labels)
            if len(got) != len(expected) or any(
                    x is not None and x != g for x, g in zip(expected, got)):
                for e in b_edits:
                    e.reason = "forward re-render mismatch — reverted"
                    file_failed.append(e)
                continue
            lines[a:b + 1] = new_block
            for e in b_edits:
                file_applied.append(
                    f"{fname}: {_clip(e.orig)} -> {_clip(e.final)}")

        new_text = "\n".join(lines)
        if file_applied and role_of.get(fname) == "procedure":
            before = _grammar_errors(slug, fname, original_text)
            after = _grammar_errors(slug, fname, new_text)
            if after - before:
                # transactional revert: the whole file's edits become notes
                for e in file_edits:
                    if not e.reason:
                        e.reason = ("apply would introduce grammar errors: "
                                    + "; ".join(sorted(after - before))[:160])
                    fallbacks.append(e)
                continue
        if file_applied:
            fpath.write_text(new_text, encoding="utf-8")
            applied.extend(file_applied)
        fallbacks.extend(file_failed)

    # ---- fallback notes ------------------------------------------------------
    noted = 0
    per_slug: dict[str, list[dict]] = {}
    for e in fallbacks:
        slug = e.entry.get("slug") or UNASSIGNED
        loc = e.entry.get("file", "unknown file")
        if e.entry.get("l2"):
            loc += f" (L2 {e.entry['l2']})"
        per_slug.setdefault(slug, []).append({
            "kind": "review",            # M6 notes-bus stamp (see notes_util)
            "type": "tracked-change",
            "location": loc,
            "anchor": _clip(e.orig, 160),
            "change": f"{_clip(e.orig, 200)} -> {_clip(e.final, 200)}"
                      + (f" [{e.reason}]" if e.reason else ""),
            "author": e.meta.get("author", ""),
            "date": e.meta.get("date", ""),
            "source": path.name,
        })
    for slug, items in sorted(per_slug.items()):
        noted += append_items(area, slug, items)

    report.setdefault("files", []).append({
        "docx": path.name, "applied": len(applied), "noted": noted,
        "details": applied,
        "fallback_reasons": sorted({e.reason for e in fallbacks if e.reason}),
    })


def _clip(s: str, n: int = 80) -> str:
    s = " ".join(s.split())
    return s if len(s) <= n else s[:n - 1] + "…"


def main(argv=None) -> int:
    ap = argparse.ArgumentParser(description="Apply tracked changes deterministically (M10)")
    ap.add_argument("area", help="area folder")
    ap.add_argument("files", nargs="*",
                    help="reviewed docx; default: _review/returned/*.docx")
    a = ap.parse_args(argv)

    area = Path(a.area).resolve()
    if not (area / "manifest.json").is_file():
        raise SystemExit(f"error: no manifest.json in {area}")
    files = [Path(f) for f in a.files] or sorted(
        (area / "_review" / "returned").glob("*.docx"))
    if not files:
        print("no reviewed documents to apply")
        return 0

    report: dict = {}
    for f in files:
        apply_docx(area, f, report)

    for s in report.get("skipped", []):
        print(f"SKIP {s}")
    tot_a = tot_n = 0
    for fr in report.get("files", []):
        tot_a += fr["applied"]
        tot_n += fr["noted"]
        print(f"{fr['docx']}: applied {fr['applied']}, noted {fr['noted']}")
        for d in fr["details"]:
            print(f"  APPLIED {d}")
        for r in fr["fallback_reasons"]:
            print(f"  fallback: {r}")
    print(f"total: {tot_a} applied deterministically, {tot_n} routed to notes")
    print("(comments untouched — run review_extract.py --comments-only next; "
          "it archives the docx)")
    return 0


if __name__ == "__main__":
    sys.exit(main())
