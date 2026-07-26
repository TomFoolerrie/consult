"""Tests for scripts/review_extract.py — extraction of reviewer comments and
tracked changes into procedure-anchored _review/{slug}.notes.yaml files.

Fixture: a real working-mode render of a tiny area (same shape as the apply
tests); comments are added with python-docx's Document.add_comment and tracked
changes are written into the XML the way Word records them.
"""
from __future__ import annotations

import copy
import json
import shutil
from pathlib import Path

import pytest
import yaml
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

import render as render_mod
import review_extract


BANK_REC = """## Bank Reconciliation

### A. Purpose

The clerk reconciles the monthly bank statement against the general ledger.

### B. Steps

- Download the bank statement from the portal.
- Compare balances and investigate any differences.

> **CONTROL — CTRL-01:** The controller signs off on the completed reconciliation.

See [[journal-entries]] for posting adjustments.
"""

JOURNAL = """## Journal Entries

### A. Purpose

Accountants post manual journal entries at month end.

> **PAIN POINT — PP-01:** Approvals are tracked by email. (Severity: Medium)
"""

MANIFEST = {
    "schema": "consult-mvp-manifest/v1",
    "area": "test-area",
    "l1": "record-to-report",
    "title": "Test Area Process Documentation",
    "subtitle": "Working Draft",
    "l2_order": ["close"],
    "components": [
        {"file": "bank-rec.md", "heading": "Bank Reconciliation",
         "role": "procedure", "slug": "bank-rec", "l2": "close", "order": 1},
        {"file": "journal-entries.md", "heading": "Journal Entries",
         "role": "procedure", "slug": "journal-entries", "l2": "close", "order": 2},
    ],
}


@pytest.fixture(scope="module")
def rendered_base(tmp_path_factory) -> Path:
    area = tmp_path_factory.mktemp("extract-base") / "area"
    area.mkdir()
    (area / "manifest.json").write_text(json.dumps(MANIFEST), encoding="utf-8")
    (area / "bank-rec.md").write_text(BANK_REC, encoding="utf-8")
    (area / "journal-entries.md").write_text(JOURNAL, encoding="utf-8")
    render_mod.render_folder(area, area / "draft.docx", emit_signal=False)
    return area


@pytest.fixture()
def area(rendered_base, tmp_path) -> Path:
    dst = tmp_path / "area"
    shutil.copytree(rendered_base, dst)
    (dst / "_review" / "returned").mkdir(parents=True, exist_ok=True)
    return dst


def para_by_text(doc, needle):
    for p in doc.paragraphs:
        if needle in p.text:
            return p
    raise AssertionError(f"no body paragraph containing {needle!r}")


def add_comment(doc, needle, text, author="SME Reviewer"):
    p = para_by_text(doc, needle)
    doc.add_comment(p.runs, text=text, author=author, initials="SR")


_wid = [500]


def track_replace(doc, needle, old, new, author="Reviewer One"):
    """Word-style tracked replacement (w:del/w:delText + w:ins) in the first
    paragraph containing `needle`."""
    for p_el in doc.element.body.iter(qn("w:p")):
        if needle not in "".join(t.text or "" for t in p_el.iter(qn("w:t"))):
            continue
        for r in p_el.iter(qn("w:r")):
            t = r.find(qn("w:t"))
            if t is None or old not in (t.text or ""):
                continue
            before, _, after = (t.text or "").partition(old)
            parent = r.getparent()
            idx = list(parent).index(r)

            def mk_run(txt):
                nr = copy.deepcopy(r)
                nt = nr.find(qn("w:t"))
                nt.text = txt
                nt.set(qn("xml:space"), "preserve")
                return nr

            def wrap(tag, txt, deltext=False):
                _wid[0] += 1
                el = OxmlElement(tag)
                el.set(qn("w:id"), str(_wid[0]))
                el.set(qn("w:author"), author)
                el.set(qn("w:date"), "2026-07-21T09:00:00Z")
                run = mk_run(txt)
                if deltext:
                    run.find(qn("w:t")).tag = qn("w:delText")
                el.append(run)
                return el

            nodes = [mk_run(before)] if before else []
            nodes.append(wrap("w:del", old, deltext=True))
            if new:
                nodes.append(wrap("w:ins", new))
            if after:
                nodes.append(mk_run(after))
            for n in nodes:
                parent.insert(idx, n)
                idx += 1
            parent.remove(r)
            return
    raise AssertionError(f"no run containing {old!r} near {needle!r}")


@pytest.fixture()
def returned(area) -> Path:
    """Reviewed copy in _review/returned/ carrying one comment on a bank-rec
    step AND one tracked change, so --comments-only filtering is observable."""
    rv = area / "_review" / "returned" / "reviewed.docx"
    shutil.copy(area / "draft.docx", rv)
    d = Document(str(rv))
    add_comment(d, "Download the bank statement",
                "Confirm which portal — Chase or the treasury workstation?")
    track_replace(d, "monthly bank statement", "monthly", "weekly")
    d.save(str(rv))
    return rv


def load_notes(area, slug):
    data = yaml.safe_load(
        (area / "_review" / f"{slug}.notes.yaml").read_text(encoding="utf-8"))
    assert data["procedure"] == slug
    return data["items"]


# --------------------------------------------------------------------------- #
# --comments-only extraction + archiving
# --------------------------------------------------------------------------- #

def test_comments_only_extracts_comment_and_archives(area, returned):
    """--comments-only writes the w:comment as an anchored note for the right
    procedure, skips the tracked-change layer, and archives the docx to
    _review/processed/."""
    rc = review_extract.main([str(returned), "--comments-only"])
    assert rc == 0
    items = load_notes(area, "bank-rec")
    assert len(items) == 1
    (item,) = items
    assert item["kind"] == "review"     # M6 notes-bus stamp (never retires a source)
    assert item["type"] == "comment"
    assert item["note"].startswith("Confirm which portal")
    assert item["author"] == "SME Reviewer"
    assert "Download the bank statement" in item["anchor"]
    assert item["location"].startswith("B")   # anchored under "### B. Steps"
    assert item["source"] == "reviewed.docx"
    # archived, no longer in returned/
    assert not returned.exists()
    assert (area / "_review" / "processed" / "reviewed.docx").is_file()


def test_comments_only_rerun_merges_and_dedupes(area, returned):
    """Re-extracting the same doc appends nothing: notes merge-append and
    dedupe on the full item fingerprint (idempotent re-runs)."""
    archived = area / "_review" / "processed" / "reviewed.docx"
    assert review_extract.main([str(returned), "--comments-only"]) == 0
    first = load_notes(area, "bank-rec")
    # doc was archived; put a copy back and extract again
    shutil.copy(archived, returned)
    assert review_extract.main([str(returned), "--comments-only"]) == 0
    assert load_notes(area, "bank-rec") == first


def test_directory_input_processes_every_docx(area, returned):
    """A directory argument (the kit-return path) processes each *.docx in it."""
    rc = review_extract.main([str(returned.parent), "--comments-only"])
    assert rc == 0
    assert len(load_notes(area, "bank-rec")) == 1
    assert not returned.exists()   # archived by the directory sweep too


def test_no_archive_flag_leaves_doc_in_place(area, returned):
    """--no-archive extracts the notes but never moves the docx."""
    rc = review_extract.main([str(returned), "--comments-only", "--no-archive"])
    assert rc == 0
    assert returned.exists()
    assert not (area / "_review" / "processed").exists()
    assert len(load_notes(area, "bank-rec")) == 1


def test_dry_run_writes_nothing(area, returned, capsys):
    """--dry-run prints the would-be notes and writes/archives nothing."""
    rc = review_extract.main([str(returned), "--dry-run"])
    assert rc == 0
    assert "Confirm which portal" in capsys.readouterr().out
    assert not (area / "_review" / "bank-rec.notes.yaml").exists()
    assert returned.exists()


# --------------------------------------------------------------------------- #
# full extraction (tracked changes included)
# --------------------------------------------------------------------------- #

def test_full_extraction_includes_tracked_changes(area, returned):
    """Without --comments-only the tracked del→ins pair is extracted as a
    single 'old → new' replacement item attributed to the same procedure."""
    rc = review_extract.main([str(returned), "--no-archive"])
    assert rc == 0
    items = load_notes(area, "bank-rec")
    kinds = sorted(it["type"] for it in items)
    assert kinds == ["comment", "tracked-change"]
    (tc,) = [it for it in items if it["type"] == "tracked-change"]
    assert tc["change"] == "monthly → weekly"
    assert tc["author"] == "Reviewer One"
    assert tc["location"].startswith("A")   # under "### A. Purpose"
    # anchor is the paragraph's VISIBLE text (insertions in, deletions out)
    assert "weekly bank statement" in tc["anchor"]


def test_comment_outside_any_procedure_routes_to_unassigned(area):
    """A comment anchored before the first procedure heading (cover text) is
    routed to _unassigned.notes.yaml rather than dropped."""
    rv = area / "_review" / "returned" / "cover-comment.docx"
    shutil.copy(area / "draft.docx", rv)
    d = Document(str(rv))
    add_comment(d, "Working Draft", "Is this the right subtitle?")
    d.save(str(rv))
    rc = review_extract.main([str(rv), "--comments-only", "--no-archive"])
    assert rc == 0
    (item,) = load_notes(area, review_extract.UNASSIGNED)
    assert item["type"] == "comment"
    assert item["note"] == "Is this the right subtitle?"
    assert not (area / "_review" / "bank-rec.notes.yaml").exists()


def test_missing_file_and_bad_area_exit_nonzero(tmp_path):
    """A nonexistent docx or an area without manifest.json exits 2."""
    assert review_extract.main([str(tmp_path / "nope.docx")]) == 2
    d = Document()
    d.add_paragraph("x")
    orphan = tmp_path / "orphan.docx"
    d.save(str(orphan))
    assert review_extract.main([str(orphan)]) == 2   # no manifest around it
