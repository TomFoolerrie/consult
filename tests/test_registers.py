"""Tests for the M30 register machinery.

The verb (`engagement.py register add|update|list` → registers.py) is the
ONLY writer of components/_client/registers/*.md — humans included. Entries
carry a class (citable | context) and mandatory provenance; briefs list them
by class, render compiles CITABLE entries into a Shared Reference appendix
and NEVER shows context entries.
"""
import json
from pathlib import Path

import pytest
import yaml
from docx import Document

import brief
import consolidate
import engagement
import registers
import render
from test_render import make_area as make_render_area, doc_text


def reg(argv):
    """Run the verb the way the orchestrator does — through engagement.py."""
    return engagement.main(["register"] + argv)


def make_root(tmp_path: Path) -> Path:
    root = tmp_path / "components"
    root.mkdir()
    return root


def add(root, name="p2p-facts", eid="approval-threshold", cls="citable",
        text="PO approval threshold is USD 10,000.",
        provenance="SRC-004 (procure-to-pay)"):
    return reg(["add", name, "--id", eid, "--class", cls, "--text", text,
                "--provenance", provenance, "--root", str(root)])


# --------------------------------------------------------------------------- #
# The verb — round trip, refusals, idempotency
# --------------------------------------------------------------------------- #

def test_add_creates_file_with_h1_and_round_trips(tmp_path, capsys):
    root = make_root(tmp_path)
    assert add(root) == 0
    path = root / "_client" / "registers" / "p2p-facts.md"
    text = path.read_text(encoding="utf-8")
    assert text.startswith("# P2P Facts\n")          # H1 title on creation
    assert "## approval-threshold  [citable]" in text
    assert "provenance: SRC-004 (procure-to-pay)" in text
    capsys.readouterr()
    assert reg(["list", "--root", str(root)]) == 0
    out = capsys.readouterr().out
    # machine-stable entry-id convention: <register>#<entry-id>
    assert "p2p-facts#approval-threshold  [citable]  PO approval " \
           "threshold is USD 10,000." in out
    assert "provenance: SRC-004 (procure-to-pay)" in out


def test_add_is_idempotent_at_identical_content(tmp_path, capsys):
    root = make_root(tmp_path)
    assert add(root) == 0
    before = (root / "_client" / "registers" / "p2p-facts.md").read_text(
        encoding="utf-8")
    capsys.readouterr()
    assert add(root) == 0
    assert "no-op" in capsys.readouterr().out
    assert (root / "_client" / "registers" / "p2p-facts.md").read_text(
        encoding="utf-8") == before


def test_add_refuses_id_collision_with_different_content(tmp_path, capsys):
    root = make_root(tmp_path)
    assert add(root) == 0
    assert add(root, text="a DIFFERENT threshold") == 2
    err = capsys.readouterr().err
    assert "collision" in err and "register update" in err


@pytest.mark.parametrize("cls", ["citable", "context"])
def test_add_refuses_missing_provenance_for_both_classes(tmp_path, capsys, cls):
    root = make_root(tmp_path)
    assert add(root, cls=cls, provenance="") == 2
    assert "provenance" in capsys.readouterr().err
    assert not (root / "_client" / "registers" / "p2p-facts.md").exists()


def test_add_refuses_unknown_class(tmp_path, capsys):
    root = make_root(tmp_path)
    assert add(root, cls="shiny") == 2
    assert "unknown --class" in capsys.readouterr().err


def test_update_replaces_text_and_optionally_provenance(tmp_path, capsys):
    root = make_root(tmp_path)
    assert add(root) == 0
    assert reg(["update", "p2p-facts", "--id", "approval-threshold",
                "--text", "Threshold raised to USD 12,500.",
                "--root", str(root)]) == 0
    text = (root / "_client" / "registers" / "p2p-facts.md").read_text(
        encoding="utf-8")
    assert "USD 12,500" in text and "USD 10,000" not in text
    assert "SRC-004" in text                    # provenance kept when omitted
    assert reg(["update", "p2p-facts", "--id", "approval-threshold",
                "--text", "Threshold raised to USD 12,500.",
                "--provenance", "SRC-009 (treasury)",
                "--root", str(root)]) == 0
    text = (root / "_client" / "registers" / "p2p-facts.md").read_text(
        encoding="utf-8")
    assert "SRC-009 (treasury)" in text and "SRC-004" not in text


def test_update_refuses_unknown_id_listing_known_ones(tmp_path, capsys):
    root = make_root(tmp_path)
    assert add(root) == 0
    assert add(root, eid="cutoff-rule", cls="context",
               text="Cutoff is disputed.", provenance="SRC-002 (ap)") == 0
    assert reg(["update", "p2p-facts", "--id", "nope", "--text", "x",
                "--root", str(root)]) == 2
    err = capsys.readouterr().err
    assert "approval-threshold" in err and "cutoff-rule" in err


def test_verb_refuses_to_write_into_unstructured_file(tmp_path, capsys):
    root = make_root(tmp_path)
    regs = root / "_client" / "registers"
    regs.mkdir(parents=True)
    (regs / "approval-matrix.md").write_text("| band | approver |\n",
                                             encoding="utf-8")
    assert add(root, name="approval-matrix") == 2
    assert "unstructured" in capsys.readouterr().err


def test_list_tolerates_unstructured_files_and_filters_loudly(tmp_path,
                                                              capsys):
    root = make_root(tmp_path)
    regs = root / "_client" / "registers"
    regs.mkdir(parents=True)
    (regs / "approval-matrix.md").write_text("| band | approver |\n",
                                             encoding="utf-8")
    assert add(root, eid="cutoff-rule", cls="context",
               text="Cutoff ownership disputed.",
               provenance="SRC-002 (ap)") == 0
    capsys.readouterr()
    assert reg(["list", "--root", str(root)]) == 0
    out = capsys.readouterr().out
    assert f"unstructured (pre-M30): {regs / 'approval-matrix.md'}" in out
    assert "p2p-facts#cutoff-rule  [context]" in out
    # a filtered view says what it filtered
    assert reg(["list", "--root", str(root), "--class", "citable"]) == 0
    out = capsys.readouterr().out
    assert "filtered view (class=citable)" in out
    assert "cutoff-rule" not in out


# --------------------------------------------------------------------------- #
# Brief consumption — drafter, consolidator, placement
# --------------------------------------------------------------------------- #

def make_brief_area(root: Path) -> Path:
    """A minimal drafter-briefable area UNDER a components/ root."""
    area = root / "ap"
    (area / "_reference").mkdir(parents=True)
    (area / "manifest.json").write_text(json.dumps({
        "schema": "consult-mvp-manifest/v1", "area": "ap", "l1": "finance",
        "title": "AP", "l2_order": ["invoices"],
        "components": [
            {"file": "10_bank-rec.md", "heading": "Bank Rec",
             "role": "procedure", "slug": "bank-rec", "l2": "invoices",
             "order": 10},
            {"file": "20_payment-run.md", "heading": "Payment Run",
             "role": "procedure", "slug": "payment-run", "l2": "invoices",
             "order": 20},
        ]}), encoding="utf-8")
    (area / "10_bank-rec.md").write_text("## Bank Rec\n\ntext\n",
                                         encoding="utf-8")
    (area / "20_payment-run.md").write_text("## Payment Run\n\ntext\n",
                                            encoding="utf-8")
    (area / "_reference" / "sources.yaml").write_text(
        yaml.safe_dump({"sources": []}), encoding="utf-8")
    return area


def seed_registers(root: Path) -> None:
    assert add(root, eid="approval-threshold") == 0
    assert add(root, eid="callback-owner", cls="context",
               text="Callback ownership is disputed (AP vs Treasury).",
               provenance="SRC-002 (procure-to-pay)") == 0


def test_drafter_brief_lists_entries_by_class_with_rules(tmp_path, capsys):
    root = make_root(tmp_path)
    area = make_brief_area(root)
    seed_registers(root)
    regs = root / "_client" / "registers"
    (regs / "legacy.md").write_text("freeform seed\n", encoding="utf-8")
    capsys.readouterr()
    assert brief.main([str(area), "--slug", "bank-rec"]) == 0
    out = capsys.readouterr().out
    # citable: the reference-don't-restate rule, entry ids + first line only
    assert "citable — reference, never restate" in out
    assert "p2p-facts#approval-threshold: PO approval threshold is " \
           "USD 10,000." in out
    # context: the pre-read rule
    assert "context — facts already established: align; cite the " \
           "provenance source if you state one, GAP if you can't; NEVER " \
           "cite the register by name" in out
    assert "p2p-facts#callback-owner: Callback ownership is disputed" in out
    # unstructured seed file keeps the old whole-file line
    assert str(regs / "legacy.md") in out and "never restate" in out


def test_consolidator_brief_lists_structured_entries(tmp_path, capsys):
    root = make_root(tmp_path)
    area = make_brief_area(root)
    seed_registers(root)
    capsys.readouterr()
    assert consolidate.main(["brief", str(area), "--bucket",
                             "invoices"]) == 0
    out = capsys.readouterr().out
    assert "p2p-facts#approval-threshold  [citable]" in out
    assert "p2p-facts#callback-owner  [context]" in out
    assert "propose NEW entries in your status as `p2p-facts#<entry-id>` " \
           "with class + provenance" in out


def test_placement_brief_lists_structured_entries(tmp_path, capsys):
    root = make_root(tmp_path)
    make_brief_area(root)
    # placement needs 2+ areas
    area2 = root / "gl"
    area2.mkdir()
    (area2 / "manifest.json").write_text(json.dumps({
        "schema": "consult-mvp-manifest/v1", "area": "gl", "l1": "finance",
        "title": "GL", "l2_order": [], "components": []}),
        encoding="utf-8")
    seed_registers(root)
    capsys.readouterr()
    assert engagement.main(["brief", str(root)]) == 0
    out = capsys.readouterr().out
    assert "propose ENTRIES as `<register>#<entry-id>`" in out
    assert "provenance" in out
    assert "p2p-facts#approval-threshold  [citable]" in out
    assert "p2p-facts#callback-owner  [context]" in out


# --------------------------------------------------------------------------- #
# Render — citable appendix in, context NEVER
# --------------------------------------------------------------------------- #

def make_engagement_render_area(tmp_path: Path) -> tuple[Path, Path]:
    root = make_root(tmp_path)
    area = make_render_area(root)     # builds root/ap with a full manifest
    return root, area


def test_render_appends_citable_and_never_context(tmp_path):
    root, area = make_engagement_render_area(tmp_path)
    seed_registers(root)
    out = tmp_path / "d.docx"
    render.render_folder(area, out, emit_signal=False)
    text = doc_text(Document(str(out)))
    assert "Shared Reference — Engagement Registers" in text
    assert "P2P Facts" in text
    assert "approval-threshold — PO approval threshold is USD 10,000." \
        in text
    # context entries NEVER reach rendered output — not id, not text
    assert "callback-owner" not in text
    assert "Callback ownership" not in text
    # provenance is the entry's chain to the ledger, not deliverable content
    assert "SRC-004" not in text


def test_render_context_only_registers_add_no_appendix(tmp_path):
    root, area = make_engagement_render_area(tmp_path)
    assert add(root, eid="callback-owner", cls="context",
               text="Callback ownership is disputed.",
               provenance="SRC-002 (p2p)") == 0
    out = tmp_path / "d.docx"
    render.render_folder(area, out, emit_signal=False)
    text = doc_text(Document(str(out)))
    assert "Shared Reference" not in text
    assert "Callback ownership" not in text


def test_kit_subset_render_skips_the_appendix(tmp_path):
    """Kit docs are lean (same rule as the L2 dividers): the appendix rides
    only full folder renders."""
    root, area = make_engagement_render_area(tmp_path)
    seed_registers(root)
    out = tmp_path / "kit.docx"
    render.render_folder(area, out, slugs=["payment-run"],
                         emit_signal=False)
    text = doc_text(Document(str(out)))
    assert "Shared Reference" not in text


def test_final_mode_render_carries_the_citable_appendix(tmp_path):
    root, area = make_engagement_render_area(tmp_path)
    seed_registers(root)
    out = tmp_path / "final.docx"
    render.render_folder(area, out, mode="final", emit_signal=False)
    text = doc_text(Document(str(out)))
    assert "approval-threshold — PO approval threshold is USD 10,000." \
        in text
    assert "callback-owner" not in text
