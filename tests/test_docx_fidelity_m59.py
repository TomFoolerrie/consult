"""M59 gate — docx text fidelity: the builder stops eating prose.

The adversarial-review repros (F-04, F-09, F-20) through real docx builds.
docs/v2/M59-docx-text-fidelity.md is the spec.
"""

from pathlib import Path

import pytest

import docx
import cfgi_markdown_to_word as cfgi

from docx import Document  # noqa: E402
from docx.shared import Inches  # noqa: E402


def build(tmp_path, body_md: str) -> Document:
    out = tmp_path / "out.docx"
    cfgi.convert_assembled(body_md, out, title="T", subtitle="")
    return Document(str(out))


def all_text(doc) -> str:
    parts = [p.text for p in doc.paragraphs]
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                parts.append(cell.text)
    return "\n".join(parts)


class TestF04AngleBracketProse:
    def test_approval_ladder_renders_every_character(self, tmp_path):
        text = all_text(build(
            tmp_path,
            "## Approvals\n\nApproval ladder: <5k the cost-center owner "
            "approves; 5–25k adds the Controller; >25k routes to the CFO.\n"))
        assert "<5k" in text
        assert "5–25k" in text
        assert ">25k routes to the CFO" in text
        assert "the cost-center owner approves" in text

    def test_placeholder_and_comparison_survive(self, tmp_path):
        text = all_text(build(
            tmp_path, "## S\n\nEnter <client name> where debit < credit.\n"))
        assert "<client name>" in text
        assert "debit < credit" in text

    def test_entity_escaped_author_also_survives(self, tmp_path):
        text = all_text(build(tmp_path, "## S\n\nThreshold: &lt;5k only.\n"))
        assert "<5k" in text

    def test_whitelisted_tags_still_strip(self, tmp_path):
        text = all_text(build(
            tmp_path, "## S\n\n<b>bold</b> and <em>em</em> "
                      "and <!-- a comment --> prose.\n"))
        assert "<b>" not in text and "<em>" not in text
        assert "comment" not in text
        assert "bold" in text and "em and" in text and "prose." in text


class TestF09EscapesStayLiteral:
    def test_paired_escaped_asterisks_render_literal(self, tmp_path):
        doc = build(tmp_path, "## S\n\nMark \\*required\\* fields.\n")
        para = next(p for p in doc.paragraphs
                    if "required" in p.text)
        assert "*required*" in para.text
        assert not any(r.italic for r in para.runs), (
            "escaped markers re-styled as emphasis")

    def test_lone_escaped_marker_stays_correct(self, tmp_path):
        text = all_text(build(tmp_path, "## S\n\nuse \\*.xlsx for exports.\n"))
        assert "*.xlsx" in text

    def test_escaped_underscores_and_backslash(self, tmp_path):
        doc = build(tmp_path, "## S\n\nName \\_field\\_ and C:\\\\temp.\n")
        para = next(p for p in doc.paragraphs if "field" in p.text)
        assert "_field_" in para.text
        assert "C:\\temp" in para.text
        assert not any(r.italic for r in para.runs)

    def test_real_emphasis_still_works(self, tmp_path):
        doc = build(tmp_path, "## S\n\nThis is *actually italic* text.\n")
        para = next(p for p in doc.paragraphs if "actually" in p.text)
        assert any(r.italic and "actually italic" in r.text
                   for r in para.runs)


class TestF20NestedListDepth:
    def test_three_level_list_has_three_distinct_indents(self, tmp_path):
        doc = build(tmp_path,
                    "## S\n\n- top level\n  - second level\n"
                    "    - third level\n")
        by_text = {}
        for p in doc.paragraphs:
            if "level" in p.text:
                by_text[p.text] = p.paragraph_format.left_indent
        assert by_text["top level"] == Inches(0.25)
        assert by_text["second level"] == Inches(0.5)
        assert by_text["third level"] == Inches(0.75)

    def test_depth_is_capped(self, tmp_path):
        doc = build(tmp_path, "## S\n\n- a\n          - very deep\n")
        p = next(p for p in doc.paragraphs if p.text == "very deep")
        assert p.paragraph_format.left_indent == Inches(0.75)

    def test_flat_list_unchanged(self, tmp_path):
        doc = build(tmp_path, "## S\n\n- one\n- two\n")
        for t in ("one", "two"):
            p = next(p for p in doc.paragraphs if p.text == t)
            assert p.paragraph_format.left_indent == Inches(0.25)
