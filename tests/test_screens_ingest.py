"""Tests for scripts/screens_ingest.py — screenshot-template return-trip (M9).

Fixturing note: the "filled template" fixture is built the same way real ones
are — kits.write_screenshot_template creates the bookmarked paste boxes and
the _review/.maps sidecar, then python-docx pastes an image into a box cell,
exactly what a reviewer's paste produces (a blip inside the bookmarked cell).
"""
import base64
import json
from pathlib import Path

import screens_ingest
from kits import write_screenshot_template

from docx import Document

PNG_1PX = base64.b64decode(
    "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAYAAAAfFcSJ"
    "AAAADUlEQVR42mP8z8BQDwAEhQGAhKmMIQAAAABJRU5ErkJggg=="
)

ITEMS = [
    {"slug": "vendor-onboarding", "local": "SC-01", "disp": "SC-01",
     "text": "Vendor master entry screen.", "step": "Step 1 - Enter vendor",
     "proc_label": "1.1 Vendor Onboarding"},
    {"slug": "payment-run", "local": "SC-01", "disp": "SC-02",
     "text": "Payment proposal screen.", "step": "",
     "proc_label": "2.1 Payment Run"},
]


def make_template(tmp_path: Path, fill_first: bool = False):
    """Area + template docx + sidecar map, optionally with box 0 filled."""
    area = tmp_path / "ap"
    (area / "_review" / "returned").mkdir(parents=True)
    maps_dir = area / "_review" / ".maps"
    maps_dir.mkdir()
    path = area / "_review" / "returned" / "screenshots_jane.docx"
    smap = write_screenshot_template(path, "Jane Junior", ITEMS, "AP Process")
    (maps_dir / f"{smap['doc_id']}.json").write_text(
        json.dumps(smap), encoding="utf-8")
    if fill_first:
        png = tmp_path / "shot.png"
        png.write_bytes(PNG_1PX)
        doc = Document(str(path))
        # paste into the first dashed box (tables are one per SC item)
        cell = doc.tables[0].cell(0, 0)
        cell.paragraphs[0].add_run().add_picture(str(png))
        doc.save(str(path))
    return area, path


def test_empty_boxes_reported_never_guessed(tmp_path):
    """A returned template with no pasted images reports every box as empty
    (by display id) and saves nothing under _assets/screens/."""
    area, path = make_template(tmp_path)
    res = screens_ingest.ingest_template(area, path)
    assert res["saved"] == []
    assert res["empty"] == ["SC-01", "SC-02"]
    assert res["stray"] == 0
    assert not (area / "_assets").exists()


def test_pasted_image_saved_to_assets_by_local_id(tmp_path):
    """An image pasted in a bookmarked box is extracted to
    _assets/screens/<slug>/<LOCAL-ID>.png — where render --mode final looks."""
    area, path = make_template(tmp_path, fill_first=True)
    res = screens_ingest.ingest_template(area, path)
    assert res["saved"] == ["vendor-onboarding/SC-01.png"]
    assert res["empty"] == ["SC-02"]
    dest = area / "_assets" / "screens" / "vendor-onboarding" / "SC-01.png"
    assert dest.read_bytes() == PNG_1PX


def test_non_template_docx_is_skipped(tmp_path):
    """A docx without the cw-screens core-property stamp is skipped, not
    ingested (returned/ also carries reviewed procedure docs)."""
    area = tmp_path / "ap"
    (area / "_review" / "returned").mkdir(parents=True)
    other = area / "_review" / "returned" / "reviewed.docx"
    d = Document()
    d.add_paragraph("just a reviewed doc")
    d.save(str(other))
    res = screens_ingest.ingest_template(area, other)
    assert "skipped" in res


def test_template_with_missing_map_is_skipped(tmp_path):
    """A stamped template whose sidecar map is gone from _review/.maps is
    skipped rather than guessed at."""
    area, path = make_template(tmp_path)
    for f in (area / "_review" / ".maps").glob("*.json"):
        f.unlink()
    res = screens_ingest.ingest_template(area, path)
    assert "skipped" in res


def test_main_sweeps_returned_and_archives(tmp_path, capsys):
    """The returned/ sweep ingests self-identifying templates, archives them
    to _review/processed/, and leaves non-template docx alone."""
    area, path = make_template(tmp_path, fill_first=True)
    other = area / "_review" / "returned" / "reviewed.docx"
    d = Document()
    d.add_paragraph("reviewed procedure doc")
    d.save(str(other))
    assert screens_ingest.main([str(area)]) == 0
    out = capsys.readouterr().out
    assert "saved _assets/screens/vendor-onboarding/SC-01.png" in out
    assert "empty box(es), not returned: SC-02" in out
    assert not path.exists()
    assert (area / "_review" / "processed" / path.name).is_file()
    assert other.exists()   # non-template stays for the review-extract script


def test_main_no_archive(tmp_path):
    """--no-archive ingests but leaves the template in returned/."""
    area, path = make_template(tmp_path, fill_first=True)
    assert screens_ingest.main([str(area), "--no-archive"]) == 0
    assert path.exists()
    assert (area / "_assets" / "screens" / "vendor-onboarding" /
            "SC-01.png").is_file()


def test_main_nothing_to_ingest(tmp_path, capsys):
    """No returned files at all exits 0 with a message."""
    area = tmp_path / "ap"
    (area / "_review" / "returned").mkdir(parents=True)
    assert screens_ingest.main([str(area)]) == 0
    assert "no templates to ingest" in capsys.readouterr().out
