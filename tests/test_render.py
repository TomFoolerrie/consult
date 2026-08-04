"""Tests for scripts/render.py and the bundled converter
skills/consult-docx-builder/scripts/cfgi_markdown_to_word.py.

All fixtures are synthetic areas built under tmp_path (see conftest.py).
`make_area` is also imported by test_kits.py so the two suites exercise the
same synthetic area shape.
"""
import base64
import json
import zipfile
from pathlib import Path

import pytest
from docx import Document
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.oxml.ns import qn

import cfgi_markdown_to_word as cfgi
import render


# 1x1 transparent PNG.
PNG_1PX = base64.b64decode(
    "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAYAAAAfFcSJ"
    "AAAADUlEQVR42mP8z8BQDwAEhQGAhKmMIQAAAABJRU5ErkJggg=="
)

MANIFEST = {
    "schema": "consult-mvp-manifest/v1",
    "area": "ap",
    "l1": "finance",
    "title": "Accounts Payable Process",
    "subtitle": "Current State Documentation",
    "l2_order": ["invoices", "payments"],
    "components": [
        {"file": "00_document-profile.md", "heading": "Document Profile",
         "role": "static", "order": 0},
        {"file": "10_vendor-onboarding.md", "heading": "Vendor Onboarding",
         "role": "procedure", "slug": "vendor-onboarding", "l2": "invoices",
         "order": 10},
        {"file": "20_payment-run.md", "heading": "Payment Run",
         "role": "procedure", "slug": "payment-run", "l2": "payments",
         "order": 20},
        {"file": "30_cash-application.md", "heading": "Cash Application",
         "role": "procedure", "slug": "cash-application", "l2": "payments",
         "order": 30},
        {"file": "90_appendix-b-gaps.md", "heading": "Appendix B - Gap Log",
         "role": "derived", "derived_kind": "gap-log", "writer": "python",
         "order": 90},
    ],
}

PROFILE_MD = """## Document Profile

| Field | Value |
|---|---|
| Client / Organization | Acme Corp |
| Version | v1.0 |
"""

VENDOR_MD = """## Vendor Onboarding

### A. Process Overview

New vendors are set up before the first invoice; see [[payment-run]].

### B. Quick Reference

- **Preparer:** AP Clerk
- **Reviewer:** Controller

### C. Pre-Requisites

- Approved vendor request form

### D. Inputs

- Vendor W-9

### E. Step-by-Step Procedure

#### Step 1 - Enter vendor

1. Open the vendor portal.
2. Enter the vendor master details.

> **VALIDATION REQUIRED - GAP-01:** Confirm the approval threshold amount.
> - **Nature:** Missing detail
> - **Owner to confirm:** AP Clerk

> **SCREENSHOT PLACEHOLDER - SC-01:** Vendor master entry screen.
> - **System:** SAP

The threshold is applied here [[GAP-01 - threshold unknown]].

#### Step 2 - Approve

1. Route the vendor for approval.
2. Confirm activation.

### F. Key Controls

> **CONTROL - CTRL-01:** Vendor changes require dual approval.
> - **Control type:** Preventive

### G. Outputs

- Active vendor master record

### H. Known Issues / Improvement Notes

- None noted.
"""

PAYMENT_MD = """## Payment Run

### A. Process Overview

Weekly payment run for approved invoices.

### B. Quick Reference

- **Preparer:** Controller

### E. Step-by-Step Procedure

#### Step 1 - Select invoices

1. Build the payment proposal.
2. Review exceptions.

> **VALIDATION REQUIRED - GAP-01:** Confirm the bank cut-off time.
> - **Nature:** Missing detail
> - **Owner to confirm:** AP Clerk

> **SCREENSHOT PLACEHOLDER - SC-01:** Payment proposal screen.

### G. Outputs

- Payment file
"""

CASH_MD = """## Cash Application

### A. Process Overview

Cash receipts are applied daily.

### B. Quick Reference

- **Preparer:** Mystery Person

### E. Step-by-Step Procedure

#### Step 1 - Apply cash

1. Match receipts to open invoices.
"""

GAPLOG_MD = """## Appendix B - Gap Log

<!-- derived: gap-log; writer: python -->

| Gap | Procedure |
|---|---|
| GAP-01 | [[#vendor-onboarding]] |
"""

ROLES_YAML = """roles:
  - slug: ap-clerk
    name: AP Clerk
    people:
      - Mark Manager
      - Jane Junior
  - slug: controller
    name: Controller
    people:
      - Carol Chief
"""

ORG_YAML = """people:
  - name: Carol Chief
    title: Controller
  - name: Mark Manager
    title: AP Manager
    reports_to: Carol Chief
  - name: Jane Junior
    title: AP Clerk
    reports_to: Mark Manager
"""


def make_area(tmp_path: Path) -> Path:
    """Build a complete synthetic area folder under tmp_path."""
    area = tmp_path / "ap"
    area.mkdir()
    (area / "manifest.json").write_text(json.dumps(MANIFEST, indent=1),
                                        encoding="utf-8")
    (area / "00_document-profile.md").write_text(PROFILE_MD, encoding="utf-8")
    (area / "10_vendor-onboarding.md").write_text(VENDOR_MD, encoding="utf-8")
    (area / "20_payment-run.md").write_text(PAYMENT_MD, encoding="utf-8")
    (area / "30_cash-application.md").write_text(CASH_MD, encoding="utf-8")
    (area / "90_appendix-b-gaps.md").write_text(GAPLOG_MD, encoding="utf-8")
    ref = area / "_reference"
    ref.mkdir()
    (ref / "roles.yaml").write_text(ROLES_YAML, encoding="utf-8")
    client = area / "_client"
    client.mkdir()
    (client / "org-chart.yaml").write_text(ORG_YAML, encoding="utf-8")
    return area


def doc_text(doc) -> str:
    parts = [p.text for p in doc.paragraphs]
    for t in doc.tables:
        for row in t.rows:
            for c in row.cells:
                parts.append(c.text)
    return "\n".join(parts)


def bookmark_names(doc) -> list[str]:
    return [bm.get(qn("w:name")) or ""
            for bm in doc.element.body.iter(qn("w:bookmarkStart"))]


# --------------------------------------------------------------------------- #
# working mode
# --------------------------------------------------------------------------- #

def test_working_render_produces_docx_with_provenance(tmp_path):
    """Working mode writes the docx, per-paragraph cw_ bookmarks, and a
    _review/.maps sidecar keyed by the doc id stamped into core properties."""
    area = make_area(tmp_path)
    out = tmp_path / "draft.docx"
    stats = render.render_folder(area, out, emit_signal=False)
    assert out.is_file()
    doc = Document(str(out))
    cw = [n for n in bookmark_names(doc) if n.startswith("cw_")]
    assert cw, "working mode must stamp cw_ provenance bookmarks"
    map_path = Path(stats["map"])
    assert map_path.parent == area / "_review" / ".maps"
    payload = json.loads(map_path.read_text(encoding="utf-8"))
    assert payload["schema"] == "consult-review-map/v1"
    assert payload["doc_id"] == stats["doc_id"]
    assert doc.core_properties.category == f"cw-map:{stats['doc_id']}"
    # every bookmark has a map entry pointing at a real fragment file
    assert set(cw) <= set(payload["entries"])
    files = {e["file"] for e in payload["entries"].values()}
    assert "10_vendor-onboarding.md" in files


def test_working_render_emits_signal_file(tmp_path):
    """Full working render (emit_signal=True) writes {area}/.render.json."""
    area = make_area(tmp_path)
    out = tmp_path / "draft.docx"
    render.render_folder(area, out, emit_signal=True)
    sig = json.loads((area / ".render.json").read_text(encoding="utf-8"))
    assert sig["docx"] == str(out)
    assert sig["awaiting_review"] is True


def test_no_signal_when_emit_signal_false(tmp_path):
    """emit_signal=False (programmatic/kit call) never writes .render.json."""
    area = make_area(tmp_path)
    render.render_folder(area, tmp_path / "d.docx", emit_signal=False)
    assert not (area / ".render.json").exists()


def test_working_render_shows_gaps_and_display_ids(tmp_path):
    """Working mode keeps VALIDATION REQUIRED callouts visible and rewrites
    procedure-local callout ids to global display ids (payment-run GAP-01 ->
    GAP-02)."""
    area = make_area(tmp_path)
    out = tmp_path / "draft.docx"
    render.render_folder(area, out, emit_signal=False)
    text = doc_text(Document(str(out)))
    assert "VALIDATION REQUIRED" in text
    assert "GAP-02" in text        # payment-run's local GAP-01, renumbered
    assert "SC-02" in text
    # cross-reference token resolved to number + title
    assert "2.1 Payment Run" in text
    # numbered procedure headings from the ONE display-number map
    assert any(p.text.startswith("1.1 Vendor Onboarding")
               for p in Document(str(out)).paragraphs)


# --------------------------------------------------------------------------- #
# final mode
# --------------------------------------------------------------------------- #

def test_final_mode_strips_gaps_and_reports_counts(tmp_path):
    """Final mode strips VR callouts, inline [[GAP-..]] tags, and the gap-log
    appendix; counts are reported and rendering never refuses."""
    area = make_area(tmp_path)
    out = tmp_path / "final.docx"
    stats = render.render_folder(area, out, mode="final", emit_signal=False)
    assert stats["gaps_stripped"] == 2
    assert stats["gap_tags_stripped"] == 1
    text = doc_text(Document(str(out)))
    assert "VALIDATION REQUIRED" not in text
    assert "[[GAP" not in text and "[GAP-" not in text
    assert "Gap Log" not in text                    # derived gap-log dropped
    assert "CTRL-01" in text                        # controls survive


def test_final_mode_reports_dangling_prose_gap_references(tmp_path):
    """A free-prose mention of a gap id ("... see GAP-01") survives the final
    strip — its definition (the callout and the gap-log row) is exactly what
    final mode removes — so the render counts and enumerates it per procedure
    instead of shipping the dangling reference silently. The prose itself is
    left untouched (no mechanical rewriting of deliverable wording), and
    working mode reports nothing: there the reference points at a callout the
    reader still has."""
    area = make_area(tmp_path)
    pay = area / "20_payment-run.md"
    pay.write_text(pay.read_text(encoding="utf-8").replace(
        "- Payment file",
        "- Payment file (cut-off unconfirmed - see GAP-01)"), encoding="utf-8")

    working = render.render_folder(area, tmp_path / "working.docx",
                                   emit_signal=False)
    assert working["dangling_gap_refs"] == {}
    assert working["dangling_gap_ref_count"] == 0

    out = tmp_path / "final.docx"
    stats = render.render_folder(area, out, mode="final", emit_signal=False)
    # payment-run's local GAP-01 is display GAP-02 (vendor-onboarding's gap
    # takes GAP-01): the detector sees what the READER sees.
    assert stats["dangling_gap_refs"] == {"payment-run": ["GAP-02"]}
    assert stats["dangling_gap_ref_count"] == 1
    text = doc_text(Document(str(out)))
    assert "see GAP-02" in text
    assert "VALIDATION REQUIRED" not in text


def test_final_mode_scrubs_unambiguous_citations(tmp_path):
    """Final mode removes the two mechanically-safe citation shapes — a
    parenthetical containing nothing but SRC/GAP ids, and a pure-citation
    "See GAP-##." sentence — and leaves ids woven into sentence meaning for
    the dangling-reference warning (now covering SRC too). Working mode is
    untouched: citations are the drafters' provenance."""
    area = make_area(tmp_path)
    pay = area / "20_payment-run.md"
    pay.write_text(pay.read_text(encoding="utf-8").replace(
        "- Payment file",
        "- The run is cut Thursday (SRC-002, SRC-005). See GAP-01.\n"
        "- Approver disputed (SRC-004; see GAP-01, which is unresolved)\n"
        "- Payment file"), encoding="utf-8")

    working = render.render_folder(area, tmp_path / "w.docx",
                                   emit_signal=False)
    assert working["citations_scrubbed"] == 0
    wtext = doc_text(Document(str(tmp_path / "w.docx")))
    assert "(SRC-002, SRC-005)" in wtext

    stats = render.render_folder(area, tmp_path / "f.docx", mode="final",
                                 emit_signal=False)
    assert stats["citations_scrubbed"] == 2
    text = doc_text(Document(str(tmp_path / "f.docx")))
    assert "SRC-002" not in text and "See GAP" not in text
    assert "The run is cut Thursday." in text     # clean seam, no double space
    assert "SRC-004" in text                      # woven — left for a human
    assert "SRC-004" in stats["dangling_gap_refs"]["payment-run"]


def test_final_mode_placeholder_kept_when_no_image(tmp_path):
    """Final mode keeps the SCREENSHOT PLACEHOLDER callout when no captured
    image exists under _assets/screens/<slug>/."""
    area = make_area(tmp_path)
    out = tmp_path / "final.docx"
    stats = render.render_folder(area, out, mode="final", emit_signal=False)
    assert stats["screens_embedded"] == 0
    assert stats["screens_placeholder"] == 2
    assert "SCREENSHOT PLACEHOLDER" in doc_text(Document(str(out)))


def test_final_mode_embeds_captured_screenshot_with_caption(tmp_path):
    """Final mode swaps a placeholder for the image at
    _assets/screens/<slug>/<LOCAL-ID>.png, with an italic caption below."""
    area = make_area(tmp_path)
    shot = area / "_assets" / "screens" / "vendor-onboarding" / "SC-01.png"
    shot.parent.mkdir(parents=True)
    shot.write_bytes(PNG_1PX)
    out = tmp_path / "final.docx"
    stats = render.render_folder(area, out, mode="final", emit_signal=False)
    assert stats["screens_embedded"] == 1
    assert stats["screens_placeholder"] == 1     # payment-run's is still open
    doc = Document(str(out))
    assert len(doc.inline_shapes) == 1
    captions = [p for p in doc.paragraphs
                if "SC-01:" in p.text and "Vendor master entry screen" in p.text]
    assert captions and all(r.italic for r in captions[0].runs)


# --------------------------------------------------------------------------- #
# subset (--slugs) render
# --------------------------------------------------------------------------- #

def test_subset_unknown_slug_raises_systemexit(tmp_path):
    """--slugs with an unknown procedure slug exits with an error."""
    area = make_area(tmp_path)
    with pytest.raises(SystemExit, match="unknown procedure slug"):
        render.render_folder(area, tmp_path / "x.docx",
                             slugs=["nope"], emit_signal=False)


def test_subset_render_matches_full_numbering(tmp_path):
    """A kit subset render carries the FULL manifest's display numbers and
    global callout display ids, contains only the requested procedure, and
    never writes the .render.json signal."""
    area = make_area(tmp_path)
    out = tmp_path / "kit.docx"
    render.render_folder(area, out, slugs=["payment-run"], emit_signal=True)
    assert not (area / ".render.json").exists()
    text = doc_text(Document(str(out)))
    assert "2.1 Payment Run" in text     # not renumbered to 1.1
    assert "GAP-02" in text              # global display id, not local GAP-01
    assert "Vendor Onboarding" not in text
    assert "Document Profile" not in text    # no cover/front matter


# --------------------------------------------------------------------------- #
# converter: ordered-list numbering
# --------------------------------------------------------------------------- #

def _num_ids(doc):
    """[(paragraph, numId)] for List Number paragraphs, in document order."""
    out = []
    for p in doc.paragraphs:
        if p.style.name != "List Number":
            continue
        numPr = p._p.find(qn("w:pPr") + "/" + qn("w:numPr"))
        if numPr is None:
            pPr = p._p.find(qn("w:pPr"))
            numPr = pPr.find(qn("w:numPr")) if pPr is not None else None
        nid = numPr.find(qn("w:numId")) if numPr is not None else None
        out.append((p, nid.get(qn("w:val")) if nid is not None else None))
    return out


def test_two_ordered_lists_restart_numbering(tmp_path):
    """Two separate '1.'-starting lists get DIFFERENT numbering instances
    (shared within a list) and each <w:num> carries a startOverride, so the
    second list restarts at 1 instead of numbering continuously."""
    md = tmp_path / "lists.md"
    md.write_text(
        "## Section\n\nFirst list:\n\n1. alpha\n2. beta\n\n"
        "Interlude paragraph.\n\n1. gamma\n2. delta\n",
        encoding="utf-8")
    out = tmp_path / "lists.docx"
    cfgi.convert(md, out, include_toc=False, landscape=False, do_cover=False)
    doc = Document(str(out))
    nums = _num_ids(doc)
    assert [p.text for p, _ in nums] == ["alpha", "beta", "gamma", "delta"]
    ids = [nid for _, nid in nums]
    assert all(ids), "every ordered item must carry an explicit numId"
    assert ids[0] == ids[1], "items of one list share a numbering instance"
    assert ids[2] == ids[3]
    assert ids[0] != ids[2], "separate lists must NOT share a numId"
    # numbering part: each allocated <w:num> has an ilvl-0 startOverride
    numbering = doc.part.part_related_by(RT.NUMBERING).element
    by_id = {n.get(qn("w:numId")): n for n in numbering.findall(qn("w:num"))}
    for nid in {ids[0], ids[2]}:
        num = by_id[nid]
        so = num.find(qn("w:lvlOverride") + "/" + qn("w:startOverride"))
        if so is None:
            ov = num.find(qn("w:lvlOverride"))
            so = ov.find(qn("w:startOverride")) if ov is not None else None
        assert so is not None and so.get(qn("w:val")) == "1"


def test_repeated_one_idiom_is_one_continuous_list(tmp_path):
    """CommonMark semantics: a contiguous run numbered '1. / 1. / 1.' is ONE
    list (renders 1, 2, 3) — no per-item restart."""
    md = tmp_path / "ones.md"
    md.write_text("## S\n\n1. alpha\n1. beta\n1. gamma\n", encoding="utf-8")
    out = tmp_path / "ones.docx"
    cfgi.convert(md, out, include_toc=False, landscape=False, do_cover=False)
    nums = _num_ids(Document(str(out)))
    ids = [nid for _, nid in nums]
    assert len(ids) == 3 and len(set(ids)) == 1, \
        "contiguous ordered items must share one numbering instance"


def test_blank_separated_restart_seeds_literal_number(tmp_path):
    """An ordered item after a break restarts a fresh instance seeded with its
    own literal number, so displayed numbers always match the source."""
    md = tmp_path / "seed.md"
    md.write_text("## S\n\n1. alpha\n2. beta\n\n3. gamma\n", encoding="utf-8")
    out = tmp_path / "seed.docx"
    cfgi.convert(md, out, include_toc=False, landscape=False, do_cover=False)
    doc = Document(str(out))
    nums = _num_ids(doc)
    ids = [nid for _, nid in nums]
    assert ids[0] == ids[1] != ids[2]
    numbering = doc.part.part_related_by(RT.NUMBERING).element
    by_id = {n.get(qn("w:numId")): n for n in numbering.findall(qn("w:num"))}
    so = by_id[ids[2]].find(qn("w:lvlOverride") + "/" + qn("w:startOverride"))
    assert so is not None and so.get(qn("w:val")) == "3"


# --------------------------------------------------------------------------- #
# converter: tracked changes
# --------------------------------------------------------------------------- #

def test_enable_track_changes_schema_position_and_lock(tmp_path):
    """enable_track_changes puts <w:trackChanges/> at a schema-valid position
    (before any later-sequenced settings child) and adds
    <w:documentProtection w:edit="trackedChanges" w:enforcement="1"/>."""
    out = tmp_path / "tc.docx"
    cfgi.convert_assembled("## S\n\nBody paragraph.\n", out,
                           title="T", subtitle="", do_cover=False,
                           track_changes=True)
    doc = Document(str(out))
    settings = doc.settings.element
    names = [c.tag.rsplit("}", 1)[-1] for c in settings]
    assert "trackChanges" in names
    tc_i = names.index("trackChanges")
    # nothing sequenced AFTER trackChanges may appear before it
    assert not any(n in cfgi._SETTINGS_AFTER_TRACK_CHANGES
                   for n in names[:tc_i])
    dp = settings.find(qn("w:documentProtection"))
    assert dp is not None
    assert dp.get(qn("w:edit")) == "trackedChanges"
    assert dp.get(qn("w:enforcement")) == "1"
    assert names.index("documentProtection") > tc_i


def test_track_changes_survives_zip_roundtrip(tmp_path):
    """The trackChanges/documentProtection elements are actually persisted in
    the saved package's settings.xml."""
    out = tmp_path / "tc.docx"
    cfgi.convert_assembled("## S\n\nBody.\n", out, title="T", subtitle="",
                           do_cover=False, track_changes=True)
    with zipfile.ZipFile(out) as z:
        xml = z.read("word/settings.xml").decode("utf-8")
    assert "trackChanges" in xml and "documentProtection" in xml


# --------------------------------------------------------------------------- #
# CLI / input resolution
# --------------------------------------------------------------------------- #

def test_main_rejects_final_flags_for_single_file(tmp_path):
    """--mode/--slugs/--track-changes require an area folder, not a lone .md."""
    md = tmp_path / "one.md"
    md.write_text("# T\n\nBody.\n", encoding="utf-8")
    with pytest.raises(SystemExit, match="require an area folder"):
        render.main([str(md), "-o", str(tmp_path / "o.docx"),
                     "--mode", "final"])


def test_main_missing_input_errors(tmp_path):
    """A nonexistent input path exits with an error."""
    with pytest.raises(SystemExit, match="input not found"):
        render.main([str(tmp_path / "nowhere")])


def test_folder_without_manifest_rejected(tmp_path):
    """A directory lacking manifest.json is not a valid area."""
    d = tmp_path / "empty"
    d.mkdir()
    with pytest.raises(SystemExit, match="no manifest.json"):
        render.main([str(d)])


@pytest.mark.parametrize("n", [2, 5, 11, 12, 17, 30])
def test_column_widths_stay_positive_and_fill_usable(n):
    """Wide tables (RACI runs one column per role) must not yield negative widths.

    With more columns than the 0.6" floor fits across the page, the floor has to
    yield to an even share: python-docx rejects a negative w:gridCol outright.
    """
    from docx.shared import Inches

    usable = int(Inches(6.5))
    tb = cfgi.TableBlock(
        header=[f"Role {i}" for i in range(n)],
        rows=[["R"] * n, ["A"] * n],
    )
    widths = cfgi._column_widths(tb, n, usable)

    assert len(widths) == n
    assert all(w > 0 for w in widths), f"non-positive width for n={n}: {widths}"
    assert sum(widths) == usable


def test_l2_chapter_dividers_open_each_subprocess_on_a_fresh_page(tmp_path):
    """A folder render emits one `# {ordinal}. {Title}` divider before the
    first procedure of each L2 bucket, plus `Reference & Appendices` before
    the back matter — each a Heading 1 with page-break-before. Display glue
    only: no fragment, manifest, or provenance change."""
    area = make_area(tmp_path)
    out = tmp_path / "d.docx"
    render.render_folder(area, out, emit_signal=False)
    h1 = [p for p in Document(str(out)).paragraphs
          if p.style.name == "Heading 1"]
    # This fixture's only front matter is the cover-lifted Document Profile,
    # so no Introduction chapter — an empty chapter head would be worse than
    # none.
    assert [p.text for p in h1] == ["Document Control", "1. Invoices",
                                    "2. Payments", "Reference & Appendices"]
    assert all(p.paragraph_format.page_break_before for p in h1)
    # Chapter weight must reach the RUNS: styled_run stamps body-default
    # direct formatting that overrides the paragraph style, so a style-table
    # change alone renders 10pt black (the "flat H1" review note).
    run = h1[0].runs[0]
    assert run.font.size.pt == 20 and run.font.bold
    # Same trap one level down: H4 (procedure steps, appendix process
    # buckets) was absent from HEADING_RUN_FMT and rendered as flat body
    # text (the "sub-heading visibility" review note).
    h4 = [p for p in Document(str(out)).paragraphs
          if p.style.name == "Heading 4"]
    assert h4, "fixture has #### step headings"
    run4 = h4[0].runs[0]
    assert run4.font.size.pt == 11 and run4.font.bold

    # With the cover off the profile renders inline as front matter, and the
    # Introduction chapter appears above it.
    out2 = tmp_path / "d2.docx"
    render.render_folder(area, out2, do_cover=False, emit_signal=False)
    texts = [p.text for p in Document(str(out2)).paragraphs
             if p.style.name == "Heading 1"]
    assert texts == ["Introduction", "1. Invoices", "2. Payments",
                     "Reference & Appendices"]


def test_toc_title_is_not_a_heading_so_the_toc_cannot_list_itself(tmp_path):
    """The TOC field collects Heading 1-3; a Heading-styled page title makes
    the table list itself as its first entry. The title is a direct-formatted
    Normal paragraph instead — same chapter look, invisible to the field."""
    area = make_area(tmp_path)
    out = tmp_path / "t.docx"
    # No flag passed: folder renders carry the TOC unconditionally.
    render.render_folder(area, out, emit_signal=False)
    toc_title = [p for p in Document(str(out)).paragraphs
                 if p.text == "Table of Contents"]
    assert len(toc_title) == 1
    assert toc_title[0].style.name == "Normal"
    run = toc_title[0].runs[0]
    assert run.font.size.pt == 20 and run.font.bold


def test_subset_kit_render_has_no_dividers(tmp_path):
    """Kit docs are lean per-owner excerpts: a chapter head over a single
    excerpted procedure is noise, so subset renders carry no H1 at all."""
    area = make_area(tmp_path)
    out = tmp_path / "k.docx"
    render.render_folder(area, out, slugs=["payment-run"], emit_signal=False)
    doc = Document(str(out))
    assert not [p for p in doc.paragraphs if p.style.name == "Heading 1"]
    assert not [p for p in doc.paragraphs if p.text == "Table of Contents"]


def test_callout_color_comes_from_the_label_not_the_prose(tmp_path):
    """A CONTROL whose prose cites a gap ("see GAP-07") stays control-green,
    and a VALIDATION REQUIRED stays gap-yellow. The old whole-text keyword
    scan recolored boxes by whatever their prose happened to mention — GAP
    outranks CONTROL in the cascade, so control boxes turned yellow."""
    md = tmp_path / "c.md"
    md.write_text(
        "# T\n\n## Key Controls\n\n"
        "> **CONTROL — CTRL-001:** Dual approval of banking changes.\n"
        "> - **Owner:** TBD — the performing role is contested; see GAP-07\n"
        "\n"
        "> **VALIDATION REQUIRED — GAP-07:** Who performs the callback.\n"
        "> - **Owner to confirm:** Controller\n",
        encoding="utf-8")
    out = tmp_path / "c.docx"
    cfgi.main([str(md), "-o", str(out)])
    fills = []
    for t in Document(str(out)).tables:
        shd = t.cell(0, 0)._tc.get_or_add_tcPr().find(qn("w:shd"))
        fills.append(shd.get(qn("w:fill")))
    assert fills == ["F3F8F4", "FCF7CC"]   # control green, then gap yellow


def test_paired_pain_improvement_table_is_not_red(tmp_path):
    """The Appendix A register pairs 'Pain Point' and 'Improvement
    Opportunity' in ONE header — the pain-table red wash must not fire on it
    (a red recommendation reads as a problem). A pain-only table keeps the
    red styling."""
    md = tmp_path / "a.md"
    md.write_text(
        "# T\n\n## Appendix\n\n"
        "| Pain Point | Impact | Severity | Improvement Opportunity |\n"
        "|---|---|---|---|\n"
        "| PP-01 — manual matching | Two days | High | IO-01 — automate |\n"
        "\n\n## Known Issues\n\n"
        "| Pain Point | Impact |\n"
        "|---|---|\n"
        "| PP-02 — re-keying | An hour a day |\n",
        encoding="utf-8")
    out = tmp_path / "a.docx"
    cfgi.main([str(md), "-o", str(out)])
    body_fills = []
    for t in Document(str(out)).tables:
        shd = t.cell(1, 0)._tc.get_or_add_tcPr().find(qn("w:shd"))
        body_fills.append(shd.get(qn("w:fill")))
    assert body_fills[0] == "FFFFFF"       # paired register: standard
    assert body_fills[1] == "FBEBEB"       # pain-only table: still red


def test_wide_raci_table_renders(tmp_path):
    """A 17-column RACI matrix converts without a width error."""
    md = tmp_path / "wide.md"
    roles = [f"Role {i}" for i in range(16)]
    header = "| Activity | " + " | ".join(roles) + " |"
    sep = "|" + "---|" * 17
    row = "| Weekly Payment Run | " + " | ".join(["R"] * 16) + " |"
    md.write_text(f"# T\n\n## RACI Matrix\n\n{header}\n{sep}\n{row}\n", encoding="utf-8")

    out = tmp_path / "wide.docx"
    cfgi.main([str(md), "-o", str(out)])

    doc = Document(str(out))
    assert doc.tables, "no table rendered"
    assert all(c.width > 0 for c in doc.tables[0].columns)


def test_raci_rows_render_in_display_number_order(tmp_path):
    """Reviewer ask: RACI activities appear sequentially (1.1, 1.2, … 2.1).
    Ordering is render-time display keyed on each row's leading [[slug]]
    token — the derived file's authored order is never rewritten. A row
    without a resolvable slug leaves the authored order untouched."""
    area = make_area(tmp_path)
    manifest = json.loads((area / "manifest.json").read_text(encoding="utf-8"))
    manifest["components"].append(
        {"file": "84_raci.md", "heading": "RACI Matrix", "role": "derived",
         "derived_kind": "raci", "writer": "agent", "order": 84})
    (area / "manifest.json").write_text(json.dumps(manifest), encoding="utf-8")
    (area / "84_raci.md").write_text(
        "## RACI Matrix\n\n"
        "| Activity | Responsible | Accountable | Consulted | Informed |\n"
        "|---|---|---|---|---|\n"
        "| [[cash-application]] | AR Clerk | Controller | — | — |\n"
        "| [[vendor-onboarding]] | AP Clerk | Controller | — | — |\n"
        "| [[payment-run]] | AP Clerk | Controller | — | — |\n",
        encoding="utf-8")

    out = tmp_path / "d.docx"
    render.render_folder(area, out, emit_signal=False)
    doc = Document(str(out))
    raci = next(t for t in doc.tables
                if t.rows[0].cells[0].text.strip() == "Activity")
    acts = [r.cells[0].text for r in raci.rows[1:]]
    assert acts == ["1.1 Vendor Onboarding", "2.1 Payment Run",
                    "2.2 Cash Application"]

    # authored file untouched (display-only ordering)
    assert (area / "84_raci.md").read_text(encoding="utf-8").split("\n")[4] \
        .startswith("| [[cash-application]]")


def test_escaped_pipe_in_table_cell_does_not_shear_the_row(tmp_path):
    """A literal pipe in cell text is written `\\|` (aggregate's cell()); the
    converter must split on UNESCAPED pipes only. Splitting on the escape
    slid every later cell one column right — a phantom fourth column and a
    stray backslash in the cell (the systems-view quirk)."""
    md = tmp_path / "t.md"
    md.write_text(
        "# T\n\n## Systems\n\n"
        "| System / Tool | Role in Process | Related Procedures |\n"
        "|---|---|---|\n"
        "| SAP S/4HANA | ERP (SAP\\|S4 job schedule) | 1.1, 1.2 |\n",
        encoding="utf-8")
    out = tmp_path / "t.docx"
    cfgi.main([str(md), "-o", str(out)])
    t = Document(str(out)).tables[0]
    assert len(t.columns) == 3
    assert t.cell(1, 1).text == "ERP (SAP|S4 job schedule)"
    assert t.cell(1, 2).text == "1.1, 1.2"


# --------------------------------------------------------------------------- #
# v1.18 — TOC depth/refresh, Document Control, final-mode register grooming,
# lexicon, readiness scorecard, back-matter page breaks
# --------------------------------------------------------------------------- #

ROLE_DICT_MD = """## Role Dictionary

<!-- derived: role-dictionary; writer: python -->

_Canonical functional roles with the procedures each appears in._

| Functional Role | Reports To | Standard Responsibilities |
|---|---|---|
| AP Clerk |  | Enters invoices |
"""

SCREENS_IDX_MD = """## Appendix C - Screenshots

<!-- derived: screenshot-index; writer: python -->

_Index of the screenshots referenced in this document._

| SC ID | Caption | Status |
|---|---|---|
| SC-01 | Vendor screen | Pending user input |
"""


def make_area_with_registers(tmp_path: Path) -> Path:
    """make_area plus a role-dictionary and a screenshot-index register."""
    area = make_area(tmp_path)
    man = json.loads((area / "manifest.json").read_text(encoding="utf-8"))
    man["components"] += [
        {"file": "80_roles-dict.md", "heading": "Role Dictionary",
         "role": "derived", "derived_kind": "role-dictionary",
         "writer": "python", "order": 80},
        {"file": "91_screens.md", "heading": "Appendix C - Screenshots",
         "role": "derived", "derived_kind": "screenshot-index",
         "writer": "python", "order": 91},
    ]
    (area / "manifest.json").write_text(json.dumps(man, indent=1),
                                        encoding="utf-8")
    (area / "80_roles-dict.md").write_text(ROLE_DICT_MD, encoding="utf-8")
    (area / "91_screens.md").write_text(SCREENS_IDX_MD, encoding="utf-8")
    return area


def test_toc_depth_1_2_and_update_fields_on_open(tmp_path):
    """The TOC field collects H1-H2 only and settings.xml carries
    <w:updateFields/> so Word populates the field on open (no F9)."""
    area = make_area(tmp_path)
    out = tmp_path / "d.docx"
    render.render_folder(area, out, emit_signal=False)
    doc = Document(str(out))
    instr = "".join(t.text or ""
                    for t in doc.element.body.iter(qn("w:instrText")))
    assert 'TOC \\o "1-2"' in instr
    uf = doc.settings.element.find(qn("w:updateFields"))
    assert uf is not None
    assert uf.get(qn("w:val")) == "true"


def test_cover_title_lost_separator_normalized():
    """A multi-space run in the title is a lost separator -> ' - '."""
    from docx import Document as Doc
    doc = Doc()
    cfgi.build_cover(doc, "Fixed Assets  Desktop Procedures", "", [])
    assert "Fixed Assets - Desktop Procedures" in doc_text(doc)


def test_document_control_front_matter(tmp_path):
    """Folder renders open with a blank fill-by-hand Document Control table;
    kit (subset) renders never carry it."""
    area = make_area(tmp_path)
    out = tmp_path / "d.docx"
    render.render_folder(area, out, emit_signal=False)
    txt = doc_text(Document(str(out)))
    assert "Document Control" in txt
    assert "Summary of Changes" in txt
    kit = tmp_path / "kit.docx"
    render.render_folder(area, kit, slugs=["vendor-onboarding"],
                         emit_signal=False)
    assert "Summary of Changes" not in doc_text(Document(str(kit)))


def test_final_suppresses_screenshot_index_and_strips_leadin(tmp_path):
    """Final mode: the screenshot appendix never ships; register lead-in
    prose is stripped; a blank required cell lands on the scorecard."""
    area = make_area_with_registers(tmp_path)
    out = tmp_path / "final.docx"
    stats = render.render_folder(area, out, mode="final", emit_signal=False)
    txt = doc_text(Document(str(out)))
    assert "Appendix C - Screenshots" not in txt
    assert "Canonical functional roles" not in txt
    assert "Functional Role" in txt          # the table itself is kept
    assert any("Reports To" in x
               for x in stats["readiness"]["blank_cells"])


def test_working_keeps_screenshot_index_and_leadin(tmp_path):
    """Working mode is untouched by the final-mode register grooming."""
    area = make_area_with_registers(tmp_path)
    out = tmp_path / "draft.docx"
    render.render_folder(area, out, emit_signal=False)
    txt = doc_text(Document(str(out)))
    assert "Appendix C - Screenshots" in txt
    assert "Canonical functional roles" in txt


def test_final_skips_empty_register(tmp_path):
    """A register with a header and zero data rows is skipped whole in a
    final render (no heading over an empty shell), and reported."""
    area = make_area_with_registers(tmp_path)
    (area / "80_roles-dict.md").write_text(
        "## Role Dictionary\n\n"
        "<!-- derived: role-dictionary; writer: python -->\n\n"
        "_Canonical functional roles._\n\n"
        "| Functional Role | Reports To | Standard Responsibilities |\n"
        "|---|---|---|\n", encoding="utf-8")
    out = tmp_path / "final.docx"
    stats = render.render_folder(area, out, mode="final", emit_signal=False)
    assert "Functional Role" not in doc_text(Document(str(out)))
    assert stats["empty_registers_skipped"] == ["Role Dictionary"]


def test_final_lexicon_normalizes_spelling(tmp_path):
    """`_client/lexicon.yaml` terms normalize case-variants in final mode."""
    area = make_area(tmp_path)
    (area / "_client" / "lexicon.yaml").write_text(
        "lexicon:\n  - BlackLine\n", encoding="utf-8")
    p = area / "30_cash-application.md"
    p.write_text(p.read_text(encoding="utf-8")
                 + "\nSign-off is recorded in Blackline.\n", encoding="utf-8")
    out = tmp_path / "final.docx"
    stats = render.render_folder(area, out, mode="final", emit_signal=False)
    txt = doc_text(Document(str(out)))
    assert "BlackLine" in txt
    assert "Blackline" not in txt
    assert stats["lexicon_normalized"] == 1
    # Working mode never rewrites prose.
    draft = tmp_path / "draft.docx"
    stats_w = render.render_folder(area, draft, emit_signal=False)
    assert "Blackline" in doc_text(Document(str(draft)))
    assert stats_w["lexicon_normalized"] == 0


def test_final_readiness_scorecard_flags_defects(tmp_path):
    """Placeholders and doubled spaces (the dropped-verb signature) land on
    the readiness scorecard."""
    area = make_area(tmp_path)
    p = area / "30_cash-application.md"
    p.write_text(p.read_text(encoding="utf-8")
                 + "\nThe operator  posts the entry. Timing TBD.\n",
                 encoding="utf-8")
    out = tmp_path / "final.docx"
    stats = render.render_folder(area, out, mode="final", emit_signal=False)
    r = stats["readiness"]
    assert any("TBD" in x for x in r["placeholders"])
    assert any("cash-application" in x for x in r["double_spaces"])


def test_backmatter_sections_page_break_after_first(tmp_path):
    """Reference & Appendices sections each open a fresh page, except the
    first (which stays under the chapter heading)."""
    area = make_area_with_registers(tmp_path)
    out = tmp_path / "d.docx"
    render.render_folder(area, out, emit_signal=False)
    doc = Document(str(out))
    breaks = {}
    for p in doc.paragraphs:
        if p.style.name == "Heading 2" and p.text in (
                "Role Dictionary", "Appendix B - Gap Log",
                "Appendix C - Screenshots"):
            breaks[p.text] = bool(p.paragraph_format.page_break_before)
    assert breaks["Role Dictionary"] is False
    assert breaks["Appendix B - Gap Log"] is True
    assert breaks["Appendix C - Screenshots"] is True


def test_strict_final_exits_one_when_dirty(tmp_path, capsys):
    """--strict turns a dirty final scorecard into exit 1 (the docx is
    still written); working mode is unaffected by the flag."""
    area = make_area(tmp_path)
    out = tmp_path / "final.docx"
    rc = render.main([str(area), "-o", str(out), "--mode", "final",
                      "--strict"])
    assert rc == 1          # SC-01 screenshots are still pending capture
    assert out.is_file()
    assert "READINESS" in capsys.readouterr().out
    rc = render.main([str(area), "-o", str(tmp_path / "w.docx"), "--strict"])
    assert rc == 0


def test_body_prose_is_justified_tables_are_not(tmp_path):
    """Body flow text (plain paragraphs, list items) is justified; table
    cell text and headings keep their own alignment."""
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    area = make_area(tmp_path)
    out = tmp_path / "d.docx"
    render.render_folder(area, out, emit_signal=False)
    doc = Document(str(out))
    body = [p for p in doc.paragraphs
            if p.text.strip() and (p.style.name in ("Normal",)
                                   or p.style.name.startswith("List"))
            and not p.text.startswith("Table of Contents")]
    assert body, "expected body paragraphs"
    justified = [p for p in body
                 if p.alignment == WD_ALIGN_PARAGRAPH.JUSTIFY]
    assert justified, "body flow text must be justified"
    for p in doc.paragraphs:
        if p.style.name.startswith("Heading"):
            assert p.alignment != WD_ALIGN_PARAGRAPH.JUSTIFY
    for t in doc.tables:
        for row in t.rows:
            for c in row.cells:
                for p in c.paragraphs:
                    assert p.alignment != WD_ALIGN_PARAGRAPH.JUSTIFY
