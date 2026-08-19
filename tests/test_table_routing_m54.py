"""M54 — table-row comment routing: Part A tests (characterization + contract).

Two halves, matching the spec's shape (docs/v2/M54-table-row-routing.md):

1. `TestV1AppendixCharacterization` — NEW tests over OLD behavior, passing
   TODAY: exactly how a comment anchored inside a v1 appendix table (the
   Appendix A pain-points table of the frozen p2p engagement) routes through
   `review_extract` as it stands. Pinned byte-for-byte so Part B provably
   changes MATRIX routing and nothing else.

   *** SHARED-CODE FINDING (the spec's stop condition) ***
   The v1 appendix path and the matrix path are THE SAME CODE: both docx
   flavors are parsed by `review_extract.extract_document`, whose
   `iter_paragraphs` descends into every `w:tbl` and whose heading-stack
   `Location` + `resolve_slug` attribute every item, table-anchored or not.
   There is no separate v1 routing function. Part B's "first-cell" rule will
   therefore execute on v1 appendix tables too. The characterization below
   pins why the v1 OUTCOME still cannot change: an appendix row's first cell
   ("PP-33 (5.1) — …") does not resolve through `resolve_slug` (letters
   before the number defeat NUMBER_PREFIX_RE; the text is no procedure
   title), and a pain-points table INSIDE a procedure section already routes
   by the heading stack — pinned here as the behavior Part B must preserve.

2. `TestRowRouting` — the desired end state, skip-gated on the Part B
   resolver (`_row_first_cell`) landing in scripts/review_extract.py. These
   SKIP today; the M38 strict xfail
   (tests/test_matrix_roundtrip_m38.py::TestRoutingGap) stays in place this
   phase and retires in Part C.

Idioms are adapted from tests/test_matrix_roundtrip_m38.py and
tests/test_review_extract.py (self-contained here by convention — test
modules do not import from each other). Owns this file only; every fixture
is a tmp copy — the frozen fixtures are never written.
"""
from __future__ import annotations

import shutil
from pathlib import Path

import pytest

REPO = Path(__file__).resolve().parents[1]
SCRIPTS = REPO / "scripts"
P2P_AREA = REPO / "tests" / "fixtures" / "p2p-complete" / "components" / "procure-to-pay"
IPO_ROOT = REPO / "tests" / "fixtures" / "ipo-engagement"
MATRIX_YAML = REPO / "kernel" / "deliverables" / "process-controls-matrix.yaml"

if not P2P_AREA.is_dir():
    pytest.skip("frozen p2p fixture missing", allow_module_level=True)

from docx import Document          # noqa: E402
import yaml                        # noqa: E402

import render as render_mod        # noqa: E402
import review_extract              # noqa: E402

# --------------------------------------------------------------------------- #
# The Part B gate: TestRowRouting is the CONTRACT for a resolver that does not
# exist yet. It lands in review_extract.py under the name `_row_first_cell`
# (the M38 module docstring's exact mapping).
# --------------------------------------------------------------------------- #
ROW_RESOLVER_LANDED = "_row_first_cell" in (SCRIPTS / "review_extract.py").read_text(
    encoding="utf-8")

# v1 appendix scenario constants (frozen p2p engagement).
APX_HEADER = ["Pain Point", "Impact", "Severity", "Improvement Opportunity"]
APX_ROW_KEY = "PP-33"              # a Payments-group appendix row
APX_COMMENT = "Who owns remediation of this?"
IN_PROC_ROW_KEY = "PP-01"          # the same pain point rendered INSIDE its
IN_PROC_SLUG = "new-vendor-onboarding"   # owning procedure's body table
IN_PROC_COMMENT = "Severity seems overstated."

# Matrix scenario constants — mirrors tests/test_matrix_roundtrip_m38.py.
IPO_MATRIX_KIND = "process-controls-matrix"
OWNING_SLUG = "approve-exceptions"
OWNING_HEADING = "Approve Exceptions"
ROW_COMMENT = ("Disposition also produces a weekly exceptions log for the "
               "Controller — add it as an output of this step.")
CELL_COMMENT = "Should this output go to the CFO as well?"
ORPHAN_COMMENT = "Whose row is this anyway?"


# --------------------------------------------------------------------------- #
# Self-contained helpers (adapted from the m38 file / test_review_extract.py)
# --------------------------------------------------------------------------- #

def cell_paragraph(cell):
    for p in cell.paragraphs:
        if p.runs:
            return p
    raise AssertionError("cell has no run-bearing paragraph")


def comment_cell(doc, cell, text, author="SME Reviewer"):
    p = cell_paragraph(cell)
    doc.add_comment(p.runs, text=text, author=author, initials="SR")


def notes(area: Path, slug: str):
    p = area / "_review" / f"{slug}.notes.yaml"
    if not p.is_file():
        return None
    data = yaml.safe_load(p.read_text(encoding="utf-8"))
    assert data["procedure"] == slug
    return data["items"]


def extract(returned_docx: Path) -> None:
    assert review_extract.main([str(returned_docx), "--no-archive"]) == 0


# ---- v1 (p2p) substrate ---------------------------------------------------- #

@pytest.fixture(scope="module")
def p2p_rendered(tmp_path_factory) -> Path:
    """One working-mode render of the frozen p2p engagement (module-scoped:
    the render is the expensive step; each test copies it)."""
    base = tmp_path_factory.mktemp("m54-p2p") / "area"
    shutil.copytree(P2P_AREA, base)
    out = base / "draft.docx"
    render_mod.render_folder(base, out, mode="working", emit_signal=False)
    return base


@pytest.fixture()
def p2p_area(p2p_rendered, tmp_path) -> Path:
    dst = tmp_path / "area"
    shutil.copytree(p2p_rendered, dst)
    (dst / "_review" / "returned").mkdir(parents=True)
    return dst


def appendix_a_row(doc, key: str):
    """The [cells] of the Appendix A pain-points row whose first cell starts
    `key`. Appendix A tables are the 4-column tables headed exactly
    Pain Point | Impact | Severity | Improvement Opportunity — this shape
    distinguishes them from the pain-point tables rendered inside procedure
    bodies (which also contain the PP-ids; see in_procedure_row)."""
    for t in doc.tables:
        if len(t.columns) != 4 or not t.rows:
            continue
        if [c.text.strip() for c in t.rows[0].cells] != APX_HEADER:
            continue
        for row in t.rows[1:]:
            if row.cells[0].text.strip().startswith(key):
                return row.cells
    raise AssertionError(f"no Appendix A row starting {key!r}")


def in_procedure_row(doc, key: str):
    """A pain-point row rendered INSIDE a procedure section (NOT the 4-column
    appendix shape) — the v1 table whose routing already works via the
    heading stack."""
    for t in doc.tables:
        if t.rows and [c.text.strip() for c in t.rows[0].cells] == APX_HEADER:
            continue  # that's the appendix shape
        for row in t.rows:
            if row.cells and row.cells[0].text.strip().startswith(
                    f"PAIN POINT — {key}"):
                return row.cells
    raise AssertionError(f"no in-procedure pain-point row for {key!r}")


# ---- matrix (ipo) substrate ------------------------------------------------ #

HAVE_IPO = MATRIX_YAML.is_file() and IPO_ROOT.is_dir()

if HAVE_IPO:
    import definitions       # noqa: E402
    import render_glue       # noqa: E402


@pytest.fixture()
def ipo_area(tmp_path) -> Path:
    if not HAVE_IPO:
        pytest.skip("M38 matrix definition / IPO fixture not built yet")
    dest = tmp_path / "eng"
    shutil.copytree(IPO_ROOT, dest)
    area = dest / "components" / "purchasing"
    d = definitions.load_definition(IPO_MATRIX_KIND)
    plan = definitions.compile_plan(d, area)
    out = area / "_review" / "returned" / "matrix.docx"
    out.parent.mkdir(parents=True, exist_ok=True)
    render_glue.render_plan(plan, area, out)
    return area


def matrix_row(doc, step_heading: str):
    for table in doc.tables:
        for row in table.rows:
            cells = row.cells
            if cells and cells[0].text.strip().startswith(step_heading):
                return cells
    raise AssertionError(f"no matrix row whose step cell starts {step_heading!r}")


# =========================================================================== #
# 1. V1 APPENDIX CHARACTERIZATION — old behavior, pinned as-is, passing today
# =========================================================================== #

class TestV1AppendixCharacterization:
    """CURRENT v1 behavior, byte-identical before and after Part B.

    NOTE (shared code — see module docstring): these scenarios flow through
    the exact functions Part B will touch (`extract_document`,
    `iter_paragraphs`, the `Location` heading stack, `resolve_slug`). There
    is no v1-only routing path to leave alone; what these tests pin is the
    v1 OUTCOME, which must not move.
    """

    def test_appendix_table_comment_routes_to_unassigned(self, p2p_area):
        """A comment anchored in an Appendix A table row lands in
        review_extract.UNASSIGNED: the appendix H2 is a `derived` component
        with no slug, so the heading stack resolves nothing."""
        rv = p2p_area / "_review" / "returned" / "reviewed.docx"
        shutil.copy(p2p_area / "draft.docx", rv)
        doc = Document(str(rv))
        comment_cell(doc, appendix_a_row(doc, APX_ROW_KEY)[0], APX_COMMENT)
        doc.save(str(rv))
        extract(rv)

        (item,) = notes(p2p_area, review_extract.UNASSIGNED)
        assert item["type"] == "comment"
        assert item["note"] == APX_COMMENT
        # anchor: the row's first-cell VISIBLE text, clipped at 160 chars with
        # a trailing ellipsis (review_extract.clip). Pinned as-is.
        assert item["anchor"].startswith("PP-33 (5.1)")
        assert item["anchor"].endswith("…") and len(item["anchor"]) <= 160
        # QUIRK, pinned as-is: `location` degrades to the appendix's H4 group
        # heading ("Payments"), NOT an A–H subsection letter and NOT the step
        # named in the cell — the Location model treats the H4 as a "step".
        assert item["location"] == "Payments"
        # and nothing leaked onto the owning procedure
        assert notes(p2p_area, "weekly-payment-run") is None

    def test_appendix_first_cell_text_never_resolves_a_slug(self, p2p_area):
        """The invariance guarantee for Part B, proven on the real render:
        the appendix row's first-cell text ("PP-33 (5.1) — The Corporate
        Controller…") resolves to NO slug through review_extract.resolve_slug
        — the PP-id prefix defeats the number regex and the text is no
        procedure title. So even once the first-cell rule executes on this
        table (shared code), the item still lands in UNASSIGNED: the v1
        outcome cannot move."""
        doc = Document(str(p2p_area / "draft.docx"))
        cell_text = appendix_a_row(doc, APX_ROW_KEY)[0].text
        manifest = review_extract.load_manifest(p2p_area)
        num2slug, title2slug = review_extract.build_slug_maps(manifest)
        assert review_extract.resolve_slug(cell_text, num2slug, title2slug) is None
        # true for the whole Payments-group appendix rows, not just one
        for key in ("PP-33", "PP-36", "PP-39"):
            txt = appendix_a_row(doc, key)[0].text
            assert review_extract.resolve_slug(txt, num2slug, title2slug) is None

    def test_in_procedure_table_comment_routes_by_heading_stack(self, p2p_area):
        """A comment in a table INSIDE a procedure section already routes to
        that procedure via the heading stack — the working cell-level routing
        the spec says Part B must leave untouched. Pinned as-is, including
        the bare-letter location ("G", the enclosing A–H subsection)."""
        rv = p2p_area / "_review" / "returned" / "reviewed.docx"
        shutil.copy(p2p_area / "draft.docx", rv)
        doc = Document(str(rv))
        comment_cell(doc, in_procedure_row(doc, IN_PROC_ROW_KEY)[0],
                     IN_PROC_COMMENT)
        doc.save(str(rv))
        extract(rv)

        (item,) = notes(p2p_area, IN_PROC_SLUG)
        assert item["type"] == "comment"
        assert item["note"] == IN_PROC_COMMENT
        assert item["anchor"].startswith("PAIN POINT — PP-01")
        assert item["location"] == "G"
        assert notes(p2p_area, review_extract.UNASSIGNED) is None


# =========================================================================== #
# 2. ROW ROUTING — the desired end state, skip-gated on Part B landing.
#    These SKIP today; the M38 strict xfail keeps owning the "not yet" claim.
# =========================================================================== #

@pytest.mark.skipif(
    not ROW_RESOLVER_LANDED,
    reason="M54 Part B not landed: no `_row_first_cell` resolver in "
           "scripts/review_extract.py yet (the M38 strict xfail owns the gap "
           "until then)")
class TestRowRouting:
    """The Part B contract on the rendered matrix (mirrors the M38
    TestRoutingGap scenario, asserting the POSITIVE outcome)."""

    def test_row_comment_routes_to_owning_step(self, ipo_area):
        """A comment anchored to the matrix row's first (step-name) cell
        routes to the owning step's slug — the display number and title the
        matrix build already writes into that cell resolve through the
        existing resolver. Nothing falls to _unassigned."""
        rv = ipo_area / "_review" / "returned" / "matrix.docx"
        doc = Document(str(rv))
        comment_cell(doc, matrix_row(doc, OWNING_HEADING)[0], ROW_COMMENT)
        doc.save(str(rv))
        extract(rv)

        items = notes(ipo_area, OWNING_SLUG)
        assert items, f"expected _review/{OWNING_SLUG}.notes.yaml"
        assert any(it["note"] == ROW_COMMENT for it in items)
        assert notes(ipo_area, review_extract.UNASSIGNED) is None

    def test_non_first_cell_comment_routes_to_owning_step(self, ipo_area):
        """A comment anchored in a NON-first cell of the same row is still
        row-anchored text: it resolves through the row's first cell to the
        same owning slug — routing does not depend on which cell of the row
        the reviewer happened to touch."""
        rv = ipo_area / "_review" / "returned" / "matrix.docx"
        doc = Document(str(rv))
        cells = matrix_row(doc, OWNING_HEADING)
        others = [c for c in cells[1:] if c.text.strip()]
        assert others, "the owning row has no non-first cell with text"
        comment_cell(doc, others[0], CELL_COMMENT)
        doc.save(str(rv))
        extract(rv)

        items = notes(ipo_area, OWNING_SLUG)
        assert items and any(it["note"] == CELL_COMMENT for it in items)
        assert notes(ipo_area, review_extract.UNASSIGNED) is None

    def test_unresolvable_first_cell_still_lands_unassigned(self, ipo_area):
        """A first cell whose text resolves to no slug (mangled here to a
        string neither the number map nor the title map knows) still routes
        to review_extract.UNASSIGNED — the fix narrows the leak, it never
        guesses."""
        rv = ipo_area / "_review" / "returned" / "matrix.docx"
        doc = Document(str(rv))
        cells = matrix_row(doc, OWNING_HEADING)
        for p in cells[0].paragraphs:      # mangle the step-name cell
            for r in p.runs:
                r.text = ""
        cell_paragraph_runs = cells[0].paragraphs[0].runs
        assert cell_paragraph_runs, "mangled cell lost its runs"
        cell_paragraph_runs[0].text = "Zz No Such Step (99.99)"
        comment_cell(doc, cells[0], ORPHAN_COMMENT)
        doc.save(str(rv))
        extract(rv)

        items = notes(ipo_area, review_extract.UNASSIGNED)
        assert items and any(it["note"] == ORPHAN_COMMENT for it in items)
        assert notes(ipo_area, OWNING_SLUG) is None

    def test_v1_appendix_outcome_unmoved_by_the_resolver(self, p2p_area):
        """The characterization scenario re-run WITH the resolver present:
        the Appendix A row comment still lands in UNASSIGNED with the same
        anchor/location shape — Part B changed matrix routing and nothing
        else. (TestV1AppendixCharacterization pins the same bytes today;
        this test makes the before/after pair explicit once Part B lands.)"""
        rv = p2p_area / "_review" / "returned" / "reviewed.docx"
        shutil.copy(p2p_area / "draft.docx", rv)
        doc = Document(str(rv))
        comment_cell(doc, appendix_a_row(doc, APX_ROW_KEY)[0], APX_COMMENT)
        doc.save(str(rv))
        extract(rv)

        (item,) = notes(p2p_area, review_extract.UNASSIGNED)
        assert item["note"] == APX_COMMENT
        assert item["anchor"].startswith("PP-33 (5.1)")
        assert item["location"] == "Payments"
        assert notes(p2p_area, "weekly-payment-run") is None
