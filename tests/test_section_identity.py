"""Tests for M23 — section identity by slug (docs/M23-section-identity.md).

One test (or parametrized family) per Acceptance bullet, plus the registry's own
resolution rules and the migration script's guarantees.

Every fixture is a synthetic engagement under tmp_path built through the real
entry points (`scaffold.confirm`, `aggregate.run`, `render.render_folder`,
`reconcile.reconcile`, `orchestrate.decide`), because the claims being tested —
"letter-identical render", "zero fragment edits", "unmigrated areas keep
working" — are claims about artifacts, not about helper return values.

The area builder writes each procedure fragment in BOTH forms:
  - `letterless=True`  the post-migration contract (`### Process Overview`)
  - `letterless=False` the pre-migration corpus (`### A. Process Overview`)
so every behavioral test can assert the two are indistinguishable downstream.
"""
import hashlib
import json
from pathlib import Path

import pytest
import yaml
from docx import Document

import aggregate
import client_config
import doc_model
import migrate_sections
import orchestrate
import reconcile
import render
import scaffold


# --------------------------------------------------------------------------- #
# Fixtures
# --------------------------------------------------------------------------- #

TAXONOMY = {"taxonomy": {"categories": [
    {"slug": "finance", "subcategories": [{"slug": "cash"}, {"slug": "payments"}]},
]}}
PROPOSED_PROCS = {"procedures": [
    {"slug": "bank-rec", "title": "Bank Reconciliation", "l2": "cash"},
]}
SYSTEMS_YAML = "systems:\n  - slug: netsuite\n    name: NetSuite\n"
ROLES_YAML = "roles:\n  - slug: controller\n    name: Controller\n"
SOURCES_YAML = ("sources:\n- id: SRC-001\n  file: _sources/processed/i.md\n"
                "  touches:\n  - bank-rec\n  state: processed\n")

SECTION_BODIES = {
    "scope": "Covers the monthly reconciliation only. (SRC-001)",
    "quick-reference": ("| Field | Value |\n|---|---|\n"
                        "| Frequency | Monthly |\n"
                        "| Preparer | Controller |"),
    "before-you-start": ("- **Bank statement** — Treasury; downloaded for the "
                         "closed period."),
    "steps": ("#### Step 1: Export\n\nExport the statement from the portal."),
    "controls": ("> **CONTROL — CTRL-001:** Controller signs off.\n"
                 "> - **Type:** Detective\n"
                 "> - **Frequency:** Monthly\n"
                 "> - **Owner:** Controller"),
    "outputs": "- **Output 1:** Signed reconciliation.",
    "issues": ("> **PAIN POINT — PP-001:** The statement is exported by hand.\n"
               "> - **Impact:** Two hours per close\n"
               "> - **Severity:** Medium"),
}

#: An 8-heading fragment as the corpus carried it BEFORE M16 move 1 — the state
#: every already-drafted area is in on the day the registry lands. Titles resolve
#: through `SECTION_TITLE_ALIASES`, and `Pre-Requisites` + `Inputs` both resolve
#: to the merged `before-you-start`.
OLD_MODEL_FRAGMENT = """## Bank Reconciliation

### A. Process Overview

The Controller performs this monthly. (SRC-001)

### B. Quick Reference

- **Frequency:** Monthly
- **Preparer:** Controller

### C. Pre-Requisites

- The period is closed in NetSuite.

### D. Inputs

- **Input 1:** Bank statement — Treasury.

### E. Step-by-Step Procedure

#### Step 1: Export

Export the statement from the portal.

### F. Key Controls

> **CONTROL — CTRL-001:** Controller signs off.
> - **Type:** Detective
> - **Frequency:** Monthly
> - **Owner:** Controller

### G. Outputs

- **Output 1:** Signed reconciliation.

### H. Known Issues & Improvement Opportunities

> **PAIN POINT — PP-001:** The statement is exported by hand.
> - **Impact:** Two hours per close
> - **Severity:** Medium

```consult-meta
systems:
  - netsuite
roles:
  - controller
```
"""
META = "```consult-meta\nsystems:\n  - netsuite\nroles:\n  - controller\n```\n"


def fragment(title: str, *, letterless: bool = True, slugs=None) -> str:
    """One drafted procedure, headings in either form."""
    letters = doc_model.section_letters()
    out = [f"## {title}", ""]
    for slug in (slugs or doc_model.SECTION_SLUGS):
        head = doc_model.SECTION_TITLES[slug]
        if not letterless:
            head = f"{letters[slug]}. {head}"
        out += [f"### {head}", "", SECTION_BODIES[slug], ""]
    return "\n".join(out) + "\n" + META


def engagement(tmp_path) -> Path:
    components = tmp_path / "components"
    components.mkdir(parents=True, exist_ok=True)
    return components


def write_profile(components: Path, **fields) -> Path:
    client = components / "_client"
    client.mkdir(parents=True, exist_ok=True)
    path = client / "profile.yaml"
    path.write_text(yaml.safe_dump({"profile": fields}, sort_keys=False),
                    encoding="utf-8")
    return path


def area_with(components: Path, *, letterless=True, name="treasury",
              slugs=None) -> Path:
    """A reconcile-clean one-procedure area, drafted in the requested form."""
    area = components / name
    area.mkdir(parents=True, exist_ok=True)
    prof = client_config.profile(area)
    comps = [{"file": "10_bank-rec.md", "role": "procedure", "slug": "bank-rec",
              "heading": "Bank Reconciliation", "l2": "cash", "order": 10}]
    for d in scaffold.profile_derived_files(prof):
        comps.append({"file": d["file"], "role": "derived",
                      "derived_kind": d["kind"], "writer": d["writer"],
                      "heading": d["heading"], "order": d["order"]})
    (area / "manifest.json").write_text(json.dumps(
        {"schema": "consult-mvp-manifest/v1", "area": name, "l1": "finance",
         "l2_order": ["cash", "payments"], "title": "Treasury",
         "subtitle": "Current state", "components": comps}, indent=1),
        encoding="utf-8")
    (area / "10_bank-rec.md").write_text(
        fragment("Bank Reconciliation", letterless=letterless, slugs=slugs),
        encoding="utf-8")
    ref = area / "_reference"
    ref.mkdir(exist_ok=True)
    (ref / "systems.yaml").write_text(SYSTEMS_YAML, encoding="utf-8")
    (ref / "roles.yaml").write_text(ROLES_YAML, encoding="utf-8")
    (ref / "sources.yaml").write_text(SOURCES_YAML, encoding="utf-8")
    for comp in comps:
        path = area / comp["file"]
        if comp["role"] == "derived" and comp["writer"] == "agent":
            path.write_text(
                f"## {comp['heading']}\n\n<!-- derived: {comp['derived_kind']}; "
                f"writer: agent -->\n\nNothing of note for [[bank-rec]].\n",
                encoding="utf-8")
    assert aggregate.run(str(area)) == 0
    return area


def rendered_text(area: Path, out_name="draft.docx", **kw) -> str:
    out = area / out_name
    render.render_folder(area, out, emit_signal=False, **kw)
    doc = Document(str(out))
    parts = [p.text for p in doc.paragraphs]
    for t in doc.tables:
        for row in t.rows:
            for c in row.cells:
                parts.append(c.text)
    return "\n".join(parts)


def sha(path: Path) -> str:
    return hashlib.sha1(path.read_bytes()).hexdigest()


# --------------------------------------------------------------------------- #
# Design point 1 — the registry
# --------------------------------------------------------------------------- #

def test_the_registry_is_the_m16_seven_section_set():
    """M16 move 1's registry half: seven slugs, in the ticket's document order,
    with the declared-job titles."""
    assert list(doc_model.SECTION_TITLES) == [
        "scope", "quick-reference", "before-you-start", "steps", "outputs",
        "controls", "issues"]
    assert doc_model.section_letters() == {
        "scope": "A", "quick-reference": "B", "before-you-start": "C",
        "steps": "D", "outputs": "E", "controls": "F", "issues": "G"}


def test_the_letter_aliases_stay_frozen_at_their_historical_meaning():
    """The letters are NOT the new positions: they are frozen at the sections
    they historically named, so a profile written before M16 move 1 keeps
    meaning the same sections (M23 design point 3). `G` still means `outputs`
    even though `outputs` now renders as E, and `C`/`D` both point at the merged
    `before-you-start`."""
    assert doc_model.SECTION_LETTER_ALIASES == {
        "A": "scope", "B": "quick-reference", "C": "before-you-start",
        "D": "before-you-start", "E": "steps", "F": "controls",
        "G": "outputs", "H": "issues"}
    assert doc_model.section_letters()["outputs"] == "E"   # display ≠ alias


def test_every_pre_m16_heading_still_resolves():
    """The rename half of M16 move 1 costs ZERO fragment edits: every 8-section
    title in the corpus resolves through `SECTION_TITLE_ALIASES`, and the two
    merged ones land on the same slug."""
    for title, slug in [("Process Overview", "scope"),
                        ("Quick Reference", "quick-reference"),
                        ("Pre-Requisites", "before-you-start"),
                        ("Inputs", "before-you-start"),
                        ("Step-by-Step Procedure", "steps"),
                        ("Outputs", "outputs"),
                        ("Key Controls", "controls"),
                        ("Known Issues & Improvement Opportunities", "issues")]:
        assert doc_model.section_of_heading(f"### {title}") == slug
        assert doc_model.section_slug(title) == slug
    # the pre-M16 SLUGS keep resolving as profile input, too
    for old, slug in [("overview", "scope"), ("prerequisites", "before-you-start"),
                      ("inputs", "before-you-start")]:
        assert doc_model.section_slug(old) == slug


@pytest.mark.parametrize("token,slug", [
    ("controls", "controls"),            # the slug
    ("Key Controls", "controls"),        # the canonical title
    ("key-controls", "controls"),        # a normalized spelling
    ("F", "controls"), ("f", "controls"), ("F.", "controls"),   # letter aliases
    ("H", "issues"), ("A", "scope"), ("D", "before-you-start"),
    ("overview", "scope"),               # a pre-M16 slug (profile input)
    ("At a Glance", "quick-reference"),  # the new title
    ("steps", "steps"), ("Steps", "steps"),   # the SLUG, case-insensitive
    ("Z", None), ("", None), (None, None), ("Appendix A", None),
])
def test_section_slug_canonicalizes_every_accepted_spelling(token, slug):
    """`section_slug` is the one canonicalizer for PROFILE INPUT: slug, title or
    letter alias in, slug out. (Headings are stricter — see
    `section_of_heading`: `### Steps` is nobody's TITLE, so it is not a
    section heading.)"""
    assert doc_model.section_slug(token) == slug


@pytest.mark.parametrize("line,slug", [
    ("### Scope", "scope"),                         # migrated
    ("### Process Overview", "scope"),              # the pre-M16 title
    ("### A. Process Overview", "scope"),           # pre-M16 AND unmigrated
    ("### Inputs", "before-you-start"),             # merged into C
    ("### Pre-Requisites", "before-you-start"),     # …and so is this
    ("### F. Local Wording For Controls", "controls"),   # letter fallback
    ("### Known Issues & Improvement Opportunities", "issues"),
    ("#### Step 1: Export", None),                  # a step is not a section
    ("### Pain Points", None),                      # a derived view's heading
    ("Process Overview", None),
])
def test_section_of_heading_is_the_one_heading_parser(line, slug):
    """THE TRANSITION CONTRACT: both heading forms parse, title first and letter
    second, so migration is never a prerequisite."""
    assert doc_model.section_of_heading(line) == slug


# --------------------------------------------------------------------------- #
# Design point 2 + Acceptance bullet 1 — letter-identical rendering
# --------------------------------------------------------------------------- #

def test_letterless_fragments_render_letter_identical_to_lettered_ones(tmp_path):
    """Acceptance bullet 1: A–H assigned from the default profile order, so a
    migrated area and an unmigrated one produce the same document."""
    migrated = rendered_text(area_with(engagement(tmp_path / "new")))
    legacy = rendered_text(area_with(engagement(tmp_path / "old"),
                                     letterless=False))
    assert migrated == legacy
    for slug, letter in doc_model.section_letters().items():
        assert f"{letter}. {doc_model.SECTION_TITLES[slug]}" in migrated


def test_the_fragment_on_disk_never_carries_a_letter(tmp_path):
    """Design point 2: the letter exists only in the rendered document."""
    area = area_with(engagement(tmp_path))
    text = (area / "10_bank-rec.md").read_text(encoding="utf-8")
    assert "### Scope" in text
    for letter in doc_model.SECTION_LETTER_ALIASES:
        assert f"### {letter}." not in text
    assert "A. Scope" in rendered_text(area)


def test_scaffold_stamps_letterless_skeletons_from_the_registry(tmp_path):
    """Design point 5: scaffold emits letterless headings, every registry title
    present, none of them lettered."""
    components = engagement(tmp_path)
    area = components / "treasury"
    proposed = area / "_reference" / ".proposed"
    proposed.mkdir(parents=True, exist_ok=True)
    (proposed / "procedures.yaml").write_text(yaml.safe_dump(PROPOSED_PROCS),
                                              encoding="utf-8")
    taxonomy = tmp_path / "taxonomy.yaml"
    taxonomy.write_text(yaml.safe_dump(TAXONOMY), encoding="utf-8")
    scaffold.confirm(area, "finance", taxonomy, None, None)

    skeleton = (area / "10_bank-rec.md").read_text(encoding="utf-8")
    for title in doc_model.SECTION_TITLES.values():
        assert f"### {title}" in skeleton
    for letter in doc_model.SECTION_LETTER_ALIASES:
        assert f"### {letter}." not in skeleton


# --------------------------------------------------------------------------- #
# Acceptance bullet 2 — reordering re-letters, zero fragment edits
# --------------------------------------------------------------------------- #

def test_reordering_sections_reletters_the_render_with_zero_fragment_edits(tmp_path):
    """Acceptance bullet 2. The fragment's bytes are asserted UNCHANGED across
    the re-letter — that is the whole point of the ticket."""
    components = engagement(tmp_path)
    area = area_with(components)
    frag = area / "10_bank-rec.md"
    before = sha(frag)
    assert "A. Scope" in rendered_text(area, "one.docx")

    reordered = ["quick-reference", "scope", "before-you-start", "steps",
                 "outputs", "controls", "issues"]
    write_profile(components, sections=reordered)
    text = rendered_text(area, "two.docx")

    assert "A. At a Glance" in text and "B. Scope" in text
    assert sha(frag) == before


def test_a_lettered_fragment_is_reletterd_too_not_left_stale(tmp_path):
    """An UNMIGRATED fragment renders with the letter the PROFILE says, never
    the stale one baked into its heading."""
    components = engagement(tmp_path)
    write_profile(components, sections=["quick-reference", "scope",
                                        "before-you-start", "steps",
                                        "outputs", "controls", "issues"])
    text = rendered_text(area_with(components, letterless=False))
    assert "A. At a Glance" in text and "B. Scope" in text
    assert "A. Scope" not in text


# --------------------------------------------------------------------------- #
# Acceptance bullet 3 — `body_omit: [controls]` == `body_omit: [F]`
# --------------------------------------------------------------------------- #

@pytest.mark.parametrize("token", ["controls", "F", "Key Controls"])
def test_body_omit_means_the_same_section_however_it_is_spelled(tmp_path, token):
    """Acceptance bullet 3: the letter alias is canonicalized to the slug, so
    the three spellings are one profile."""
    components = engagement(tmp_path / token)
    write_profile(components, body_omit=[token],
                  derived=client_config.DEFAULT_DERIVED + ["appendix-controls"])
    prof = client_config.profile(components / "treasury")
    assert prof.body_omit == ["controls"]
    text = rendered_text(area_with(components))
    assert "F. Key Controls" not in text
    assert "CTRL-" in text                 # the register still carries it


def test_sections_accepts_letters_and_slugs_interchangeably(tmp_path):
    """Design point 3: an existing letter-based profile and its slug-based twin
    resolve to the same Profile."""
    a = engagement(tmp_path / "letters")
    write_profile(a, sections=[s for s in "ABCDE"] + ["G", "H"],
                  derived=client_config.DEFAULT_DERIVED + ["appendix-controls"])
    b = engagement(tmp_path / "slugs")
    write_profile(b, sections=["scope", "quick-reference", "before-you-start",
                               "steps", "outputs", "issues"],  # letters A B C D E G H
                  derived=client_config.DEFAULT_DERIVED + ["appendix-controls"])
    pa = client_config.profile(a / "treasury")
    pb = client_config.profile(b / "treasury")
    assert pa.sections == pb.sections
    assert pa.letters() == pb.letters()
    # Dropping `controls` closes the gap: `issues` becomes F, not G.
    assert pa.letters()["issues"] == "F"


# --------------------------------------------------------------------------- #
# Design point 4 — home-section rules re-key to slugs; no letter literals
# --------------------------------------------------------------------------- #

def test_home_section_rules_are_slugs(tmp_path):
    """CONTROL→`controls`, PP/IO→`issues`, inline kinds→`steps`."""
    import callouts
    assert callouts.home_section("CONTROL") == "controls"
    assert callouts.home_section("PAIN POINT") == "issues"
    assert callouts.home_section("IMPROVEMENT OPPORTUNITY") == "issues"
    assert callouts.home_section("VALIDATION REQUIRED") == "steps"
    assert callouts.home_section("SCREENSHOT PLACEHOLDER") == "steps"
    assert callouts.home_section("VIBE CHECK") is None


def test_aggregate_keys_section_bodies_by_slug_in_both_forms():
    """The engines read sections by slug; the heading form is irrelevant."""
    for letterless in (True, False):
        subs = aggregate.split_subsections(
            fragment("Bank Reconciliation", letterless=letterless))
        assert set(subs) == set(doc_model.SECTION_SLUGS)
        assert "Frequency" in subs["quick-reference"]
        assert "CTRL-001" in subs["controls"]


def test_no_letter_literal_survives_in_scripts_home_section_logic():
    """Acceptance bullet 5, mechanized: nothing in scripts/ may key a section on
    a letter. Grepping for the two shapes that encode one — a `### [A-H].`
    heading pattern and a bare A–H section-key literal — must find nothing."""
    import re
    banned = [
        re.compile(r'###\\s\*\\s\+\(\?P<letter>'),          # letter-capturing regex
        re.compile(r'"\^###\\\\s\+\[A-Z'),
        re.compile(r'###\s+\[A-H\]'),
        re.compile(r'###\\s\+\(\[A-Z\]\)'),
        re.compile(r'sections\.get\(\s*"[A-H]"'),
        re.compile(r'subs\.get\(\s*"[A-H]"'),
        re.compile(r'(CONTROLS_SECTION|ISSUES_SECTION)\s*=\s*"[A-H]"'),
        re.compile(r'ALL_SECTIONS\s*=\s*\[\s*"A"'),
        re.compile(r'MANDATORY_SECTIONS\s*=\s*\[\s*"A"'),
    ]
    root = Path(__file__).resolve().parents[1] / "scripts"
    offenders = []
    for path in sorted(root.glob("*.py")):
        text = path.read_text(encoding="utf-8")
        for pat in banned:
            for m in pat.finditer(text):
                offenders.append(f"{path.name}: {m.group(0)!r}")
    assert offenders == []


def test_the_advisors_profile_drift_guard_keys_on_slugs(tmp_path):
    """Design point 5: the `reprofile` detector keeps working, and the migration
    itself can never read as drift (both heading forms count as present)."""
    components = engagement(tmp_path)
    write_profile(components, sections=["scope", "quick-reference", "steps"])
    area = area_with(components)
    orchestrate.emit_aggregate(str(area), warnings=[])
    for letterless in (True, False):
        (area / "10_bank-rec.md").write_text(
            fragment("Bank Reconciliation", letterless=letterless),
            encoding="utf-8")
        assert orchestrate.decide(str(area))["action"] != "reprofile"

    # A section the profile requires and the fragment lacks IS drift, by slug.
    write_profile(components, sections=["scope", "quick-reference", "steps"])
    (area / "10_bank-rec.md").write_text(
        fragment("Bank Reconciliation", slugs=["scope", "quick-reference"])
        .replace("<!-- unfilled -->", ""), encoding="utf-8")
    d = orchestrate.decide(str(area))
    assert d["action"] == "reprofile"
    assert d["details"]["missing"] == {"bank-rec": ["steps"]}


# --------------------------------------------------------------------------- #
# M10 provenance — the ordering that is load-bearing
# --------------------------------------------------------------------------- #

def test_the_letter_stamp_runs_before_the_body_is_emitted(tmp_path):
    """M10's provenance maps anchor on RENDERED text, which keeps its letters,
    so the stamp has to happen inside the per-procedure transform chain — before
    any body line is emitted and therefore before the .maps sidecar is derived
    from those lines. This is that ordering's guard."""
    area = area_with(engagement(tmp_path))
    raw = (area / "10_bank-rec.md").read_text(encoding="utf-8")
    prof = client_config.profile(area)
    assert "### A. Scope" in render._apply_profile(raw, prof)
    assert "A. Scope" in rendered_text(area)


def test_the_review_map_still_anchors_on_the_right_fragment_lines(tmp_path):
    """The other half of the ordering claim: the stamp is line-count-preserving,
    so every map entry's recorded fragment line range still hashes to the
    paragraph the reviewer will comment on."""
    area = area_with(engagement(tmp_path))
    stats = render.render_folder(area, area / "draft.docx", emit_signal=False)
    payload = json.loads(Path(stats["map"]).read_text(encoding="utf-8"))
    lines = (area / "10_bank-rec.md").read_text(encoding="utf-8").split("\n")

    checked = 0
    for entry in payload["entries"].values():
        if entry["file"] != "10_bank-rec.md" or entry["kind"] != "para":
            continue
        lo, hi = entry["lines"]
        source = "\n".join(lines[lo:hi + 1]).strip()
        assert entry["hash"] == hashlib.sha1(
            source.encode("utf-8")).hexdigest(), (
                f"map entry points at {source!r}, which is not what it hashed")
        checked += 1
    assert checked, "no paragraph entries to verify"


def test_lettering_preserves_the_line_count_provenance_depends_on():
    """Every M14/M16 body transform is line-count-preserving; so is this one."""
    text = fragment("Bank Reconciliation")
    letters = doc_model.section_letters()
    stamped = render._letter_sections(text, letters)
    assert len(stamped.split("\n")) == len(text.split("\n"))
    assert stamped.count("### ") == text.count("### ")


# --------------------------------------------------------------------------- #
# Acceptance bullet 4 — the migration script
# --------------------------------------------------------------------------- #

def test_migration_strips_letters_and_changes_nothing_else(tmp_path):
    """Acceptance bullet 4: converts an area, reconcile clean, render
    letter-identical."""
    area = area_with(engagement(tmp_path), letterless=False)
    orchestrate.emit_aggregate(str(area), warnings=[])
    before_render = rendered_text(area, "before.docx")
    before = (area / "10_bank-rec.md").read_text(encoding="utf-8")

    assert migrate_sections.main([str(area), "--write"]) == 0

    after = (area / "10_bank-rec.md").read_text(encoding="utf-8")
    assert after == "\n".join(
        f"### {ln.split('. ', 1)[1]}" if ln.startswith("### ") else ln
        for ln in before.split("\n"))
    assert len(after.split("\n")) == len(before.split("\n"))
    assert aggregate.run(str(area)) == 0
    orchestrate.emit_aggregate(str(area), warnings=[])
    assert reconcile.reconcile(area) == 0
    assert rendered_text(area, "after.docx") == before_render


def test_migration_is_a_dry_run_by_default(tmp_path):
    """Nothing is written without `--write`, and the report says what would be."""
    area = area_with(engagement(tmp_path), letterless=False)
    frag = area / "10_bank-rec.md"
    before = sha(frag)
    assert migrate_sections.main([str(area)]) == 0
    assert sha(frag) == before


def test_migration_is_idempotent(tmp_path):
    """A second run finds nothing and writes nothing."""
    area = area_with(engagement(tmp_path), letterless=False)
    frag = area / "10_bank-rec.md"
    assert migrate_sections.main([str(area), "--write"]) == 0
    once = sha(frag)
    stats = migrate_sections.migrate_area(area, write=True)
    assert stats["headings"] == 0 and stats["written"] == []
    assert sha(frag) == once
    assert migrate_sections.strip_letters(
        frag.read_text(encoding="utf-8"))[1] == 0


def test_migration_leaves_an_unknown_title_alone(tmp_path):
    """A heading whose title is not in the registry keeps its letter: the letter
    is the only thing naming its section, so dropping it would lose identity."""
    text = "## P\n\n### F. Locally Reworded Controls\n\nbody\n"
    migrated, changed = migrate_sections.strip_letters(text)
    assert changed == 0 and migrated == text


def test_migration_touches_only_procedure_fragments(tmp_path):
    """Derived views have `###` sub-headings of their own ("### Pain Points");
    the migration must never open them."""
    area = area_with(engagement(tmp_path), letterless=False)
    appendix = area / "88_appendix-a.md"
    before = sha(appendix)
    migrate_sections.migrate_area(area, write=True)
    assert sha(appendix) == before


# --------------------------------------------------------------------------- #
# Acceptance bullet 6 — M16 move 1's re-letter is a registry + profile edit
# --------------------------------------------------------------------------- #

def test_a_registry_rename_and_merge_relettes_without_touching_fragments(
        tmp_path, monkeypatch):
    """Acceptance bullet 6, and the reason this ticket exists.

    A future reshape (M16 move 1) is performed here as what it should be: an
    edit to `SECTION_TITLES` (rename), an entry in `SECTION_TITLE_ALIASES` (so
    already-drafted fragments keep resolving), and a profile whose `sections:`
    drops a merged section. No fragment is opened — asserted by hash — and the
    render re-letters accordingly. The CONTENT mergers stay a drafter wave, out
    of scope here.
    """
    components = engagement(tmp_path)
    area = area_with(components)
    frag = area / "10_bank-rec.md"
    before = sha(frag)

    renamed = dict(doc_model.SECTION_TITLES)
    renamed["controls"] = "Controls & Checks"          # a further rename
    monkeypatch.setattr(doc_model, "SECTION_TITLES", renamed)
    monkeypatch.setattr(doc_model, "SECTION_TITLE_ALIASES",
                        dict(doc_model.SECTION_TITLE_ALIASES,
                             **{"Key Controls": "controls"}))
    write_profile(components, sections=["scope", "quick-reference",
                                        "before-you-start", "steps",
                                        "controls", "issues"])

    text = rendered_text(area, "reshaped.docx")
    assert "D. Procedure" in text                       # was E pre-M16
    assert "E. Controls & Checks" in text               # renamed AND re-lettered
    assert "F. Outputs & Evidence" not in text          # dropped by the profile
    assert sha(frag) == before                          # zero fragment edits


def test_the_letter_of_a_body_omitted_section_does_not_shift(tmp_path):
    """`body_omit` hides a section from the body without removing it from
    `sections:`, so it holds its letter position and its neighbours keep
    theirs — the M14 behavior, preserved."""
    components = engagement(tmp_path)
    write_profile(components, body_omit=["controls"],
                  derived=client_config.DEFAULT_DERIVED + ["appendix-controls"])
    text = rendered_text(area_with(components))
    assert "D. Procedure" in text
    assert "E. Outputs & Evidence" in text and "F. Outputs" not in text
    assert "F. Key Controls" not in text


# --------------------------------------------------------------------------- #
# Transition contract — an UNMIGRATED area is fully functional
# --------------------------------------------------------------------------- #

def test_an_unmigrated_area_renders_aggregates_and_reconciles_clean(tmp_path):
    """The documented transition contract: parsing accepts BOTH forms, so the
    migration is a tidy-up and never a prerequisite."""
    area = area_with(engagement(tmp_path), letterless=False)
    text = rendered_text(area)
    assert "A. Scope" in text and "G. Known Issues" in text
    assert "PP-01" in text          # the display id
    orchestrate.emit_aggregate(str(area), warnings=[])
    assert reconcile.reconcile(area) == 0


def test_a_half_migrated_area_is_still_coherent(tmp_path):
    """Mid-migration (one fragment converted, one not) is a real state: both
    fragments must render with the same lettering."""
    area = area_with(engagement(tmp_path), letterless=False)
    mixed = (area / "10_bank-rec.md").read_text(encoding="utf-8")
    mixed = mixed.replace("### A. Scope", "### Scope")
    (area / "10_bank-rec.md").write_text(mixed, encoding="utf-8")
    text = rendered_text(area)
    for slug, letter in doc_model.section_letters().items():
        assert f"{letter}. {doc_model.SECTION_TITLES[slug]}" in text


# --------------------------------------------------------------------------- #
# M16 move 1 — the TRANSITION for an OLD-MODEL (8-heading) fragment
# --------------------------------------------------------------------------- #

def test_an_old_model_fragment_parses_renders_and_reconciles(tmp_path):
    """The registry half's whole promise: a fragment drafted against the 8-section
    model — untouched, letters and pre-M16 titles and all — still parses,
    renders with the SEVEN-section lettering and the NEW titles, aggregates and
    reconciles (exit 0). No drafter has run yet."""
    area = area_with(engagement(tmp_path))
    frag = area / "10_bank-rec.md"
    frag.write_text(OLD_MODEL_FRAGMENT, encoding="utf-8")

    text = rendered_text(area)
    assert "A. Scope" in text                      # was `A. Process Overview`
    assert "B. At a Glance" in text                # was `B. Quick Reference`
    assert "C. Before You Start" in text           # was `C. Pre-Requisites`
    assert "D. Procedure" in text                  # was `E. Step-by-Step …`
    assert "E. Outputs & Evidence" in text         # was `G. Outputs`
    assert "F. Key Controls" in text
    assert "G. Known Issues" in text               # was H
    assert "H. " not in text                       # seven sections, not eight

    orchestrate.emit_aggregate(str(area), warnings=[])
    assert reconcile.reconcile(area) == 0           # WARNING, never an error


def test_two_headings_resolving_to_one_slug_tolerate_report_and_keep_facts(
        tmp_path):
    """The C+D merge, pre-content-wave. THE TRANSITION BEHAVIOUR:

    - aggregate CONCATENATES the two bodies, so no fact is lost;
    - render letters both headings C and re-titles only the FIRST, so the
      un-merged pair is visible instead of printed twice;
    - reconcile WARNS, naming the fragment and the content wave (exit stays 0).
    """
    area = area_with(engagement(tmp_path))
    frag = area / "10_bank-rec.md"
    frag.write_text(OLD_MODEL_FRAGMENT, encoding="utf-8")

    dupes = doc_model.duplicate_sections(OLD_MODEL_FRAGMENT)
    assert dupes == {"before-you-start": ["Pre-Requisites", "Inputs"]}

    body = aggregate.split_subsections(OLD_MODEL_FRAGMENT)["before-you-start"]
    assert "The period is closed in NetSuite." in body       # from C
    assert "Bank statement — Treasury." in body              # from D — kept

    text = rendered_text(area)
    assert "C. Before You Start" in text
    assert "C. Inputs" in text                # the repeat keeps its own title
    assert text.count("C. Before You Start") == 1

    warns = reconcile.parse_procedure("bank-rec", "10_bank-rec.md",
                                      OLD_MODEL_FRAGMENT).warnings
    assert any("AWAITING THE M16 CONTENT WAVE" in w for w in warns)
    assert any("Before You Start" in w and "Inputs" in w for w in warns)
    orchestrate.emit_aggregate(str(area), warnings=[])
    assert reconcile.reconcile(area) == 0          # loud, not blocking
