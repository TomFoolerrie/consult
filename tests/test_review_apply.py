"""Tests for scripts/review_apply.py — the M10 deterministic tracked-changes
apply (verified-or-fallback: precision over recall).

Fixture strategy: build a tiny two-procedure area under tmp, render it with
scripts/render.py in working mode (real bookmarks + sidecar review map), then
simulate reviewer tracked changes by editing the docx XML the way Word writes
them (w:ins runs for insertions; deleted text wrapped as w:delText inside
w:del), drop the doc into _review/returned/ and run the apply.
"""
from __future__ import annotations

import copy
import json
import shutil
from pathlib import Path

import pytest
import yaml
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

import render as render_mod
import review_apply
import reconcile as reconcile_mod


# --------------------------------------------------------------------------- #
# area fixture (rendered ONCE per module, copied per test)
# --------------------------------------------------------------------------- #

BANK_REC = """## Bank Reconciliation

### A. Purpose

The clerk reconciles the monthly bank statement against the general ledger.

### B. Steps

- Download the bank statement from the portal.
- Compare balances and investigate any differences.

> **CONTROL — CTRL-01:** The controller signs off on the completed reconciliation.

See [[journal-entries]] for posting adjustments.
"""

JOURNAL = """## Journal Entries

### A. Purpose

Accountants post manual journal entries at month end.

> **PAIN POINT — PP-01:** Approvals are tracked by email. (Severity: Medium)
"""

MANIFEST = {
    "schema": "consult-mvp-manifest/v1",
    "area": "test-area",
    "l1": "record-to-report",
    "title": "Test Area Process Documentation",
    "subtitle": "Working Draft",
    "l2_order": ["close"],
    "components": [
        {"file": "bank-rec.md", "heading": "Bank Reconciliation",
         "role": "procedure", "slug": "bank-rec", "l2": "close", "order": 1},
        {"file": "journal-entries.md", "heading": "Journal Entries",
         "role": "procedure", "slug": "journal-entries", "l2": "close", "order": 2},
    ],
}


@pytest.fixture(scope="module")
def rendered_base(tmp_path_factory) -> Path:
    """One area, rendered in working mode (bookmarks + sidecar map stamped)."""
    area = tmp_path_factory.mktemp("apply-base") / "area"
    area.mkdir()
    (area / "manifest.json").write_text(json.dumps(MANIFEST), encoding="utf-8")
    (area / "bank-rec.md").write_text(BANK_REC, encoding="utf-8")
    (area / "journal-entries.md").write_text(JOURNAL, encoding="utf-8")
    render_mod.render_folder(area, area / "draft.docx", emit_signal=False)
    return area


@pytest.fixture()
def area(rendered_base, tmp_path) -> Path:
    """Per-test private copy of the rendered area (map travels inside it)."""
    dst = tmp_path / "area"
    shutil.copytree(rendered_base, dst)
    (dst / "_review" / "returned").mkdir(parents=True, exist_ok=True)
    return dst


@pytest.fixture()
def returned(area) -> Path:
    """The rendered doc copied into _review/returned/ as a reviewer return."""
    rv = area / "_review" / "returned" / "reviewed.docx"
    shutil.copy(area / "draft.docx", rv)
    return rv


# --------------------------------------------------------------------------- #
# Word-style tracked-change surgery helpers
# --------------------------------------------------------------------------- #

def para_text(p_el) -> str:
    return "".join(t.text or "" for t in p_el.iter(qn("w:t")))


def find_para(doc, needle):
    for p_el in doc.element.body.iter(qn("w:p")):
        if needle in para_text(p_el):
            return p_el
    raise AssertionError(f"no paragraph containing {needle!r}")


_wid = [100]


def track_replace(doc, needle, old, new, author="Reviewer One"):
    """Replace `old` with `new` in the paragraph containing `needle`, written
    the way Word records it: original text wrapped as w:delText inside w:del,
    inserted text as a run inside w:ins. Empty new = pure deletion."""
    p_el = find_para(doc, needle)
    for r in p_el.iter(qn("w:r")):
        t = r.find(qn("w:t"))
        if t is None or old not in (t.text or ""):
            continue
        before, _, after = (t.text or "").partition(old)
        parent = r.getparent()
        idx = list(parent).index(r)

        def mk_run(txt):
            nr = copy.deepcopy(r)
            nt = nr.find(qn("w:t"))
            nt.text = txt
            nt.set(qn("xml:space"), "preserve")
            return nr

        def wrap(tag, txt, deltext=False):
            _wid[0] += 1
            el = OxmlElement(tag)
            el.set(qn("w:id"), str(_wid[0]))
            el.set(qn("w:author"), author)
            el.set(qn("w:date"), "2026-07-21T09:00:00Z")
            run = mk_run(txt)
            if deltext:
                run.find(qn("w:t")).tag = qn("w:delText")
            el.append(run)
            return el

        nodes = []
        if before:
            nodes.append(mk_run(before))
        if old:
            nodes.append(wrap("w:del", old, deltext=True))
        if new:
            nodes.append(wrap("w:ins", new))
        if after:
            nodes.append(mk_run(after))
        for n in nodes:
            parent.insert(idx, n)
            idx += 1
        parent.remove(r)
        return p_el
    raise AssertionError(f"no run containing {old!r} near {needle!r}")


def untracked_replace(doc, needle, old, new):
    """Silently mutate run text WITHOUT tracked-change markup (a reviewer who
    turned tracking off) — invalidates the paragraph's recorded hash."""
    p_el = find_para(doc, needle)
    for t in p_el.iter(qn("w:t")):
        if old in (t.text or ""):
            t.text = (t.text or "").replace(old, new)
            return p_el
    raise AssertionError(f"no run text containing {old!r}")


def apply_run(area, rv):
    report: dict = {}
    review_apply.apply_docx(area, rv, report)
    return report


def notes_files(area):
    return sorted((area / "_review").glob("*.notes.yaml"))


def load_notes(area, slug):
    data = yaml.safe_load(
        (area / "_review" / f"{slug}.notes.yaml").read_text(encoding="utf-8"))
    return data["items"]


# --------------------------------------------------------------------------- #
# clean applies
# --------------------------------------------------------------------------- #

def test_clean_tracked_edit_updates_fragment(area, returned):
    """A clean tracked replacement on an anchored body paragraph is written
    back into the markdown fragment, with nothing routed to notes."""
    d = Document(str(returned))
    track_replace(d, "monthly bank statement", "monthly", "weekly")
    d.save(str(returned))

    report = apply_run(area, returned)
    frag = (area / "bank-rec.md").read_text(encoding="utf-8")
    assert "The clerk reconciles the weekly bank statement" in frag
    assert "monthly" not in frag
    assert report["files"][0]["applied"] == 1
    assert report["files"][0]["noted"] == 0
    assert notes_files(area) == []


def test_edit_preserves_tokens_and_formatting_outside_the_edit(area, returned):
    """Splicing an edit into a paragraph carrying a [[slug]] cross-reference
    keeps the token (and all untouched lines) byte-for-byte intact."""
    before = (area / "bank-rec.md").read_text(encoding="utf-8")
    d = Document(str(returned))
    track_replace(d, "for posting adjustments", "posting", "recording")
    d.save(str(returned))

    report = apply_run(area, returned)
    frag = (area / "bank-rec.md").read_text(encoding="utf-8")
    assert report["files"][0]["applied"] == 1
    assert "See [[journal-entries]] for recording adjustments." in frag
    # every other line survives verbatim
    changed = [(a, b) for a, b in zip(before.split("\n"), frag.split("\n"))
               if a != b]
    assert changed == [("See [[journal-entries]] for posting adjustments.",
                        "See [[journal-entries]] for recording adjustments.")]


def test_applied_fragment_stays_reconcile_clean(area, returned):
    """An applied edit never introduces reconcile grammar errors."""
    d = Document(str(returned))
    track_replace(d, "Compare balances", "investigate", "chase down")
    d.save(str(returned))
    report = apply_run(area, returned)
    assert report["files"][0]["applied"] == 1
    frag = reconcile_mod.parse_procedure(
        "bank-rec", "bank-rec.md",
        (area / "bank-rec.md").read_text(encoding="utf-8"))
    assert frag.errors == []


def test_pure_deletion_removes_the_paragraph_line(area, returned):
    """A tracked deletion of a whole paragraph deletes its markdown line."""
    d = Document(str(returned))
    track_replace(d, "for posting adjustments",
                  "See 1.2 Journal Entries for posting adjustments.", "")
    d.save(str(returned))
    report = apply_run(area, returned)
    assert report["files"][0]["applied"] == 1
    frag = (area / "bank-rec.md").read_text(encoding="utf-8")
    assert "[[journal-entries]]" not in frag


def test_callout_body_edit_applies_keeping_label_line(area, returned):
    """An edit to a callout BODY (after the label/ID prefix) is appliable —
    callout boxes are anchored 1x1 tables — and the label survives."""
    d = Document(str(returned))
    track_replace(d, "controller signs off", "signs off", "approves and signs off")
    d.save(str(returned))
    report = apply_run(area, returned)
    assert report["files"][0]["applied"] == 1
    frag = (area / "bank-rec.md").read_text(encoding="utf-8")
    assert ("> **CONTROL — CTRL-01:** The controller approves and signs off "
            "on the completed reconciliation." in frag)


def test_apply_never_archives_the_returned_doc(area, returned):
    """review_apply leaves the returned docx in place (review_extract
    --comments-only archives it afterwards) and never creates processed/."""
    d = Document(str(returned))
    track_replace(d, "monthly bank statement", "monthly", "weekly")
    d.save(str(returned))
    apply_run(area, returned)
    assert returned.exists()
    assert not (area / "_review" / "processed").exists()


# --------------------------------------------------------------------------- #
# verified-or-fallback: nothing unverified is ever written
# --------------------------------------------------------------------------- #

def test_untracked_change_in_same_paragraph_routes_to_notes(area, returned):
    """A tracked edit sharing a paragraph with an UNtracked change (hash
    mismatch after reject-all) writes nothing and lands in notes.yaml with
    the anchor context."""
    before = (area / "bank-rec.md").read_text(encoding="utf-8")
    d = Document(str(returned))
    track_replace(d, "monthly bank statement", "clerk", "analyst")
    untracked_replace(d, "against the general ledger", "ledger", "register")
    d.save(str(returned))

    report = apply_run(area, returned)
    assert (area / "bank-rec.md").read_text(encoding="utf-8") == before
    assert report["files"][0]["applied"] == 0
    assert report["files"][0]["noted"] == 1
    (item,) = load_notes(area, "bank-rec")
    assert item["type"] == "tracked-change"
    assert item["location"] == "bank-rec.md (L2 close)"
    assert "hash mismatch" in item["change"]
    assert item["author"] == "Reviewer One"
    assert item["source"] == "reviewed.docx"


def test_callout_label_id_edit_routes_to_notes(area, returned):
    """An edit touching the callout label/ID line prefix is never applied —
    ID minting belongs to a drafter — and falls back to notes."""
    before = (area / "bank-rec.md").read_text(encoding="utf-8")
    d = Document(str(returned))
    track_replace(d, "CONTROL — CTRL-01", "CTRL-01", "CTRL-99")
    d.save(str(returned))

    report = apply_run(area, returned)
    assert (area / "bank-rec.md").read_text(encoding="utf-8") == before
    assert report["files"][0]["applied"] == 0
    assert report["files"][0]["noted"] == 1
    (item,) = load_notes(area, "bank-rec")
    assert "callout label/ID line changed" in item["change"]


def test_corrupted_bookmark_routes_edit_to_notes(area, returned):
    """An edit in a paragraph whose bookmark no longer matches the sidecar map
    is unanchored: nothing is written and the item goes to _unassigned notes."""
    before = (area / "bank-rec.md").read_text(encoding="utf-8")
    d = Document(str(returned))
    p_el = track_replace(d, "monthly bank statement", "monthly", "weekly")
    for bm in p_el.iter(qn("w:bookmarkStart")):
        bm.set(qn("w:name"), "cw_99999")   # corrupt the anchor
    d.save(str(returned))

    report = apply_run(area, returned)
    assert (area / "bank-rec.md").read_text(encoding="utf-8") == before
    assert report["files"][0]["applied"] == 0
    assert report["files"][0]["noted"] == 1
    (item,) = load_notes(area, review_apply.UNASSIGNED)
    assert "no verified anchor" in item["change"]


def test_stripped_bookmark_routes_edit_to_notes(area, returned):
    """An edited paragraph with NO cw_ bookmark at all is unanchored → notes."""
    d = Document(str(returned))
    p_el = track_replace(d, "Download the bank statement", "portal", "bank portal")
    for tag in ("w:bookmarkStart", "w:bookmarkEnd"):
        for bm in list(p_el.iter(qn(tag))):
            bm.getparent().remove(bm)
    d.save(str(returned))

    report = apply_run(area, returned)
    assert report["files"][0]["applied"] == 0
    (item,) = load_notes(area, review_apply.UNASSIGNED)
    assert "no verified anchor" in item["change"]


def test_fragment_changed_since_render_routes_to_notes(area, returned):
    """If the source fragment drifted after the render (re-synthesis no longer
    matches the reviewed original), the edit falls back instead of splicing."""
    frag_path = area / "bank-rec.md"
    drifted = frag_path.read_text(encoding="utf-8").replace(
        "The clerk reconciles the monthly", "The clerk now reconciles the monthly")
    frag_path.write_text(drifted, encoding="utf-8")

    d = Document(str(returned))
    track_replace(d, "monthly bank statement", "general ledger", "sub-ledger")
    d.save(str(returned))
    report = apply_run(area, returned)
    assert frag_path.read_text(encoding="utf-8") == drifted
    assert report["files"][0]["applied"] == 0
    (item,) = load_notes(area, "bank-rec")
    # NOTE: the recorded hash also no longer matches, so the earlier anchor
    # check fires first — either fallback reason honours the guarantee.
    assert "mismatch" in item["change"]


def test_docx_without_review_map_is_skipped(area, tmp_path):
    """A returned doc with no cw-map category / sidecar map is left for
    review_extract — reported as skipped, nothing applied or noted."""
    plain = tmp_path / "plain.docx"
    d = Document()
    d.add_paragraph("Not one of ours.")
    d.save(str(plain))
    report = apply_run(area, plain)
    assert "files" not in report
    assert any("no review map" in s for s in report["skipped"])
    assert notes_files(area) == []


# --------------------------------------------------------------------------- #
# CLI
# --------------------------------------------------------------------------- #

def test_cli_sweeps_returned_and_applies(area, returned, capsys):
    """The CLI with only an area sweeps _review/returned/*.docx, applies clean
    edits, exits 0, and leaves the doc in place for review_extract."""
    d = Document(str(returned))
    track_replace(d, "monthly bank statement", "monthly", "weekly")
    d.save(str(returned))

    rc = review_apply.main([str(area)])
    out = capsys.readouterr().out
    assert rc == 0
    assert "applied 1, noted 0" in out
    assert "weekly bank statement" in (area / "bank-rec.md").read_text(encoding="utf-8")
    assert returned.exists()


def test_cli_with_nothing_returned(area, capsys):
    """No returned docs is a clean no-op exit."""
    rc = review_apply.main([str(area)])
    assert rc == 0
    assert "no reviewed documents" in capsys.readouterr().out
