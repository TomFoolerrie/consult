"""M38 WP-M3 — the review round-trip for a TABLE-FIRST deliverable.

What this file proves, over a tmp copy of the frozen IPO engagement:

  render the matrix docx  ->  reviewer marks up a ROW (a comment anchored to
  the row's step-name cell, plus a Word tracked change inside a cell of the
  same row)  ->  review_extract  ->  the notes bus  ->  (routing)  ->  the
  owning step's fragment is edited  ->  aggregate rebuilds the matrix VIEW
  ->  re-render.

Two halves, reported honestly and separately:

1. THE VIEW-REBUILD HALF — a real pass. The matrix is a view: the reviewer's
   substance is applied to the OWNING STEP's fragment, `aggregate` rebuilds
   `20_process-controls-matrix.md` from it, and the re-rendered docx carries
   the change. No reviewer edit is ever written into the matrix fragment
   (asserted).

2. THE ROUTING HALF — THE DOCUMENTED GAP the ticket anticipated. See
   `TestRoutingGap`. review_extract attributes an item to a procedure by the
   H2 heading stack (`resolve_slug` over the area manifest's *procedure*
   components). The matrix renders as ONE H2 section ("Process & Controls
   Matrix", a `role: derived` component with no slug) whose H3s are TAXONOMY
   NODES, not steps — so every row-anchored item resolves to no slug and is
   routed to `_review/_unassigned.notes.yaml`. Nothing is dropped, nothing
   lands on the matrix; but the OWNING STEP is not yet the note target for a
   table-first document.

   THE EXACT MAPPING THE MACHINERY WOULD NEED (and no more): when the
   anchoring `w:p` is inside a `w:tbl` row, resolve the slug from the ROW's
   first cell instead of the heading stack — the rendered "Process Step" cell
   already carries both keys `resolve_slug` accepts (the display number, e.g.
   "1.3", and the step title, e.g. "Approve Exceptions"), so the resolution
   needs no new authored data. `test_row_first_cell_carries_both_resolution_
   keys` proves that on the real render; the desired end state is pinned as a
   strict xfail so it flips to a pass the day the mapping lands.

Owns this file only. The fixtures are frozen (fingerprinted below).
"""
from __future__ import annotations

import copy
import shutil
from pathlib import Path

import pytest

REPO = Path(__file__).resolve().parents[1]
MATRIX_YAML = REPO / "kernel" / "deliverables" / "process-controls-matrix.yaml"
IPO_ROOT = Path(__file__).resolve().parent / "fixtures" / "ipo-engagement"

if not MATRIX_YAML.is_file() or not IPO_ROOT.is_dir():
    pytest.skip("M38 matrix definition / IPO fixture not built yet — the target",
                allow_module_level=True)

from docx import Document              # noqa: E402
from docx.oxml import OxmlElement      # noqa: E402
from docx.oxml.ns import qn            # noqa: E402
import yaml                            # noqa: E402

import aggregate                        # noqa: E402
import definitions                      # noqa: E402
import doc_model                        # noqa: E402
import render_glue                      # noqa: E402
import review_extract                   # noqa: E402

IPO_AREA = IPO_ROOT / "components" / "purchasing"
MATRIX_FRAGMENT = "20_process-controls-matrix.md"
MATRIX_KIND = "process-controls-matrix"

# The row under review and the step that owns it.
OWNING_SLUG = "approve-exceptions"
OWNING_HEADING = "Approve Exceptions"
REVIEW_COMMENT = ("Disposition also produces a weekly exceptions log for the "
                  "Controller — add it as an output of this step.")
# The substance of the comment, applied by hand to the OWNING step (standing in
# for the consult-drafter update run the orchestrator would dispatch).
OWNING_OUTPUT_LINE = "- Returned invoice — sent back to the supplier by email"
NEW_OUTPUT_LINE = ("- Weekly exceptions log — to the Controller (Microsoft "
                   "Excel, filed on the finance shared drive)")
NEW_OUTPUT_MARKER = "Weekly exceptions log"


# --------------------------------------------------------------------------- #
# Substrate: a tmp copy of the frozen engagement with the matrix rendered
# --------------------------------------------------------------------------- #

def ipo_copy(tmp_path):
    dest = tmp_path / "eng"
    shutil.copytree(IPO_ROOT, dest)
    return dest, dest / "components" / "purchasing"


def render_matrix(area: Path, out: Path) -> Path:
    """The M38 gate's invocation: load definition -> compile plan -> render."""
    d = definitions.load_definition(MATRIX_KIND)
    plan = definitions.compile_plan(d, area)
    out.parent.mkdir(parents=True, exist_ok=True)
    render_glue.render_plan(plan, area, out)
    return out


@pytest.fixture()
def area(tmp_path) -> Path:
    _root, a = ipo_copy(tmp_path)
    render_matrix(a, a / "_review" / "returned" / "matrix.docx")
    return a


# --------------------------------------------------------------------------- #
# Reviewer-markup synthesis — the SAME approach tests/test_review_extract.py
# uses (python-docx `add_comment`; w:del/w:delText + w:ins written by hand),
# retargeted from body paragraphs to TABLE CELL paragraphs. The v1 helpers
# themselves reach body paragraphs only (`para_by_text` walks
# `doc.paragraphs`, which excludes table cells), so the lookup below is the
# one thing this file re-does; the markup XML is v1's.
# --------------------------------------------------------------------------- #

def matrix_row(doc, step_heading: str):
    """The [cells] of the matrix row whose FIRST cell names `step_heading`."""
    for table in doc.tables:
        for row in table.rows:
            cells = row.cells
            if cells and cells[0].text.strip().startswith(step_heading):
                return cells
    raise AssertionError(f"no matrix row whose step cell starts {step_heading!r}")


def cell_paragraph(cell):
    for p in cell.paragraphs:
        if p.runs:
            return p
    raise AssertionError("cell has no run-bearing paragraph")


def add_comment(doc, cell, text, author="SME Reviewer"):
    p = cell_paragraph(cell)
    doc.add_comment(p.runs, text=text, author=author, initials="SR")
    return p


_wid = [900]


def track_replace_in(cell, old: str, new: str, author="Reviewer One"):
    """v1's Word-style tracked replacement (w:del + w:ins), scoped to one cell."""
    p_el = cell_paragraph(cell)._p
    for r in p_el.iter(qn("w:r")):
        t = r.find(qn("w:t"))
        if t is None or old not in (t.text or ""):
            continue
        before, _, after = (t.text or "").partition(old)
        parent = r.getparent()
        idx = list(parent).index(r)

        def mk_run(txt):
            nr = copy.deepcopy(r)
            nt = nr.find(qn("w:t"))
            nt.text = txt
            nt.set(qn("xml:space"), "preserve")
            return nr

        def wrap(tag, txt, deltext=False):
            _wid[0] += 1
            el = OxmlElement(tag)
            el.set(qn("w:id"), str(_wid[0]))
            el.set(qn("w:author"), author)
            el.set(qn("w:date"), "2026-08-14T09:00:00Z")
            run = mk_run(txt)
            if deltext:
                run.find(qn("w:t")).tag = qn("w:delText")
            el.append(run)
            return el

        nodes = [mk_run(before)] if before else []
        nodes.append(wrap("w:del", old, deltext=True))
        if new:
            nodes.append(wrap("w:ins", new))
        if after:
            nodes.append(mk_run(after))
        for n in nodes:
            parent.insert(idx, n)
            idx += 1
        parent.remove(r)
        return
    raise AssertionError(f"no run containing {old!r} in this cell")


def mark_up_matrix(area: Path) -> Path:
    """Comment the owning step's row + a tracked change inside the same row."""
    rv = area / "_review" / "returned" / "matrix.docx"
    doc = Document(str(rv))
    cells = matrix_row(doc, OWNING_HEADING)
    add_comment(doc, cells[0], REVIEW_COMMENT)          # anchored to the row
    outputs = [c for c in cells if "Returned invoice" in c.text]
    assert outputs, "the owning row has no Outputs cell to edit"
    track_replace_in(outputs[0], "by email", "by secure email")
    doc.save(str(rv))
    return rv


def notes(area: Path, slug: str):
    p = area / "_review" / f"{slug}.notes.yaml"
    if not p.is_file():
        return None
    data = yaml.safe_load(p.read_text(encoding="utf-8"))
    assert data["procedure"] == slug
    return data["items"]


def all_notes_files(area: Path):
    return sorted(p.name for p in (area / "_review").glob("*.notes.yaml"))


# --------------------------------------------------------------------------- #
# 1. What the v1 synthesis machinery CAN reach on a table-first docx
# --------------------------------------------------------------------------- #

class TestMarkupReachesTheTable:
    def test_comment_and_tracked_change_land_inside_a_matrix_row(self, area):
        """Both v1 markup mechanisms work INSIDE the matrix table: a comment
        anchored to the row's step-name cell and a del→ins in the same row's
        Outputs cell. Only the v1 *lookup* helper needed replacing (it walks
        body paragraphs, which excludes cells) — the markup XML is unchanged,
        and review_extract already descends into w:tbl (`iter_paragraphs`)."""
        rv = mark_up_matrix(area)
        xml = Document(str(rv)).element.body.xml
        assert "commentRangeStart" in xml
        assert "<w:ins " in xml and "<w:del " in xml and "delText" in xml

    def test_extraction_sees_both_items(self, area):
        mark_up_matrix(area)
        assert review_extract.main([str(area / "_review" / "returned"),
                                    "--no-archive"]) == 0
        items = [it for f in all_notes_files(area)
                 for it in yaml.safe_load(
                     (area / "_review" / f).read_text(encoding="utf-8"))["items"]]
        types = sorted(it["type"] for it in items)
        assert types == ["comment", "tracked-change"], items
        assert all(it["kind"] == "review" for it in items)   # M6 bus stamp
        (c,) = [it for it in items if it["type"] == "comment"]
        assert c["note"] == REVIEW_COMMENT
        (tc,) = [it for it in items if it["type"] == "tracked-change"]
        assert tc["change"] == "by email → by secure email"


# --------------------------------------------------------------------------- #
# 2. The matrix is a VIEW — nothing ever targets it (a real pass)
# --------------------------------------------------------------------------- #

class TestNoNoteTargetsTheMatrix:
    def test_no_note_file_or_item_names_the_matrix(self, area):
        """No reviewer item is routed to the matrix fragment or its kind: the
        matrix is a view, so it is never a drafter target."""
        mark_up_matrix(area)
        assert review_extract.main([str(area / "_review" / "returned"),
                                    "--no-archive"]) == 0
        forbidden = {MATRIX_FRAGMENT, MATRIX_KIND,
                     Path(MATRIX_FRAGMENT).stem,          # 20_process-controls-matrix
                     "process-controls-matrix.notes.yaml"}
        for name in all_notes_files(area):
            assert name.replace(".notes.yaml", "") not in forbidden, name
            data = yaml.safe_load(
                (area / "_review" / name).read_text(encoding="utf-8"))
            assert data["procedure"] not in forbidden

    def test_reviewer_markup_never_edits_the_matrix_fragment(self, area):
        """Extraction is read-only over the source of truth: the derived
        matrix fragment is byte-identical after the round trip."""
        frag = area / MATRIX_FRAGMENT
        before = frag.read_bytes()
        mark_up_matrix(area)
        assert review_extract.main([str(area / "_review" / "returned"),
                                    "--no-archive"]) == 0
        assert frag.read_bytes() == before


# --------------------------------------------------------------------------- #
# 3. THE DOCUMENTED GAP — routing a row-anchored item to the owning step
# --------------------------------------------------------------------------- #

class TestRoutingGap:
    """The gap WP-M3 anticipated, CLOSED by M54 (Amendment A0): when the H2
    heading stack resolves no slug and the anchoring w:p sits inside a w:tbl
    row, review_extract now resolves the slug from the row's FIRST CELL
    (`_row_first_cell`) through the existing `resolve_slug` — exactly the
    mapping `test_row_first_cell_carries_both_resolution_keys` proved
    available. The former strict xfail below asserts positively; the former
    current-behavior test pins the NEW routed outcome."""

    def test_current_behavior_row_items_route_to_unassigned(self, area):
        """(M54: was the pinned OLD behavior — row items in _unassigned.)
        Row-anchored items now route to the OWNING step's slug via the
        row's first cell; the anchor still carries the first-cell text and
        nothing falls to _unassigned."""
        mark_up_matrix(area)
        assert review_extract.main([str(area / "_review" / "returned"),
                                    "--no-archive"]) == 0
        assert notes(area, review_extract.UNASSIGNED) is None
        items = notes(area, OWNING_SLUG)
        assert items and len(items) == 2
        # the anchor still names the owning row — the first-cell text — so
        # nothing is lost in the re-dispatch.
        assert any(OWNING_HEADING in it["anchor"] for it in items)
        # location still degrades to the taxonomy-node H3, not the step
        assert all(it["location"] for it in items)

    def test_row_first_cell_carries_both_resolution_keys(self, area):
        """THE MAPPING THAT WOULD CLOSE THE GAP, proven available: the rendered
        "Process Step" cell of the owning row contains the step's display
        number AND its title — the two keys `review_extract.resolve_slug`
        already accepts. So the fix is 'when the anchoring w:p is inside a
        w:tbl row, resolve the slug from the row's first cell', with no new
        authored data and no renderer change."""
        doc = Document(str(area / "_review" / "returned" / "matrix.docx"))
        cell_text = matrix_row(doc, OWNING_HEADING)[0].text
        manifest = doc_model.load_manifest(area)
        num2slug, title2slug = review_extract.build_slug_maps(manifest)
        assert OWNING_HEADING in cell_text
        assert doc_model.display_numbers(manifest)[OWNING_SLUG] in cell_text
        # both keys resolve through the existing resolver
        assert review_extract.resolve_slug(
            OWNING_HEADING, num2slug, title2slug) == OWNING_SLUG
        assert review_extract.resolve_slug(
            cell_text.split("(")[-1].rstrip(") ") + " " + OWNING_HEADING,
            num2slug, title2slug) == OWNING_SLUG

    def test_desired_row_comment_targets_the_owning_step(self, area):
        """Formerly the suite's one strict xfail — flipped positive by M54
        Part B (`_row_first_cell` in scripts/review_extract.py)."""
        mark_up_matrix(area)
        assert review_extract.main([str(area / "_review" / "returned"),
                                    "--no-archive"]) == 0
        items = notes(area, OWNING_SLUG)
        assert items, f"expected _review/{OWNING_SLUG}.notes.yaml"
        assert any(it["note"] == REVIEW_COMMENT for it in items)


# --------------------------------------------------------------------------- #
# 4. The view-rebuild half — a real pass, independent of the routing outcome
# --------------------------------------------------------------------------- #

class TestViewRebuild:
    def test_step_edit_rebuilds_the_matrix_and_the_docx(self, area):
        """The reviewer's substance is applied to the OWNING STEP (standing in
        for the consult-drafter update run), `aggregate` rebuilds the matrix
        view plan-driven, and the re-rendered docx carries the change — the
        view-rebuild loop for a table-first deliverable."""
        mark_up_matrix(area)
        assert review_extract.main([str(area / "_review" / "returned"),
                                    "--no-archive"]) == 0

        step = area / f"10_{OWNING_SLUG}.md"
        text = step.read_text(encoding="utf-8")
        assert OWNING_OUTPUT_LINE in text
        matrix = area / MATRIX_FRAGMENT
        before_frag = matrix.read_text(encoding="utf-8")
        assert NEW_OUTPUT_MARKER not in before_frag

        step.write_text(
            text.replace(OWNING_OUTPUT_LINE,
                         OWNING_OUTPUT_LINE + "\n" + NEW_OUTPUT_LINE),
            encoding="utf-8")

        assert aggregate.run(str(area)) == 0
        after_frag = matrix.read_text(encoding="utf-8")
        assert after_frag != before_frag
        assert NEW_OUTPUT_MARKER in after_frag
        # the change landed on the OWNING step's row, nowhere else
        owning_rows = [ln for ln in after_frag.splitlines()
                       if NEW_OUTPUT_MARKER in ln]
        assert len(owning_rows) == 1
        assert OWNING_SLUG in owning_rows[0]

        out = render_matrix(area, area / "matrix-v2.docx")
        doc = Document(str(out))
        assert any(NEW_OUTPUT_MARKER in c.text
                   for c in matrix_row(doc, OWNING_HEADING))


# --------------------------------------------------------------------------- #
# 5. The frozen fixture is untouched by any of the above
# --------------------------------------------------------------------------- #

def test_frozen_ipo_fixture_untouched(tmp_path):
    fingerprints = {p: (p.stat().st_mtime_ns, p.stat().st_size)
                    for p in IPO_ROOT.rglob("*") if p.is_file()}
    _root, a = ipo_copy(tmp_path)
    render_matrix(a, a / "_review" / "returned" / "matrix.docx")
    mark_up_matrix(a)
    assert review_extract.main([str(a / "_review" / "returned"),
                                "--no-archive"]) == 0
    assert aggregate.run(str(a)) == 0
    for p, fp in fingerprints.items():
        assert (p.stat().st_mtime_ns, p.stat().st_size) == fp, p
