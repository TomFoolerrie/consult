"""Tests for scripts/kits.py — per-procedure (L3) review kit emission (M9).

Reuses the synthetic area from test_render (same fixture shape: roles.yaml
people lists + org-chart ranks; three procedures with resolvable, resolvable-
to-another-person, and unresolvable preparers).
"""
import json
from pathlib import Path

import pytest
from docx import Document
from docx.oxml.ns import qn

import kits
from xlsx_min import read_xlsx
from test_render import make_area


@pytest.fixture(scope="module")
def built(tmp_path_factory):
    """Build kits once for the read-only assertions below."""
    area = make_area(tmp_path_factory.mktemp("kits"))
    kits.build_kits(str(area))
    return area, area / "_review" / "kits"


def test_one_folder_per_procedure(built):
    """Kits are grouped one folder per PROCEDURE (L3) — review travels by
    topic; contacts are annotations inside the kit, not the grouping key."""
    _, out = built
    dirs = sorted(p.name for p in out.iterdir() if p.is_dir())
    assert dirs == ["1-1_vendor-onboarding", "2-1_payment-run",
                    "2-2_cash-application"]


def test_contact_resolution_biases_lowest_ranked_person(built):
    """The 'AP Clerk' role resolves to Jane Junior (deepest reports_to chain),
    not the first-listed Mark Manager — org-chart rank wins over list order."""
    _, out = built
    idx = (out / "index.md").read_text(encoding="utf-8")
    assert "| Jane Junior |" in idx and "Mark Manager" not in idx.split("Escalation")[0].split("## By person")[0]
    assert (out / "1-1_vendor-onboarding" / "1.1_vendor-onboarding.docx").is_file()


def test_gap_workbook_rows_carry_ref_column(built):
    """Gap workbook rows carry a 'Ref (do not edit)' column of slug#LOCAL-ID
    (local id, not display id) so ingest maps rows without guessing."""
    _, out = built
    rows = read_xlsx(out / "1-1_vendor-onboarding" / "gaps_vendor-onboarding.xlsx")
    assert [r["Ref (do not edit)"] for r in rows] == ["vendor-onboarding#GAP-01"]
    # display ids in the visible Gap ID column stay global
    assert [r["Gap ID"] for r in rows] == ["GAP-01"]
    assert all(r["Answer"] == "" and r["Status"] == "" for r in rows)


def test_gap_stays_in_its_procedure_kit_naming_its_own_contact(built):
    """A gap whose 'Owner to confirm' resolves to Jane STAYS in the
    payment-run kit (Carol's procedure) — the Contact column names Jane, the
    kit set stays whole, and the index flags the extra contact."""
    _, out = built
    rows = read_xlsx(out / "2-1_payment-run" / "gaps_payment-run.xlsx")
    assert [r["Ref (do not edit)"] for r in rows] == ["payment-run#GAP-01"]
    assert rows[0]["Contact"] == "Jane Junior"
    # escalation is the resolved contact's manager from the org chart
    assert rows[0]["Escalation"] == "Mark Manager"
    idx = (out / "index.md").read_text(encoding="utf-8")
    assert "Carol Chief (+ Jane Junior)" in idx


def test_readme_and_index_written(built):
    """Every kit gets a README.md; index.md lists each kit with its person."""
    _, out = built
    for k in ("1-1_vendor-onboarding", "2-1_payment-run", "2-2_cash-application"):
        assert (out / k / "README.md").is_file()
    idx = (out / "index.md").read_text(encoding="utf-8")
    assert "| `1-1_vendor-onboarding/` | 1.1 Vendor Onboarding | Jane Junior |" in idx
    assert "## By person" in idx
    assert "- **Jane Junior**: `1-1_vendor-onboarding/`" in idx
    readme = (out / "1-1_vendor-onboarding" / "README.md").read_text(encoding="utf-8")
    assert "Review kit — 1.1 Vendor Onboarding" in readme
    assert "**Send to: Jane Junior**" in readme
    assert "1.1_vendor-onboarding.docx" in readme
    # the cross-contact kit says who else is involved
    readme2 = (out / "2-1_payment-run" / "README.md").read_text(encoding="utf-8")
    assert "**Send to: Carol Chief** · also involves: Jane Junior" in readme2


def test_unresolvable_procedure_gets_an_unassigned_contact(built):
    """The procedure whose preparer matches no role still gets ITS kit — the
    contact is flagged unassigned for human triage in the index."""
    _, out = built
    assert (out / "2-2_cash-application" / "2.2_cash-application.docx").is_file()
    idx = (out / "index.md").read_text(encoding="utf-8")
    assert "| `2-2_cash-application/` | 2.2 Cash Application | (unassigned) |" in idx


def test_kit_docs_have_tracked_changes_and_global_ids(built):
    """Kit procedure docs open with tracked changes enforced and use the same
    display numbers / global callout ids as the full draft."""
    _, out = built
    doc = Document(str(out / "2-1_payment-run" / "2.1_payment-run.docx"))
    settings = doc.settings.element
    assert settings.find(qn("w:trackChanges")) is not None
    dp = settings.find(qn("w:documentProtection"))
    assert dp is not None and dp.get(qn("w:edit")) == "trackedChanges"
    text = "\n".join(p.text for p in doc.paragraphs)
    assert "2.1 Payment Run" in text
    # callouts render inside 1x1 table cells, so look there for the id
    cell_text = "\n".join(c.text for t in doc.tables
                          for row in t.rows for c in row.cells)
    assert "GAP-02" in cell_text     # global display id, not local GAP-01
    assert "GAP-01" not in cell_text


def test_screenshot_template_boxes_and_sidecar_map(built):
    """The screenshot doc has one bookmarked (scr_NNN) paste box per SC item,
    and a consult-screens-map sidecar under _review/.maps keyed by the doc id
    stamped into core properties."""
    area, out = built
    sdoc_path = out / "1-1_vendor-onboarding" / "screenshots_vendor-onboarding.docx"
    assert sdoc_path.is_file()
    doc = Document(str(sdoc_path))
    names = [bm.get(qn("w:name"))
             for bm in doc.element.body.iter(qn("w:bookmarkStart"))]
    scr = [n for n in names if n and n.startswith("scr_")]
    assert scr == ["scr_000"]        # vendor-onboarding's one SC item
    text = "\n".join(par.text for par in doc.paragraphs)
    assert "Contact: Jane Junior" in text
    cat = doc.core_properties.category
    assert cat.startswith("cw-screens:")
    doc_id = cat.split(":", 1)[1]
    smap = json.loads((area / "_review" / ".maps" / f"{doc_id}.json")
                      .read_text(encoding="utf-8"))
    assert smap["schema"] == "consult-screens-map/v1"
    assert smap["entries"]["scr_000"] == {
        "slug": "vendor-onboarding", "local": "SC-01", "disp": "SC-01"}


def test_build_kits_regenerates_from_scratch(tmp_path):
    """Rebuilding removes stale kit content — kits are derived artifacts."""
    area = make_area(tmp_path)
    out = area / "_review" / "kits"
    stale = out / "old-person"
    stale.mkdir(parents=True)
    (stale / "junk.txt").write_text("x", encoding="utf-8")
    kits.build_kits(str(area))
    assert not stale.exists()
    assert (out / "index.md").is_file()


def test_build_kits_requires_manifest(tmp_path):
    """An area without manifest.json is rejected."""
    d = tmp_path / "noarea"
    d.mkdir()
    with pytest.raises(SystemExit, match="no manifest.json"):
        kits.build_kits(str(d))


# --------------------------------------------------------------------------- #
# v1.18.2 — filename budget: long slugs shorten to whole words in kit paths
# --------------------------------------------------------------------------- #

def test_long_slug_shortened_in_kit_paths(tmp_path):
    """Kit folder/file names carry the slug twice; a long slug is cut to
    whole hyphen words (~24 chars) so OneDrive-deep paths stay under Word's
    ~255-char full-path ceiling. Display labels and IDs are untouched."""
    area = make_area(tmp_path)
    man = json.loads((area / "manifest.json").read_text(encoding="utf-8"))
    long_slug = "vendor-master-data-maintenance-and-banking-change"
    for c in man["components"]:
        if c.get("slug") == "vendor-onboarding":
            c["slug"] = long_slug
            break
    (area / "manifest.json").write_text(json.dumps(man, indent=1),
                                        encoding="utf-8")
    # The gap-log's [[#vendor-onboarding]] token must follow the rename.
    gl = area / "90_appendix-b-gaps.md"
    gl.write_text(gl.read_text(encoding="utf-8").replace(
        "vendor-onboarding", long_slug), encoding="utf-8")
    assert kits.build_kits(str(area)) == 0
    out = area / "_review" / "kits"
    dirs = sorted(d.name for d in out.iterdir() if d.is_dir())
    key = "1-1_vendor-master-data"
    assert key in dirs, dirs
    kdir = out / key
    assert (kdir / "1.1_vendor-master-data.docx").is_file()
    files = {f.name for f in kdir.iterdir()}
    assert not any(long_slug in f for f in files)
