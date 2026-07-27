"""Tests for the M14 document profile — docs/M14-document-profile.md.

One test (or one parametrized family) per Acceptance bullet, plus the profile
validator's own fail-loud rules. Every fixture is a synthetic engagement under
tmp_path: a `components/` dir with an optional `_client/profile.yaml` and one
area folder inside it, so `client_config.load` resolves the profile exactly as
it does in a real engagement.

The two enforcement points are exercised through their real entry points —
`scaffold.confirm` for the skeletons and the manifest, `render.render_folder`
for the document — because "byte-identical to today" is a claim about the
artifacts, not about a helper's return value.
"""
import json
from pathlib import Path

import pytest
import yaml
from docx import Document

import aggregate
import client_config
import doc_model
import orchestrate
import reconcile
import render
import scaffold
from client_config import ProfileError


# --------------------------------------------------------------------------- #
# Fixtures: an engagement, a profile, a scaffolded area, a drafted area
# --------------------------------------------------------------------------- #

TAXONOMY = {"taxonomy": {"categories": [
    {"slug": "finance", "subcategories": [{"slug": "cash"}, {"slug": "payments"}]},
]}}

PROPOSED_PROCS = {"procedures": [
    {"slug": "bank-rec", "title": "Bank Reconciliation", "l2": "cash"},
    {"slug": "payment-run", "title": "Payment Run", "l2": "payments"},
]}

SYSTEMS_YAML = "systems:\n  - slug: netsuite\n    name: NetSuite\n"
ROLES_YAML = "roles:\n  - slug: controller\n    name: Controller\n"
SOURCES_YAML = ("sources:\n- id: SRC-001\n  file: _sources/processed/i.md\n"
                "  touches:\n  - bank-rec\n  - payment-run\n  state: processed\n")

DEP_MARKER = "<!-- derived: dependencies; writer: agent -->"


def engagement(tmp_path) -> Path:
    """`components/` — the engagement root `_client/` resolution walks up to."""
    components = tmp_path / "components"
    components.mkdir(parents=True, exist_ok=True)
    return components


def write_profile(components: Path, **fields) -> Path:
    """Write `components/_client/profile.yaml` with `profile:` at top level."""
    client = components / "_client"
    client.mkdir(parents=True, exist_ok=True)
    path = client / "profile.yaml"
    path.write_text(yaml.safe_dump({"profile": fields}, sort_keys=False),
                    encoding="utf-8")
    return path


def scaffold_area(components: Path, name="treasury"):
    """Run the real confirm gate on a fresh area inside `components`."""
    area = components / name
    proposed = area / "_reference" / ".proposed"
    proposed.mkdir(parents=True, exist_ok=True)
    (proposed / "procedures.yaml").write_text(yaml.safe_dump(PROPOSED_PROCS),
                                              encoding="utf-8")
    (proposed / "systems.yaml").write_text(SYSTEMS_YAML, encoding="utf-8")
    taxonomy = components.parent / "taxonomy.yaml"
    taxonomy.write_text(yaml.safe_dump(TAXONOMY), encoding="utf-8")
    scaffold.confirm(area, "finance", taxonomy, None, None)
    return area


def manifest_of(area: Path) -> dict:
    return json.loads((area / "manifest.json").read_text(encoding="utf-8"))


def derived_kinds(area: Path) -> list[str]:
    return [c["derived_kind"] for c in manifest_of(area)["components"]
            if c.get("role") == "derived"]


def fragment(title: str, slug: str, ctrl: str, pp: str) -> str:
    """A drafted procedure (the M16 seven-section model) with a real controls
    control, a real known-issues pain point, and a `[[slug]]` cross-reference
    INSIDE the controls section (bullet: a reference in an omitted section still
    reconciles)."""
    other = "payment-run" if slug == "bank-rec" else "bank-rec"
    return f"""## {title}

### Scope

Covers the monthly reconciliation only. (SRC-001)

### At a Glance

| Field | Value |
|---|---|
| Frequency | Monthly |
| Preparer | Controller |

### Before You Start

- **Bank statement** — Treasury; downloaded for the closed period.

### Procedure

#### Step 1: Export

Export the statement from the portal.

### Outputs & Evidence

- **Output 1:** Signed reconciliation.

### Key Controls

> **CONTROL — {ctrl}:** Controller signs off on the reconciliation; see [[{other}]].
> - **Type:** Detective
> - **Frequency:** Monthly
> - **Owner:** Controller

### Known Issues & Improvement Opportunities

> **PAIN POINT — {pp}:** The statement is exported by hand every month.
> - **Impact:** Two hours per close
> - **Severity:** Medium

```consult-meta
systems:
  - netsuite
roles:
  - controller
```
"""


def drafted_area(components: Path, name="treasury", profile=None) -> Path:
    """A reconcile-clean area with two drafted procedures, both carrying F
    controls and H pain points, and the python-owned registers the profile
    asks for (built by the real aggregator)."""
    area = components / name
    area.mkdir(parents=True, exist_ok=True)
    # By default the area's own resolved profile decides which derived
    # components the manifest lists — exactly what scaffold would have done.
    prof = profile if profile is not None else client_config.profile(area)

    comps = [
        {"file": "10_bank-rec.md", "role": "procedure", "slug": "bank-rec",
         "heading": "Bank Reconciliation", "l2": "cash", "order": 10},
        {"file": "20_payment-run.md", "role": "procedure", "slug": "payment-run",
         "heading": "Payment Run", "l2": "payments", "order": 20},
    ]
    for d in scaffold.profile_derived_files(prof):
        comps.append({"file": d["file"], "role": "derived",
                      "derived_kind": d["kind"], "writer": d["writer"],
                      "heading": d["heading"], "order": d["order"]})
    (area / "manifest.json").write_text(json.dumps(
        {"schema": "consult-mvp-manifest/v1", "area": name, "l1": "finance",
         "l2_order": ["cash", "payments"], "title": "Treasury",
         "subtitle": "Current state", "components": comps}, indent=1),
        encoding="utf-8")

    (area / "10_bank-rec.md").write_text(
        fragment("Bank Reconciliation", "bank-rec", "CTRL-001", "PP-001"),
        encoding="utf-8")
    (area / "20_payment-run.md").write_text(
        fragment("Payment Run", "payment-run", "CTRL-001", "PP-001"),
        encoding="utf-8")
    ref = area / "_reference"
    ref.mkdir(exist_ok=True)
    (ref / "systems.yaml").write_text(SYSTEMS_YAML, encoding="utf-8")
    (ref / "roles.yaml").write_text(ROLES_YAML, encoding="utf-8")
    (ref / "sources.yaml").write_text(SOURCES_YAML, encoding="utf-8")

    # The static + agent-owned files the manifest may name.
    for comp in comps:
        path = area / comp["file"]
        if path.exists() or comp["role"] != "derived":
            continue
        if comp["writer"] == "agent":
            path.write_text(
                f"## {comp['heading']}\n\n<!-- derived: {comp['derived_kind']}; "
                f"writer: agent -->\n\nNothing of note for [[bank-rec]].\n",
                encoding="utf-8")
    assert aggregate.run(str(area)) == 0
    return area


def rendered_text(area: Path, out_name="draft.docx", **kw) -> str:
    """Render the area and return every paragraph + table cell as one string."""
    out = area / out_name
    render.render_folder(area, out, emit_signal=False, **kw)
    doc = Document(str(out))
    parts = [p.text for p in doc.paragraphs]
    for t in doc.tables:
        for row in t.rows:
            for c in row.cells:
                parts.append(c.text)
    return "\n".join(parts)


def register_of(area: Path, filename: str) -> str:
    return (area / filename).read_text(encoding="utf-8")


# --------------------------------------------------------------------------- #
# Bullet 1 — no profile.yaml: manifest, skeletons and render unchanged
# --------------------------------------------------------------------------- #

def test_absent_profile_resolves_to_todays_full_shape(tmp_path):
    """No `profile:` key anywhere: the full A–H set, an EMPTY body_omit,
    today's derived set (no controls register), and `configured` False."""
    prof = client_config.profile(engagement(tmp_path) / "treasury")
    assert prof.configured is False
    assert prof.sections == client_config.ALL_SECTIONS   # M23: slugs
    assert prof.body_omit == []
    assert prof.callouts == client_config.ALL_CALLOUTS
    assert prof.derived == client_config.DEFAULT_DERIVED
    assert "appendix-controls" not in prof.derived
    assert prof.hidden_sections() == set() and prof.dropped_callouts() == set()


def test_absent_profile_scaffolds_byte_identical_manifest_and_skeletons(tmp_path):
    """Bullet 1, scaffold half: an engagement with no profile.yaml and one whose
    profile spells out today's shape produce byte-identical output — and both
    match the shape scaffold produced before M14 (eight derived views, no
    controls register, every section heading — LETTERLESS since M23)."""
    plain = scaffold_area(engagement(tmp_path / "plain"))
    explicit_root = engagement(tmp_path / "explicit")
    write_profile(explicit_root, sections=list(client_config.ALL_SECTIONS), body_omit=[])
    explicit = scaffold_area(explicit_root)

    assert ((plain / "manifest.json").read_bytes()
            == (explicit / "manifest.json").read_bytes())
    assert ((plain / "10_bank-rec.md").read_bytes()
            == (explicit / "10_bank-rec.md").read_bytes())
    skeleton = (plain / "10_bank-rec.md").read_text(encoding="utf-8")
    for title in doc_model.SECTION_TITLES.values():
        assert f"### {title}" in skeleton
    assert derived_kinds(plain) == client_config.DEFAULT_DERIVED


def test_absent_profile_renders_byte_identical_body(tmp_path):
    """Bullet 1 + bullet 12 (empty/absent body_omit), render half: the
    extracted document text is identical with no profile and with a profile
    that excludes nothing."""
    plain = drafted_area(engagement(tmp_path / "plain"))
    explicit_root = engagement(tmp_path / "explicit")
    write_profile(explicit_root, sections=list(client_config.ALL_SECTIONS), body_omit=[])
    explicit = drafted_area(explicit_root)

    assert rendered_text(plain) == rendered_text(explicit)


# --------------------------------------------------------------------------- #
# Bullet 2 — a profile without F
# --------------------------------------------------------------------------- #

def test_profile_without_f_scaffolds_no_f_section(tmp_path):
    """Enforcement point 1: the skeleton has no Key Controls heading and every
    other section survives, sentinel and end matter intact."""
    root = engagement(tmp_path)
    write_profile(root, sections=[s for s in client_config.ALL_SECTIONS if s != "controls"])
    area = scaffold_area(root)

    skeleton = (area / "10_bank-rec.md").read_text(encoding="utf-8")
    assert "Key Controls" not in skeleton
    for slug, title in doc_model.SECTION_TITLES.items():
        if slug != "controls":
            assert f"### {title}" in skeleton
    assert "<!-- unfilled -->" in skeleton
    assert "consult-meta" in skeleton


def test_profile_without_f_renders_no_f_and_reconciles_clean(tmp_path):
    """Enforcement point 2, on fragments that DO carry F (the profile changed
    after drafting): the render shows no F section and no control, while
    reconcile — which reads fragments, not renders — stays clean."""
    root = engagement(tmp_path)
    write_profile(root, sections=[s for s in client_config.ALL_SECTIONS if s != "controls"])
    area = drafted_area(root)

    text = rendered_text(area)
    assert "F. Key Controls" not in text
    assert "CTRL-" not in text
    assert "Procedure" in text        # the rest is untouched

    orchestrate.emit_aggregate(str(area), warnings=[])
    assert reconcile.reconcile(area) == 0


# --------------------------------------------------------------------------- #
# Bullet 3 — adding a section after drafting quotes the migration
# --------------------------------------------------------------------------- #

def _walk_to_steady_state(area: Path):
    """Aggregate + a clean reconcile: the free stages below the cost gates."""
    orchestrate.emit_aggregate(str(area), warnings=[])
    orchestrate.emit_reconcile(str(area), clean=True, failing_files=[])


def test_adding_a_section_after_drafting_returns_reprofile_with_the_count(tmp_path):
    """Bullet 3: fragments drafted WITHOUT F, then F added to the profile — the
    advisor returns `reprofile`, names every procedure and the dispatch count,
    and gates (it is a cost gate)."""
    root = engagement(tmp_path)
    write_profile(root, sections=[s for s in client_config.ALL_SECTIONS if s != "controls"])
    area = scaffold_area(root)
    for slug in ("bank-rec", "payment-run"):
        path = area / f"10_{slug}.md"
        path.write_text(path.read_text(encoding="utf-8")
                        .replace("<!-- unfilled -->", "Drafted. (SRC-001)"),
                        encoding="utf-8")
    assert orchestrate.decide(str(area))["action"] != "reprofile"

    write_profile(root, sections=list(client_config.ALL_SECTIONS))
    d = orchestrate.decide(str(area))
    assert d["action"] == "reprofile"
    assert d["human_gate"] is True
    assert d["details"]["dispatches"] == 2
    assert d["details"]["missing"] == {"bank-rec": ["controls"],
                                      "payment-run": ["controls"]}
    assert "controls" in d["reason"]


def test_reprofile_clears_once_the_drafters_write_the_heading(tmp_path):
    """Bullet 3, second half: the guard clears when the heading appears — even
    when the answer is "none identified", which is what makes it terminate."""
    root = engagement(tmp_path)
    write_profile(root, sections=[s for s in client_config.ALL_SECTIONS if s != "controls"])
    area = scaffold_area(root)
    for slug in ("bank-rec", "payment-run"):
        path = area / f"10_{slug}.md"
        path.write_text(path.read_text(encoding="utf-8")
                        .replace("<!-- unfilled -->", "Drafted. (SRC-001)"),
                        encoding="utf-8")
    write_profile(root, sections=list(client_config.ALL_SECTIONS))
    assert orchestrate.decide(str(area))["action"] == "reprofile"

    # the drafter adds the heading to ONE procedure: partial migration is fine
    path = area / "10_bank-rec.md"
    path.write_text(path.read_text(encoding="utf-8")
                    + "\n### F. Key Controls\n\nNo controls identified in the "
                      "current state.\n", encoding="utf-8")
    d = orchestrate.decide(str(area))
    assert d["action"] == "reprofile"
    assert d["details"]["dispatches"] == 1
    assert list(d["details"]["missing"]) == ["payment-run"]

    path = area / "10_payment-run.md"
    path.write_text(path.read_text(encoding="utf-8")
                    + "\n### F. Key Controls\n\nNo controls identified in the "
                      "current state.\n", encoding="utf-8")
    assert orchestrate.decide(str(area))["action"] != "reprofile"


# --------------------------------------------------------------------------- #
# Bullet 4 — removing a section is free and reversible
# --------------------------------------------------------------------------- #

def test_removing_a_section_never_dispatches_and_re_adding_does_not_either(tmp_path):
    """Bullet 4: drop F from the profile after drafting — no `reprofile`, the
    render omits F, the fragments are untouched, and re-adding F does not
    re-dispatch, because the headings never left."""
    root = engagement(tmp_path)
    area = drafted_area(root)
    before = (area / "10_bank-rec.md").read_bytes()
    _walk_to_steady_state(area)

    write_profile(root, sections=[s for s in client_config.ALL_SECTIONS if s != "controls"])
    assert orchestrate.decide(str(area))["action"] != "reprofile"
    assert "F. Key Controls" not in rendered_text(area, "without-f.docx")
    assert (area / "10_bank-rec.md").read_bytes() == before

    write_profile(root, sections=list(client_config.ALL_SECTIONS))
    assert orchestrate.decide(str(area))["action"] != "reprofile"
    assert "F. Key Controls" in rendered_text(area, "with-f.docx")
    assert (area / "10_bank-rec.md").read_bytes() == before


# --------------------------------------------------------------------------- #
# Bullet 5 — `derived` without `raci`
# --------------------------------------------------------------------------- #

def test_derived_without_raci_drops_the_component_and_the_dispatch(tmp_path):
    """Bullet 5: no RACI component in the manifest, and `synthesize` never
    names the raci kind — absent from the manifest is absent from the signal."""
    root = engagement(tmp_path)
    write_profile(root, derived=[k for k in client_config.DEFAULT_DERIVED
                                if k != "raci"])
    area = scaffold_area(root)

    assert "raci" not in derived_kinds(area)
    assert not (area / "84_raci.md").exists()
    assert orchestrate.AreaState(str(area)).agent_derived_kinds == ["dependencies"]

    drafted = drafted_area(engagement(tmp_path / "drafted"),
                           profile=client_config.profile(area))
    _walk_to_steady_state(drafted)
    orchestrate.accept_draft(str(drafted))
    d = orchestrate.decide(str(drafted))
    assert d["action"] == "synthesize"
    assert d["details"]["stale_kinds"] == ["dependencies"]
    assert "raci" not in json.dumps(d["details"])


# --------------------------------------------------------------------------- #
# Bullet 6 — `body_omit: [H]`
# --------------------------------------------------------------------------- #

def test_body_omit_h_hides_the_section_but_not_appendix_a(tmp_path):
    """Bullet 6: the rendered procedure has no H, Appendix A is unchanged and
    non-empty, the fragments still carry their pain-point callouts, and no
    drafter is dispatched."""
    plain = drafted_area(engagement(tmp_path / "plain"))
    root = engagement(tmp_path / "omitted")
    write_profile(root, body_omit=["H"])
    area = drafted_area(root)

    text = rendered_text(area)
    assert "H. Known Issues" not in text          # the SECTION is gone...
    assert "PAIN POINT" not in text               # ...and so is its callout

    # Appendix A is byte-identical to the un-omitted engagement's, and real
    assert (register_of(area, "88_appendix-a.md")
            == register_of(plain, "88_appendix-a.md"))
    assert "PP-01" in register_of(area, "88_appendix-a.md")
    assert "No pain points recorded" not in register_of(area, "88_appendix-a.md")
    # ...and the appendix still reaches the reader
    assert "Pain Points" in text and "PP-01" in text

    # the fragment keeps its callout, and nothing is dispatched
    assert "PAIN POINT — PP-001" in (area / "10_bank-rec.md").read_text(
        encoding="utf-8")
    _walk_to_steady_state(area)
    assert orchestrate.decide(str(area))["action"] != "reprofile"


def test_removing_the_body_omit_entry_restores_h_with_no_redraft(tmp_path):
    """Bullet 6, second half: delete the entry and H is back, from the text
    that was always in the fragment — no drafter, no fragment write."""
    root = engagement(tmp_path)
    write_profile(root, body_omit=["H"])
    area = drafted_area(root)
    before = (area / "10_bank-rec.md").read_bytes()
    assert "G. Known Issues" not in rendered_text(area, "omitted.docx")

    write_profile(root, body_omit=[])
    text = rendered_text(area, "restored.docx")
    assert "G. Known Issues" in text
    assert "The statement is exported by hand every month." in text
    assert (area / "10_bank-rec.md").read_bytes() == before
    _walk_to_steady_state(area)
    assert orchestrate.decide(str(area))["action"] != "reprofile"


# --------------------------------------------------------------------------- #
# Bullet 7 — `body_omit: [F]` with the controls register
# --------------------------------------------------------------------------- #

def test_body_omit_f_with_controls_register_collects_every_control(tmp_path):
    """Bullet 7: no F in the procedure body, every `CTRL-` id present exactly
    once in the register, and each row naming its owning procedure."""
    root = engagement(tmp_path)
    write_profile(root, body_omit=["F"],
                  derived=client_config.DEFAULT_DERIVED + ["appendix-controls"])
    area = drafted_area(root)

    register = register_of(area, "89_appendix-controls.md")
    assert "<!-- derived: appendix-controls; writer: python -->" in register
    # M23: the caption names the section by TITLE (display), never by letter.
    assert "aggregated mechanically from the Key Controls section" in register
    # two procedures, one control each -> the two document-global display ids,
    # each exactly once, each row carrying its procedure's [[#slug]] token
    assert register.count("CTRL-01") == 1 and register.count("CTRL-02") == 1
    assert "CTRL-01 ([[#bank-rec]])" in register
    assert "CTRL-02 ([[#payment-run]])" in register
    assert "No controls recorded" not in register

    text = rendered_text(area)
    assert "F. Key Controls" not in text       # ...in the procedure body
    assert "CONTROL —" not in text             # ...nor the callout itself
    assert "Key Controls Register" in text     # ...but the register is there
    assert "CTRL-01" in text and "CTRL-02" in text
    # the register names the owning procedure by resolved display number
    assert "1.1" in text and "2.1" in text

    orchestrate.emit_aggregate(str(area), warnings=[])
    assert reconcile.reconcile(area) == 0


def test_controls_register_is_empty_but_present_with_no_controls(tmp_path):
    """The register's own generated header states its dependency, so an area
    with no CONTROL callouts says so rather than rendering a bare table."""
    root = engagement(tmp_path)
    write_profile(root, derived=client_config.DEFAULT_DERIVED
                  + ["appendix-controls"])
    area = drafted_area(root)
    for slug, fname in (("bank-rec", "10_bank-rec.md"),
                        ("payment-run", "20_payment-run.md")):
        path = area / fname
        text = path.read_text(encoding="utf-8")
        head, _, tail = text.partition("### Key Controls")
        path.write_text(head + "### Known Issues"
                        + tail.split("### Known Issues", 1)[1],
                        encoding="utf-8")
    assert aggregate.run(str(area)) == 0
    assert "_No controls recorded._" in register_of(area,
                                                    "89_appendix-controls.md")


# --------------------------------------------------------------------------- #
# Bullets 8 + 9 — the two named profile errors
# --------------------------------------------------------------------------- #

def test_body_omit_f_without_controls_register_is_a_named_error(tmp_path):
    """Bullet 8: hiding F with no destination would make the controls vanish
    from the document, so the profile is refused — naming both halves."""
    root = engagement(tmp_path)
    write_profile(root, body_omit=["F"])
    with pytest.raises(ProfileError) as exc:
        client_config.profile(root / "treasury")
    message = str(exc.value)
    assert "body_omit" in message and "appendix-controls" in message
    assert "vanish" in message


def test_body_omit_naming_a_section_absent_from_sections_is_a_named_error(tmp_path):
    """Bullet 9: omitting something that does not exist is a profile error, not
    a no-op — and the message names the section and the section set."""
    root = engagement(tmp_path)
    write_profile(root, sections=[s for s in "ABCDEFGH" if s != "H"],
                  body_omit=["H"])
    with pytest.raises(ProfileError) as exc:
        client_config.profile(root / "treasury")
    message = str(exc.value)
    assert "body_omit" in message and "issues" in message
    assert "outputs" in message          # the `sections:` it is absent from


def test_a_profile_error_stops_the_confirm_before_it_touches_the_folder(tmp_path):
    """Fail-loud in practice: scaffold resolves the profile before promoting
    anything, so a bad profile leaves `.proposed/` intact and writes no
    manifest — the human fixes the YAML and re-confirms."""
    root = engagement(tmp_path)
    write_profile(root, body_omit=["F"])
    area = root / "treasury"
    proposed = area / "_reference" / ".proposed"
    proposed.mkdir(parents=True)
    (proposed / "procedures.yaml").write_text(yaml.safe_dump(PROPOSED_PROCS),
                                              encoding="utf-8")
    taxonomy = tmp_path / "taxonomy.yaml"
    taxonomy.write_text(yaml.safe_dump(TAXONOMY), encoding="utf-8")

    with pytest.raises(ProfileError):
        scaffold.confirm(area, "finance", taxonomy, None, None)
    assert not (area / "manifest.json").exists()
    assert (proposed / "procedures.yaml").is_file()


# --------------------------------------------------------------------------- #
# Manifest drift — the trap the profile validator cannot see
#
# The validator vouches for the PROFILE (`body_omit: [controls]` requires the
# register in `derived:`), but render builds from the MANIFEST, written at the
# confirm gate. A profile that acquires the omission AFTER scaffolding leaves
# a manifest with no register component: render must refuse (or every control
# silently leaves the document), and `scaffold.sync_profile` is the supported
# proposal-free path that reconciles the manifest.
# --------------------------------------------------------------------------- #

def drifted_area(root):
    """An area scaffolded BEFORE the profile existed (default derived set, no
    controls register), whose engagement then adopted `body_omit: [F]` with
    the register in `derived:` — a profile the validator accepts."""
    area = drafted_area(root, profile=client_config.Profile())
    write_profile(root, body_omit=["F"],
                  derived=client_config.DEFAULT_DERIVED + ["appendix-controls"])
    return area


def test_render_refuses_a_manifest_missing_the_controls_register(tmp_path):
    """The valid-profile / stale-manifest combination fails the render loudly,
    naming the register, the loss, and the sync command — never a document
    with the controls silently gone."""
    root = engagement(tmp_path)
    area = drifted_area(root)
    assert "appendix-controls" not in derived_kinds(area)

    with pytest.raises(SystemExit) as exc:
        render.render_folder(area, area / "draft.docx", emit_signal=False)
    message = str(exc.value)
    assert "appendix-controls" in message and "drop them" in message
    assert "--sync-profile" in message
    assert not (area / "draft.docx").exists()


def test_sync_profile_adds_the_register_and_render_recovers(tmp_path):
    """`sync_profile` adds the register component + stub without a `.proposed/`
    round; after aggregate the render succeeds and collects every control —
    the same end state a fresh confirm would have produced."""
    root = engagement(tmp_path)
    area = drifted_area(root)

    assert scaffold.sync_profile(area) == 0
    assert "appendix-controls" in derived_kinds(area)
    assert (area / "89_appendix-controls.md").is_file()
    # non-derived components survived byte-for-byte
    assert [c["slug"] for c in manifest_of(area)["components"]
            if c.get("role") == "procedure"] == ["bank-rec", "payment-run"]

    assert aggregate.run(str(area)) == 0
    text = rendered_text(area)
    assert "CONTROL —" not in text             # still out of the body...
    assert "Key Controls Register" in text     # ...but the register caught them
    assert "CTRL-01" in text and "CTRL-02" in text


def test_sync_profile_is_idempotent_and_only_touches_the_derived_set(tmp_path):
    """A second sync is a byte-level no-op; dropping a kind from `derived:`
    removes only the manifest entry — the file stays on disk, exactly the
    removal semantics a fresh confirm has."""
    root = engagement(tmp_path)
    area = drifted_area(root)
    assert scaffold.sync_profile(area) == 0
    before = (area / "manifest.json").read_bytes()
    assert scaffold.sync_profile(area) == 0
    assert (area / "manifest.json").read_bytes() == before

    write_profile(root, body_omit=["F"],
                  derived=[k for k in client_config.DEFAULT_DERIVED
                           if k != "raci"] + ["appendix-controls"])
    assert scaffold.sync_profile(area) == 0
    assert "raci" not in derived_kinds(area)
    assert (area / "84_raci.md").is_file()     # scaffold never deletes


def test_controls_homeless_render_reports_dangling_control_references(tmp_path):
    """A profile may deliberately drop `controls` from `sections:` (the "this
    document carries no controls" shape — sanctioned, unlike the body_omit
    trap). A prose mention of a control id then dangles for the reader, in
    EVERY mode, so render counts and enumerates it like final-mode gap refs."""
    root = engagement(tmp_path)
    write_profile(root, sections=[s for s in client_config.ALL_SECTIONS
                                  if s != "controls"])
    area = drafted_area(root)
    frag = area / "10_bank-rec.md"
    frag.write_text(frag.read_text(encoding="utf-8").replace(
        "Covers the monthly reconciliation only.",
        "Covers the monthly reconciliation only; constrained by CTRL-001."),
        encoding="utf-8")

    stats = render.render_folder(area, area / "d.docx", emit_signal=False)
    assert stats["dangling_refs"] == {"bank-rec": ["CTRL-01"]}
    assert stats["dangling_ref_count"] == 1
    text = rendered_text(area, "d2.docx")
    assert "Key Controls" not in text          # the section really is gone
    assert "CTRL-01" in text                   # ...and the prose ref dangles


def test_controls_in_the_register_do_not_dangle(tmp_path):
    """The register shape (body_omit + appendix-controls) gives every control
    id a home the reader can look up, so the same prose reference is NOT
    reported."""
    root = engagement(tmp_path)
    write_profile(root, body_omit=["F"],
                  derived=client_config.DEFAULT_DERIVED + ["appendix-controls"])
    area = drafted_area(root)
    frag = area / "10_bank-rec.md"
    frag.write_text(frag.read_text(encoding="utf-8").replace(
        "Covers the monthly reconciliation only.",
        "Covers the monthly reconciliation only; constrained by CTRL-001."),
        encoding="utf-8")

    stats = render.render_folder(area, area / "d.docx", emit_signal=False)
    assert stats["dangling_refs"] == {}
    assert stats["dangling_ref_count"] == 0


def test_body_omit_issues_without_appendix_a_is_a_named_error(tmp_path):
    """The cross-field rule covers the pain-point register too: hiding
    `issues` from the body while removing `appendix-a` from `derived:` would
    make every PAIN POINT / IMPROVEMENT OPPORTUNITY vanish — refused, named,
    exactly like the controls rule."""
    root = engagement(tmp_path)
    write_profile(root, body_omit=["issues"],
                  derived=[k for k in client_config.DEFAULT_DERIVED
                           if k != "appendix-a"])
    with pytest.raises(ProfileError) as exc:
        client_config.profile(root / "treasury")
    message = str(exc.value)
    assert "body_omit" in message and "appendix-a" in message
    assert "vanish" in message and "PAIN POINT" in message


def test_render_refuses_a_manifest_missing_the_pain_point_register(tmp_path):
    """Manifest drift, issues flavor: the manifest was built without
    appendix-a (a profile that dropped it), then the profile switched to
    body_omit issues WITH the register — valid profile, stale manifest.
    Render refuses and names --sync-profile, exactly like the controls case."""
    root = engagement(tmp_path)
    area = drafted_area(root, profile=client_config.Profile(
        derived=[k for k in client_config.DEFAULT_DERIVED
                 if k != "appendix-a"]))
    write_profile(root, body_omit=["issues"],
                  derived=client_config.DEFAULT_DERIVED)
    assert "appendix-a" not in derived_kinds(area)

    with pytest.raises(SystemExit) as exc:
        render.render_folder(area, area / "d.docx", emit_signal=False)
    message = str(exc.value)
    assert "appendix-a" in message and "--sync-profile" in message

    assert scaffold.sync_profile(area) == 0
    assert "appendix-a" in derived_kinds(area)
    assert aggregate.run(str(area)) == 0
    render.render_folder(area, area / "d.docx", emit_signal=False)


def test_issues_homeless_render_reports_dangling_pp_references(tmp_path):
    """`issues` dropped from `sections:` with no appendix-a register: a prose
    mention of a pain-point id dangles and is reported, same as controls."""
    root = engagement(tmp_path)
    write_profile(root,
                  sections=[s for s in client_config.ALL_SECTIONS
                            if s != "issues"],
                  derived=[k for k in client_config.DEFAULT_DERIVED
                           if k != "appendix-a"])
    area = drafted_area(root)
    frag = area / "10_bank-rec.md"
    frag.write_text(frag.read_text(encoding="utf-8").replace(
        "Covers the monthly reconciliation only.",
        "Covers the monthly reconciliation only; the export burden is "
        "recorded as PP-001."), encoding="utf-8")

    stats = render.render_folder(area, area / "d.docx", emit_signal=False)
    assert stats["dangling_refs"] == {"bank-rec": ["PP-01"]}
    assert stats["dangling_ref_count"] == 1


def test_sync_profile_refuses_an_unscaffolded_area(tmp_path):
    """A fresh area has no manifest to reconcile — that is the confirm gate's
    job, and the refusal says so."""
    root = engagement(tmp_path)
    area = root / "treasury"
    area.mkdir()
    with pytest.raises(SystemExit) as exc:
        scaffold.sync_profile(area)
    assert "confirm" in str(exc.value)


# --------------------------------------------------------------------------- #
# Bullet 10 — `body_omit` never causes `reprofile`
# --------------------------------------------------------------------------- #

@pytest.mark.parametrize("omit", [["H"], ["F"], ["F", "H"]])
def test_body_omit_never_fires_reprofile(tmp_path, omit):
    """Bullet 10: the headings are present, so the drift detector has nothing
    to detect — `body_omit` and the migration guard do not interact."""
    root = engagement(tmp_path)
    write_profile(root, body_omit=omit,
                  derived=client_config.DEFAULT_DERIVED + ["appendix-controls"])
    area = drafted_area(root)
    assert orchestrate.AreaState(str(area)).profile_drift() == {}
    _walk_to_steady_state(area)
    assert orchestrate.decide(str(area))["action"] != "reprofile"


# --------------------------------------------------------------------------- #
# Bullet 11 — a cross-reference inside an omitted section still reconciles
# --------------------------------------------------------------------------- #

def test_xref_inside_an_omitted_section_still_passes_reconcile(tmp_path):
    """Bullet 11: reconcile validates `[[slug]]` against fragment SOURCES, not
    against the rendered output, so a reference from an omitted section
    resolves — verified with the reference living inside the omitted F."""
    root = engagement(tmp_path)
    write_profile(root, body_omit=["F"],
                  derived=client_config.DEFAULT_DERIVED + ["appendix-controls"])
    area = drafted_area(root)
    assert "see [[payment-run]]" in (area / "10_bank-rec.md").read_text(
        encoding="utf-8")

    rendered_text(area)                      # the render omits that section...
    orchestrate.emit_aggregate(str(area), warnings=[])
    assert reconcile.reconcile(area) == 0    # ...and reconcile is still clean


# --------------------------------------------------------------------------- #
# The profile itself: resolution, shadowing, reporting, and the validator
# --------------------------------------------------------------------------- #

def test_area_profile_shadows_the_engagement_profile_whole(tmp_path):
    """`profile:` is one top-level key, so M13's per-key rule applies: an
    area's own profile.yaml wins entirely, with provenance reported."""
    root = engagement(tmp_path)
    write_profile(root, sections=list(client_config.ALL_SECTIONS), body_omit=["H"])
    area = root / "treasury"
    (area / "_client").mkdir(parents=True)
    (area / "_client" / "profile.yaml").write_text(
        yaml.safe_dump({"profile": {"sections": list("ABE")}}), encoding="utf-8")

    prof = client_config.profile(area)
    assert prof.sections == ["scope", "quick-reference", "steps"]
    assert prof.body_omit == []              # shadowed WHOLE, not merged
    assert prof.layer == "area"
    assert prof.report_line().startswith("document profile: _client/ (area)")


def test_report_line_states_the_shape_when_a_profile_is_in_play(tmp_path):
    """The one-line report every stage prints: layer first (M13's shape), then
    the sections and body_omit a surprised reader needs to see."""
    root = engagement(tmp_path)
    assert client_config.profile(root / "t").report_line() == (
        "document profile: none (full A–G, nothing omitted)")

    write_profile(root, sections=[s for s in client_config.ALL_SECTIONS
                                  if s != "outputs"], body_omit=["H"])
    line = client_config.profile(root / "t").report_line()
    assert line == ("document profile: _client/ (engagement) — sections "
                    "A=scope B=quick-reference C=before-you-start D=steps "
                    "E=controls F=issues, body_omit issues")


def test_sections_are_canonicalized_to_slugs_in_the_profiles_own_order(tmp_path):
    """M23 replaces M14's "reordering is out of scope": section tokens are
    canonicalized to SLUGS, but the profile's ORDER is kept, because that order
    is what assigns the letters."""
    root = engagement(tmp_path)
    write_profile(root, sections=["E", "A", "H", "B"])
    prof = client_config.profile(root / "t")
    assert prof.sections == ["steps", "scope", "issues", "quick-reference"]
    assert prof.letters() == {"steps": "A", "scope": "B",
                              "issues": "C", "quick-reference": "D"}


def test_derived_shorthand_names_canonicalize_to_manifest_kinds(tmp_path):
    """The spec writes the derived set in reader shorthand; the manifest speaks
    `derived_kind`. Both are accepted and mean one thing."""
    root = engagement(tmp_path)
    write_profile(root, derived=["index", "roles", "systems", "dependencies",
                                 "raci", "appendix-a", "appendix-b",
                                 "appendix-c", "appendix-controls"])
    prof = client_config.profile(root / "t")
    assert prof.derived == client_config.ALL_DERIVED
    assert prof.wants("procedure-index") and prof.wants("gap-log")


VALIDATOR_CASES = [
    ("a section that is not a section",
     {"sections": ["A", "B", "E", "Z"]}, "is not a section"),
    ("a mandatory section dropped",
     {"sections": ["A", "B", "C"]}, "mandatory section"),
    ("an unknown callout kind",
     {"callouts": ["CONTROL", "VIBE CHECK"]}, "unknown callout kind"),
    ("an unknown derived component",
     {"derived": ["index", "appendix-z"]}, "unknown component"),
    ("an unknown profile field",
     {"sectionz": ["A"]}, "unknown profile field"),
    ("a non-mapping profile",
     None, "must be a mapping"),
]


@pytest.mark.parametrize("label,fields,expected",
                         [c for c in VALIDATOR_CASES],
                         ids=[c[0] for c in VALIDATOR_CASES])
def test_the_validator_names_what_is_wrong(label, fields, expected):
    """Every profile defect fails loud and NAMED: the reader is hand-editing a
    five-line YAML file, and the failure mode this replaces is a silently
    reshaped client deliverable."""
    raw = ["not", "a", "mapping"] if fields is None else fields
    with pytest.raises(ProfileError) as exc:
        client_config.parse_profile(raw, "_client/profile.yaml")
    assert expected in str(exc.value)


def test_a_second_same_layer_file_claiming_profile_is_already_fatal(tmp_path):
    """M13 gives M14 this for free: two files in one `_client/` layer both
    claiming `profile:` stops the run rather than picking a winner."""
    root = engagement(tmp_path)
    write_profile(root, sections=list(client_config.ALL_SECTIONS))
    (root / "_client" / "shape.yaml").write_text(
        yaml.safe_dump({"profile": {"sections": list("ABE")}}), encoding="utf-8")
    with pytest.raises(client_config.ClientConfigError) as exc:
        client_config.profile(root / "treasury")
    assert "duplicate top-level key 'profile'" in str(exc.value)


# --------------------------------------------------------------------------- #
# The advisor stays pure, and the guard sits where the ticket puts it
# --------------------------------------------------------------------------- #

def test_reprofile_is_a_pure_read_that_writes_nothing(tmp_path):
    """Advisor purity (M7/M18): the new guard adds no state file and no write —
    same answer twice, not one byte changed."""
    root = engagement(tmp_path)
    write_profile(root, sections=[s for s in client_config.ALL_SECTIONS if s != "controls"])
    area = scaffold_area(root)
    for slug in ("bank-rec", "payment-run"):
        path = area / f"10_{slug}.md"
        path.write_text(path.read_text(encoding="utf-8")
                        .replace("<!-- unfilled -->", "Drafted. (SRC-001)"),
                        encoding="utf-8")
    write_profile(root, sections=list(client_config.ALL_SECTIONS))

    def snapshot():
        return {p.relative_to(area): p.read_bytes()
                for p in sorted(area.rglob("*")) if p.is_file()}

    before = snapshot()
    first = orchestrate.decide(str(area))
    second = orchestrate.decide(str(area))
    assert first == second == {**first}
    assert first["action"] == "reprofile"
    assert snapshot() == before


def test_reprofile_sits_below_fill_and_above_aggregate(tmp_path):
    """Precedence: skeletons first (an `unfilled` fragment is `fill`, not
    `reprofile`), and the migration is quoted before any derived view is built
    from a shape that is about to change."""
    root = engagement(tmp_path)
    write_profile(root, sections=[s for s in client_config.ALL_SECTIONS if s != "controls"])
    area = scaffold_area(root)
    write_profile(root, sections=list(client_config.ALL_SECTIONS))

    # both fragments still carry the sentinel -> fill outranks reprofile
    assert orchestrate.decide(str(area))["action"] == "fill"

    for slug in ("bank-rec", "payment-run"):
        path = area / f"10_{slug}.md"
        path.write_text(path.read_text(encoding="utf-8")
                        .replace("<!-- unfilled -->", "Drafted. (SRC-001)"),
                        encoding="utf-8")
    # no .aggregate.json yet: guard 6 would fire, but reprofile precedes it
    assert not (area / ".aggregate.json").exists()
    assert orchestrate.decide(str(area))["action"] == "reprofile"


def test_no_profile_means_no_reprofile_however_odd_the_fragment(tmp_path):
    """The guard is inert without a resolved profile: with no `profile:` key
    there is no recorded requirement to drift from, so a fragment with no A–H
    headings at all walks the pre-M14 ladder unchanged."""
    root = engagement(tmp_path)
    area = root / "treasury"
    area.mkdir(parents=True)
    (area / "manifest.json").write_text(json.dumps(
        {"schema": "consult-mvp-manifest/v1", "area": "treasury", "l1": "fin",
         "l2_order": ["cash"], "title": "T", "subtitle": "S",
         "components": [{"file": "10_bank-rec.md", "role": "procedure",
                         "slug": "bank-rec", "heading": "Bank Rec",
                         "l2": "cash", "order": 10}]}), encoding="utf-8")
    (area / "10_bank-rec.md").write_text("## Bank Rec\n\nNo sections at all.\n",
                                         encoding="utf-8")
    assert orchestrate.AreaState(str(area)).profile_drift() == {}
    assert orchestrate.decide(str(area))["action"] == "aggregate"


# --------------------------------------------------------------------------- #
# The strip helpers, directly: line counts and the end-matter boundary
# --------------------------------------------------------------------------- #

def test_render_strips_are_line_count_preserving(tmp_path):
    """Render blanks rather than deletes, so assembled body line k still maps to
    fragment line k — the provenance the M10 review-apply pipeline depends on."""
    text = fragment("Bank Reconciliation", "bank-rec", "CTRL-001", "PP-001")
    prof = client_config.parse_profile({"sections": list("ABCDEFGH"),
                                        "body_omit": ["F", "H"],
                                        "derived": ["appendix-controls",
                                                    "appendix-a"]})
    stripped = render._apply_profile(text, prof)
    assert stripped.count("\n") == text.count("\n")
    assert "Key Controls" not in stripped and "Known Issues" not in stripped
    # end matter belongs to no section, so H's omission cannot take it
    assert "consult-meta" in stripped and "netsuite" in stripped


def test_render_strips_dropped_callout_kinds_only(tmp_path):
    """A callout kind the profile does not carry loses its whole blockquote;
    the kinds it does carry are untouched, and line counts still hold."""
    text = fragment("Bank Reconciliation", "bank-rec", "CTRL-001", "PP-001")
    prof = client_config.parse_profile(
        {"callouts": [c for c in client_config.ALL_CALLOUTS
                      if c != "PAIN POINT"]})
    stripped = render._apply_profile(text, prof)
    assert stripped.count("\n") == text.count("\n")
    assert "PAIN POINT" not in stripped
    assert "Two hours per close" not in stripped     # the sub-fields go too
    assert "CONTROL — CTRL-001" in stripped
    assert "Known Issues" in stripped                # the SECTION stays


def test_scaffold_keeps_the_end_matter_when_the_last_section_goes(tmp_path):
    """`keep_sections` drops whole `### X.` blocks but never the `consult-meta`
    end matter, which belongs to no section."""
    kept = scaffold.keep_sections(
        scaffold._FALLBACK_SKELETON.format(heading="Bank Rec"),
        [s for s in client_config.ALL_SECTIONS if s != "issues"])
    assert "### Known Issues" not in kept and "PAIN POINT" not in kept
    assert "consult-meta" in kept
    assert "\n\n\n" not in kept                      # no gap map left behind


def test_keep_sections_returns_the_text_unchanged_when_nothing_drops():
    """The byte-identity short circuit, directly: a skeleton naming no excluded
    section is returned as-is, not reformatted."""
    text = scaffold._FALLBACK_SKELETON.format(heading="Bank Rec")
    assert scaffold.keep_sections(text, client_config.ALL_SECTIONS) is text
